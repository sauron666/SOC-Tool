import os
import sys
import glob
def set_tcl_tk_env():
    python_dir = os.path.dirname(sys.executable)
    tcl_paths = glob.glob(os.path.join(python_dir, "tcl", "tcl*"))
    tk_paths = glob.glob(os.path.join(python_dir, "tcl", "tk*"))

    if tcl_paths:
        os.environ["TCL_LIBRARY"] = tcl_paths[0]
    if tk_paths:
        os.environ["TK_LIBRARY"] = tk_paths[0]
set_tcl_tk_env()

def get_tkdnd_path():
    """Връща пътя до tkdnd вътре в exe или извън него"""
    if getattr(sys, 'frozen', False):
        # Ако е в .exe – работим с _MEIxxxx папката
        base_path = sys._MEIPASS
    else:
        # При стартиране на .py – използвай текущата директория
        base_path = os.path.dirname(__file__)

    return os.path.join(base_path, "tkdnd2.9")
import os
import sys
import glob
import base64
import configparser
import threading
import concurrent.futures
from threading import Lock, Semaphore
from tkinterdnd2 import TkinterDnD, DND_FILES
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, Menu, simpledialog, scrolledtext
from PIL import Image, ImageTk
import requests
import io
import webbrowser
import subprocess
import csv
import json
import hashlib
import urllib.parse
from datetime import datetime
import socket
import time
from ldap3 import Server, Connection, ALL
#import winreg
from bs4 import BeautifulSoup
import whois
import dns.resolver
import geoip2.database
from ipwhois import IPWhois
import shodan
import ipaddress
import re
from virustotal_python import Virustotal
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from queue import Queue
import queue
from functools import lru_cache
import functools
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
import subprocess
import tkinter.scrolledtext as scrolledtext
from tkinter import filedialog
from tkinterdnd2 import DND_FILES, TkinterDnD
import asyncio
import socket
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
from ldap3 import Server, Connection, ALL
from ldap3 import Server, Connection, ALL, NTLM, SUBTREE, ALL_ATTRIBUTES
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
from datetime import datetime
from pyad import aduser, adquery, pyad
import win32api
import win32con
import win32security
import aiohttp
import ipaddress
from bs4 import BeautifulSoup
import base64
import hashlib
from urllib.parse import quote_plus
from tkinter import ttk, messagebox, filedialog, scrolledtext
from requests.auth import HTTPBasicAuth
import webbrowser
try:
    import extract_msg
except ImportError:
    extract_msg = None

try:
    import pyperclip
except ImportError:
    pyperclip = None

import openai
import os
import re
import email
from email import policy
from email.parser import BytesParser
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
#from openai import OpenAI
#from httpx import Client as HTTPXClient
from requests_negotiate_sspi import HttpNegotiateAuth
import pandas as pd
import docx
import fitz  # PyMuPDF
import httpx
import openai
import certifi

# Константи
CONFIG_FILE = "config.ini"
GEOIP_DB_PATH = "GeoLite2-City.mmdb"
ICON_FILE = "temp_icon.ico"
MAX_WORKERS = 100000  # Увеличаваме максималния брой нишки
MAX_IP_PER_BATCH = 1000000  # Максимален брой IP адреси за пакетна обработка
CACHE_TIMEOUT = 3600  # 5 минути
MAX_SCAN_THREADS = 1000 

class GeoDataSource:
    def __init__(self, config):

        self.ps_history = []
        self.ps_history_index = -1
        self.config = config
        try:
            self.geoip_reader = geoip2.database.Reader(GEOIP_DB_PATH)
        except Exception as e:
            print(f"Грешка при зареждане на GeoIP база: {e}")
            self.geoip_reader = None
    
    def get_geo_data(self, ip, force_source=None):
        """
        Взима гео данни за IP адрес :param ip: IP адрес за проверка :param force_source: опционален параметър за конкретен източник :return: резултати от гео проверка"""
        if force_source:
            return self._get_from_specific_source(force_source, ip)
        
        sources_order = ['abuseipdb', 'geoip', 'ipapi', 'ipgeolocation', 'geojs',]
        
        for source in sources_order:
            result = self._get_from_specific_source(source, ip)
            if result and not result.get('error'):
                return result
        return {"error": "Неуспешно взимане на данни"}
    
    def _get_from_specific_source(self, source_name, ip):
            """Взима данни от конкретен източник"""
            source_methods = {
                'abuseipdb': self._get_abuseipdb_data,
                'geoip': self._get_geoip_data,
                'ipapi': self._get_ipapi_data,
                'ipgeolocation': self._get_ipgeolocation_data,
                'geojs': self._get_geojs_data
            }
            
            if source_name in source_methods:
                return source_methods[source_name](ip)
            return {"error": "Невалиден източник"}
    
    def _get_abuseipdb_data(self, ip):
        if not self.config.get_api_key("ABUSEIP_API_KEY"):
            return None
        
        url = "https://api.abuseipdb.com/api/v2/check"
        headers = {
            "Key": self.config.get_api_key("ABUSEIP_API_KEY"),
            "Accept": "application/json"
        }
        params = {"ipAddress": ip, "maxAgeInDays": "90"}
        
        try:
            response = requests.get(url, headers=headers, params=params, timeout=2)
            if response.status_code == 200:
                data = response.json()["data"]
                return {
                    "country": data.get("countryCode", "N/A"),
                    "isp": data.get("isp", "N/A"),
                    "source": "AbuseIPDB"
                }
            return {"error": f"API грешка: {response.status_code}"}
        except Exception as e:
            return {"error": str(e)}
    
    def _get_ip2location_data(self, ip):
        api_key = self.config.get_api_key("IP2LOCATION_API_KEY")
        if not api_key:
            return None
        try:
            url = f"https://api.ip2location.io/?key={api_key}&ip={ip}&format=json"
            response = requests.get(
                url,
                proxies=self.config.get_proxy_settings(),
                timeout=2
            )
            if response.status_code == 200:
                data = response.json()
                return {
                    "country": data.get("country_name", "N/A"),
                    "isp": data.get("isp", "N/A"),
                    "source": "IP2Location"
                }
            return {"error": f"API грешка: {response.status_code}"}
        except Exception as e:
            return {"error": str(e)}
    
    def _get_geoip_data(self, ip):
        if not self.geoip_reader:
            return None
        
        try:
            response = self.geoip_reader.city(ip)
            return {
                "country": f"{response.country.name} ({response.country.iso_code})",
                "city": response.city.name,
                "source": "GeoIP"
            }
        except Exception as e:
            return {"error": str(e)}
    
    def _get_ipapi_data(self, ip):
        try:
            response = requests.get(
                f"http://ip-api.com/json/{ip}",
                proxies=self.config.get_proxy_settings(),
                timeout=2
            )
            if response.status_code == 200:
                data = response.json()
                return {
                    "country": data.get("country", "N/A"),
                    "isp": data.get("isp", "N/A"),
                    "source": "IP-API"
                }
            return {"error": f"API грешка: {response.status_code}"}
        except Exception as e:
            return {"error": str(e)}
    
    def _get_ipgeolocation_data(self, ip):
        api_key = self.config.get_api_key("GEO_API_KEY")
        if not api_key:
            return None
        
        try:
            response = requests.get(f"https://api.ipgeolocation.io/ipgeo?apiKey={api_key}&ip={ip}",proxies=self.config.get_proxy_settings(), timeout=1)
            if response.status_code == 200:
                data = response.json()
                return {
                    "country": data.get("country_name", "N/A"),
                    "isp": data.get("isp", "N/A"),
                    "source": "IPGeolocation"
                }
            return {"error": f"API грешка: {response.status_code}"}
        except Exception as e:
            return {"error": str(e)}
    
    def _get_geojs_data(self, ip):
        try:
            response = requests.get(
                f"https://get.geojs.io/v1/ip/geo/{ip}.json",
                proxies=self.config.get_proxy_settings(),
                timeout=2
            )
            if response.status_code == 200:
                data = response.json()
                return {
                    "country": data.get("country", "N/A"),
                    "isp": data.get("organization_name", "N/A"),
                    "source": "GeoJS"
                }
            return {"error": f"API грешка: {response.status_code}"}
        except Exception as e:
            return {"error": str(e)}

class ConfigManager:
    """Управление на конфигурациите и API ключове"""
    """Разширен мениджър на конфигурации"""

    def __init__(self):
        self.config = configparser.ConfigParser()
        self.load_config()

    def load_config(self):
        if not os.path.exists(CONFIG_FILE):
            self.create_default_config()
        else:
            self.config.read(CONFIG_FILE)
            created = False
            for section in ["API", "OSINT_API_KEYS", "SETTINGS", "PROXY", "EMAIL", "UI"]:
                if section not in self.config:
                    self.config[section] = {}
                    created = True
            if created:
                with open(CONFIG_FILE, "w") as f:
                    self.config.write(f)

    def create_default_config(self):
        """Създаване на конфигурационен файл по подразбиране"""
        self.config["API"] = {
            "VIRUSTOTAL_API_KEY": "",
            "HYBRIDANALYSIS_API_KEY": "",
            "ABUSEIP_API_KEY": "",
            "ABSTRACT_API_KEY": "",
            "GEO_API_KEY": "",
            "URLSCAN_API_KEY": "",
            "SHODAN_API_KEY": "",
            "MXTOOLBOX_API_KEY": "",
            "IPINFO_API_KEY": "",
            "FINDIP_API_KEY": "",
            "IP2LOCATION_API_KEY": "",
            "MAXMIND_ACCOUNT_ID": "",
            "MAXMIND_LICENSE_KEY": "",
            "PHISHTANK_API_KEY": "",
            "OPENAI_API_KEY": "",

        }

        self.config["SETTINGS"] = {
            "PROXY_ENABLED": "False",
            "THEME": "dark",
            "FONT": "Arial",
            "FONT_SIZE": "10",
            "LANGUAGE": "bg",
            "MAX_THREADS": str(MAX_WORKERS),
            "CACHE_ENABLED": "True",
            "CACHE_TIMEOUT": str(CACHE_TIMEOUT),
        }

        self.config["PROXY"] = {
            "HTTP_PROXY": "",
            "HTTPS_PROXY": "",
            "PROXY_USER": "",
            "PROXY_PASS": ""
        }

        self.config["EMAIL"] = {
            "SMTP_SERVER": "",
            "SMTP_PORT": "587",
            "EMAIL_USER": "",
            "EMAIL_PASSWORD": "",
            "DEFAULT_RECIPIENT": "",
            "ANALYSIS_NOTIFICATION_EMAIL": "",
            "USE_OUTLOOK": "False",
            "OUTLOOK_SIGNATURE": "True"
        }

        self.config["UI"] = {
            "ICON_BASE64": "",
            "LAST_USED_TAB": "0"
        }

        self.config["OSINT_API_KEYS"] = {
            "HUNTER_IO": "",
            "INTELLIGENCE_X": "",
            "GITHUB_TOKEN": "",
            "CENSYS_ID": "",
            "CENSYS_SECRET": "",
            "ONYPHE_KEY": "",
            "USE_INTELLIGENCE_X": "True",
            "BREACHDIRECTORY_KEY": "" 
        }


        with open(CONFIG_FILE, "w") as configfile:
            self.config.write(configfile)

    def get_api_key(self, key_name, fallback=""):
        return self.config.get("API", key_name, fallback=fallback)
    
    def get_osint_key(self, key_name, fallback=""):
        return self.config.get("OSINT_API_KEYS", key_name, fallback=fallback)

    def get_proxy_settings(self):
        if self.config.getboolean("SETTINGS", "PROXY_ENABLED", fallback=False):
            proxy_settings = {
                "http": self.config.get("PROXY", "HTTP_PROXY", fallback=""),
                "https": self.config.get("PROXY", "HTTPS_PROXY", fallback=""),
            }
            
            # Добавяне на proxy аутентикация, ако има
            proxy_user = self.config.get("PROXY", "PROXY_USER", fallback="")
            proxy_pass = self.config.get("PROXY", "PROXY_PASS", fallback="")
            
            if proxy_user and proxy_pass:
                proxy_settings["http"] = proxy_settings["http"].replace(
                    "://", f"://{proxy_user}:{proxy_pass}@"
                )
                proxy_settings["https"] = proxy_settings["https"].replace(
                    "://", f"://{proxy_user}:{proxy_pass}@"
                )
            
            return proxy_settings
        return None

    def get_ui_settings(self):
        """Връща настройките за интерфейса"""
        return {
            "theme": self.config.get("SETTINGS", "THEME", fallback="dark"),
            "font": self.config.get("SETTINGS", "FONT", fallback="Arial"),
            "font_size": self.config.getint("SETTINGS", "FONT_SIZE", fallback=10),
            "language": self.config.get("SETTINGS", "LANGUAGE", fallback="bg"),
            "last_used_tab": self.config.getint("UI", "LAST_USED_TAB", fallback=0),
        }

    def get_email_settings(self):
        """Връща настройките за имейл"""
        return {
            "smtp_server": self.config.get("EMAIL", "SMTP_SERVER", fallback=""),
            "smtp_port": self.config.getint("EMAIL", "SMTP_PORT", fallback=587),
            "email_user": self.config.get("EMAIL", "EMAIL_USER", fallback=""),
            "email_password": self.config.get("EMAIL", "EMAIL_PASSWORD", fallback=""),
            "default_recipient": self.config.get("EMAIL", "DEFAULT_RECIPIENT", fallback=""),
            "analysis_notification_email": self.config.get("EMAIL", "ANALYSIS_NOTIFICATION_EMAIL", fallback=""),
            "use_outlook": self.config.getboolean("EMAIL", "USE_OUTLOOK", fallback=False),
            "outlook_signature": self.config.getboolean("EMAIL", "OUTLOOK_SIGNATURE", fallback=True),
        }

    def get_max_threads(self):
        """Връща максималния брой нишки"""
        return self.config.getint("SETTINGS", "MAX_THREADS", fallback=MAX_WORKERS)
    
    def is_cache_enabled(self):
        return self.config.getboolean("SETTINGS", "CACHE_ENABLED", fallback=True)
    
    def get_cache_timeout(self):
        return self.config.getint("SETTINGS", "CACHE_TIMEOUT", fallback=CACHE_TIMEOUT)

    #<Пач кодът е вмъкнат тук – виж следващата стъпка>

class SecurityTools:


    async def async_batch_ip_analysis(self, ips, services, proxy_settings=None):
        """Асинхронна пакетна проверка на IP адреси"""
        
        # Определяме кои услуги ще използваме
        if not services or "all" in services:
            services = ["abuseipdb", "virustotal", "shodan", "geo", "whois"]
        
        # Създаваме задачи за всяка услуга
        tasks = []
        if 'abuseipdb' in services and self.config.get_api_key("ABUSEIP_API_KEY"):
            tasks.append(self.async_batch_abuseipdb(ips, proxy_settings))
        if 'ipapi' in services:
            tasks.append(self.async_batch_ipapi(ips, proxy_settings))
        if 'ipgeolocation' in services and self.config.get_api_key("GEO_API_KEY"):
            tasks.append(self.async_batch_ipgeolocation(ips, proxy_settings, self.config.get_api_key("GEO_API_KEY")))
        if 'geojs' in services:
            tasks.append(self.async_batch_geojs(ips, proxy_settings))
        if 'ipinfo' in services and self.config.get_api_key("IPINFO_API_KEY"):
            tasks.append(self.async_batch_ipinfo(ips, proxy_settings))
        if 'findip' in services and self.config.get_api_key("FINDIP_API_KEY"):
            tasks.append(self.async_batch_findip(ips, proxy_settings))
        if 'ip2location' in services and self.config.get_api_key("IP2LOCATION_API_KEY"):
            tasks.append(self.async_batch_ip2location(ips, proxy_settings))
        if 'maxmind' in services and self.config.get_api_key("MAXMIND_ACCOUNT_ID") and self.config.get_api_key("MAXMIND_LICENSE_KEY"):
            tasks.append(self.async_batch_maxmind(ips, proxy_settings))

        
        # Изпълняваме задачите паралелно
        completed = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Комбинираме резултатите
        results = []
        for ip in ips:
            combined = {'ip': ip, 'country': 'N/A', 'isp': 'N/A', 'source': 'N/A'}
            
            for service_results in completed:
                if isinstance(service_results, Exception):
                    continue
                    
                for result_ip, data in service_results:
                    #print(f"[DEBUG] IP: {ip}, Data: {data}")
                    if result_ip != ip or data.get('error'):
                        continue
                    # Обработка на AbuseIPDB
                    if 'countryCode' in data:
                        combined.update({
                            'country': data.get('countryCode', 'N/A'),
                            'isp': data.get('isp', 'N/A'),
                            #'source': 'AbuseIPDB'
                        })
                    elif data.get("source") == "ip2location":
                        print("IP2LOCATION RAW:", data)
                        if combined.get("country") in [None, "N/A", ""]:
                            combined["country"] = data.get("country_name", "N/A")
                        
                        # Винаги попълваме ISP с "as"
                        combined["isp"] = data.get("as", "N/A")
                        #combined["source"] = "ip2location"
                        
                    elif 'country_name' in data:
                        combined.update({
                            'country': data.get('country_name', 'N/A'),
                            'isp': data.get('organization') or data.get('isp') or 'N/A',
                            #'source': 'ipgeolocation'
                        })
                    elif data.get("source") == "ipinfo" or "org" in data:
                        combined.update({
                            'country': data.get('country', 'N/A'),
                            'isp': data.get('org', 'N/A'),
                            #'source': 'IPinfo'
                        })
                    elif "traits" in data and "country" in data:
                        combined.update({
                            'country': data.get("country", {}).get("names", {}).get("en", "N/A"),
                            'isp': data.get("traits", {}).get("isp", "N/A"),
                            #'source': 'FindIP'
                        })
                    elif data.get("source") == "maxmind":
                        if combined.get("country") in [None, "N/A", ""]:
                            combined["country"] = data.get("country", {}).get("names", {}).get("en", "N/A")
                        if combined.get("isp") in [None, "N/A", ""]:
                            combined["isp"] = data.get("traits", {}).get("isp", "N/A")
                        combined["source"] = "maxmind"

                    # Обработка на GeoIP/ipapi/ipgeolocation
                    elif 'country' in data:
                        country = data['country']
                        if '(' in country:
                            country = country.split('(')[0].strip()
                        
                        isp = data.get('isp') or data.get('organization') or data.get('org') or 'N/A'
                        
                        combined.update({
                            'country': country or 'N/A',
                            'isp': isp,
                            #'source': data.get('source', 'N/A')
                            })
                    elif data.get('source') == 'geoip':
                        combined.update({
                            'country': data.get('country', 'N/A'),
                            'isp': data.get('isp', 'N/A'),
                            #'source': 'geoip'
                        })
            
            results.append((ip, combined))
        
        return results
    

    async def fetch_with_aiohttp(self, session, url, ip, headers=None, params=None, semaphore=None):
        try:
            async with semaphore:
                async with session.get(url, headers=headers, params=params, timeout=2.0) as response:
                    if response.status == 200:
                        data = await response.json()
                        return (ip, data)
                    return (ip, {"error": f"HTTP Status: {response.status}"})
        except Exception as e:
            return (ip, {"error": str(e)})

    async def async_batch_geojs(self, ips, proxy=None):
        results = []
        connector = aiohttp.TCPConnector(ssl=False, limit=0)
        headers = {
            "User-Agent": "SOC-Tool"
        }

        proxy_url= proxy.get("http") or proxy.get("https") if proxy else None
        semaphore = asyncio.Semaphore(300)

        async with aiohttp.ClientSession(connector=connector, proxy=proxy_url) as session:
            tasks = []
            for ip in ips:
                url = f"https://get.geojs.io/v1/ip/geo/{ip}.json"
                tasks.append(self.fetch_with_aiohttp(session, url, ip, headers=headers, semaphore=semaphore))
            results = await asyncio.gather(*tasks)
        return results

    async def async_batch_ipinfo(self, ips, proxy=None):
        results = []
        connector = aiohttp.TCPConnector(ssl=False, limit=0)
        proxy_url = proxy.get("http") or proxy.get("https") if proxy else None
        headers = {
            "User-Agent": "SOC-Tool"
        }
        token = self.config.get_api_key("IPINFO_API_KEY")
        if token:
            headers["Authorization"] = f"Bearer {token}"

        semaphore = asyncio.Semaphore(300)

        async with aiohttp.ClientSession(connector=connector, proxy=proxy_url) as session:
            tasks = []
            for ip in ips:
                url = f"https://ipinfo.io/{ip}?token={token}" if token else f"https://ipinfo.io/{ip}"
                tasks.append(self.fetch_with_aiohttp(session, url, ip, headers=headers, semaphore=semaphore))
            results = await asyncio.gather(*tasks)
        return results

    async def async_batch_findip(self, ips, proxy=None):
        results = []
        connector = aiohttp.TCPConnector(ssl=False, limit=0)
        proxy_url = proxy.get("http") or proxy.get("https") if proxy else None
        headers = {"User-Agent": "SOC-Tool"}

        token = self.config.get_api_key("FINDIP_API_KEY")
        if not token:
            return []

        semaphore = asyncio.Semaphore(300)

        async with aiohttp.ClientSession(connector=connector, proxy=proxy_url) as session:
            tasks = []
            for ip in ips:
                url = f"https://api.findip.net/{ip}/?token={token}"
                tasks.append(self.fetch_with_aiohttp(session, url, ip, headers=headers, semaphore=semaphore))
            results = await asyncio.gather(*tasks)
        return results

    async def async_batch_ip2location(self, ips, proxy=None):
        results = []
        api_key = self.config.get_api_key("IP2LOCATION_API_KEY")
        if not api_key:
            return []

        connector = aiohttp.TCPConnector(ssl=False, limit=0)
        proxy_url = proxy.get("http") or proxy.get("https") if proxy else None
        headers = {"User-Agent": "SOC-Tool"}
        semaphore = asyncio.Semaphore(50)

        async with aiohttp.ClientSession(connector=connector, proxy=proxy_url) as session:
            tasks = []
            for ip in ips:
                url = f"https://api.ip2location.io/?key={api_key}&ip={ip}&format=json"
                tasks.append(self.fetch_with_aiohttp(session, url, ip, headers=headers, semaphore=semaphore))
            results = await asyncio.gather(*tasks)
        return results
    
    async def async_batch_maxmind(self, ips, proxy=None):
        results = []
        account_id = self.config.get_api_key("MAXMIND_ACCOUNT_ID")
        license_key = self.config.get_api_key("MAXMIND_LICENSE_KEY")
        if not account_id or not license_key:
            return []

        proxy_url = proxy.get("http") or proxy.get("https") if proxy else None
        connector = aiohttp.TCPConnector(ssl=False, limit=0)
        auth = aiohttp.BasicAuth(account_id, password=license_key)
        semaphore = asyncio.Semaphore(50)

        async with aiohttp.ClientSession(connector=connector, proxy=proxy_url, auth=auth) as session:
            tasks = []
            for ip in ips:
                url = f"https://geoip.maxmind.com/geoip/v2.1/city/{ip}"
                tasks.append(self.fetch_with_aiohttp(session, url, ip, headers={"User-Agent": "SOC-Tool"}, semaphore=semaphore))
            results = await asyncio.gather(*tasks)
        return results

    async def async_batch_ipapi(self, ips, proxy=None):
        results = []
        connector = aiohttp.TCPConnector(ssl=False, limit=0)
        proxy_url= proxy.get("http") or proxy.get("https") if proxy else None
        headers = {"User-Agent": "SOC-Tool"}
        semaphore = asyncio.Semaphore(5)

        async with aiohttp.ClientSession(connector=connector, proxy=proxy_url) as session:
            tasks = []
            for ip in ips:
                url = f"http://ip-api.com/json/{ip}"
                await asyncio.sleep(0.05)  # 20 заявки/секунда
                tasks.append(self.fetch_with_aiohttp(session, url, ip, headers=headers, semaphore=semaphore))
            results = await asyncio.gather(*tasks)
        return results

    async def async_batch_ipgeolocation(self, ips, proxy=None, api_key=None):
        results = []
        connector = aiohttp.TCPConnector(ssl=False, limit=0)
        proxy_url= proxy.get("http") or proxy.get("https") if proxy else None
        headers = {"User-Agent": "SOC-Tool"}
        self.config.get_api_key("GEO_API_KEY")
        
        semaphore = asyncio.Semaphore(20)
        async with aiohttp.ClientSession(connector=connector, proxy=proxy_url) as session:
            tasks = []
            for ip in ips:
                url = f"https://api.ipgeolocation.io/ipgeo?apiKey={api_key}&ip={ip}"
                await asyncio.sleep(0.05)  # 20 заявки/секунда
                tasks.append(self.fetch_with_aiohttp(session, url, ip, headers=headers, semaphore=semaphore))
            results = await asyncio.gather(*tasks)
        return results

    async def async_batch_abuseipdb(self, ips, proxy=None, api_key=None):
        results = []
        connector = aiohttp.TCPConnector(ssl=False, limit=0)
        proxy_url= proxy.get("http") or proxy.get("https") if proxy else None
        headers = {
            "Key": self.config.get_api_key("ABUSEIP_API_KEY"),
            "Accept": "application/json"
        } 
        semaphore = asyncio.Semaphore(300)

        async with aiohttp.ClientSession(connector=connector, proxy=proxy_url) as session:
            tasks = []
            for ip in ips:
                url = "https://api.abuseipdb.com/api/v2/check"
                params = {"ipAddress": ip, "maxAgeInDays": "90"}
                tasks.append(self.fetch_with_aiohttp(session, url, ip, headers=headers, params=params, semaphore=semaphore))
            raw_results = await asyncio.gather(*tasks)
            results = [(ip, data.get("data", {})) for ip, data in raw_results]
            
        return results
    
    async def async_batch_geoip(self, ips):
        results = []
        for ip in ips:
            try:
                response = self.geoip_reader.city(ip)
                country = response.country.name or "N/A"
                isp = response.traits.isp or "N/A"  # ако базата го поддържа
                results.append((ip, {
                    'country': country,
                    'isp': isp,
                    'source': 'geoip'
                }))
            except Exception as e:
                results.append((ip, {'error': str(e)}))
        return results

    def batch_ip_analysis_fast(self, ips, services, proxy_settings=None, threads=1000):
        proxy_settings = self.config.get_proxy_settings()
        results = batch_ip_analysis_fast(ip_list, proxy=proxy_settings)
        if not services:
            services = ["abuseipdb", "ipapi", "ipgeolocation", "geojs", "ipinfo", "findip", "ip2location", "maxmind"]
        
        try:
            return asyncio.run(self.async_batch_ip_analysis(ips, services, proxy_settings))
        except Exception as e:
            return [(ip, {"error": str(e)}) for ip in ips]






    """Основни инструменти за сигурност и анализ"""
    """Разширени инструменти за сигурност с подобрения"""

    def __init__(self, config, log_activity=None):
        self.config = config
        self.log_activity = log_activity or (lambda msg: print(f"[LOG] {msg}"))
        self.api_lock = Lock()
        self.cache = {}
        self.geo_data_source = GeoDataSource(config) 

        
        # Оптимизирана HTTP сесия
        self.session = requests.Session()
        retries = Retry(
            total=3,
            backoff_factor=1,
            status_forcelist=[500, 502, 503, 504]
        )
        self.session.mount('https://', HTTPAdapter(
            max_retries=retries,
            pool_connections=100,
            pool_maxsize=100
        ))

        try:
            self.geoip_reader = geoip2.database.Reader(GEOIP_DB_PATH)
        except Exception as e:
            print(f"Грешка при зареждане на GeoIP база: {e}")
            self.geoip_reader = None

    def _get_cache_key(self, func_name, *args):
        return f"{func_name}_{'_'.join(str(arg) for arg in args)}"
    
    def _cache_result(self, key, data):
        if self.config.is_cache_enabled():
            self.cache[key] = {
                'data': data,
                'timestamp': time.time()
            }
    
    def _get_cached_result(self, key):
        if not self.config.is_cache_enabled():
            return None
        
        cached = self.cache.get(key)
        if cached:
            if time.time() - cached['timestamp'] < self.config.get_cache_timeout():
                return cached['data']
        return None

    def scan_file_virustotal(self, file_path):
        """Подобрено сканиране на файлове с VirusTotal"""
        cache_key = self._get_cache_key("scan_file_vt", file_path)
        cached = self._get_cached_result(cache_key)
        if cached:
            return cached
        
        with self.api_lock:
            if not self.config.get_api_key("VIRUSTOTAL_API_KEY"):
                return {"error": "Липсва VirusTotal API ключ"}

            url = "https://www.virustotal.com/api/v3/files"
            headers = {"x-apikey": self.config.get_api_key("VIRUSTOTAL_API_KEY")}

            try:
                # Проверка за размер на файла (650MB за VirusTotal v3 API)
                file_size = os.path.getsize(file_path)
                if file_size > 650 * 1024 * 1024:
                    return {"error": "Файлът надвишава максималния размер от 650MB"}

                with open(file_path, "rb") as file:
                    files = {"file": file}
                    response = self.session.post(
                        url,
                        headers=headers,
                        files=files,
                        proxies=self.config.get_proxy_settings(),
                        verify=False,
                        timeout=120,
                    )

                if response.status_code == 200:
                    data = response.json()
                    analysis_id = data.get("data", {}).get("id", "")
                    
                    # Взимаме пълната информация за файла, за да получим SHA256
                    file_url = f"https://www.virustotal.com/api/v3/files/{analysis_id}"
                    file_response = self.session.get(
                        file_url,
                        headers=headers,
                        proxies=self.config.get_proxy_settings(),
                        verify=False,
                        timeout=30,
                    )
                    
                    if file_response.status_code == 200:
                        file_data = file_response.json()
                        attributes = file_data.get("data", {}).get("attributes", {})
                        sha256 = attributes.get("sha256", "")
                        last_analysis = attributes.get("last_analysis_stats", {})
                        results = attributes.get("last_analysis_results", {})
                        
                        vendors = [
                            f"{k}: {v['result']}" 
                            for k, v in results.items()
                            if v.get('category') == 'malicious'
                        ][:10]
                        
                        result = {
                            "id": analysis_id,
                            "sha256": sha256,
                            "positives": last_analysis.get("malicious", 0),
                            "total": sum(last_analysis.values()),
                            "permalink": f"https://www.virustotal.com/gui/file/{sha256}" if sha256 else f"https://www.virustotal.com/gui/analysis/{analysis_id}",
                            "vendors": vendors,
                            "detections_str": f"{last_analysis.get('malicious', 0)}/{sum(last_analysis.values())} детекции",
                            "type": "instant"
                        }
                        self._cache_result(cache_key, result)
                        return result
                    
                    # Ако не успеем да вземем SHA256, използваме анализа
                    analysis_url = f"https://www.virustotal.com/api/v3/analyses/{analysis_id}"
                    analysis_response = self.session.get(
                        analysis_url,
                        headers=headers,
                        proxies=self.config.get_proxy_settings(),
                        verify=False,
                        timeout=30,
                    )
                    
                    if analysis_response.status_code == 200:
                        analysis_data = analysis_response.json()
                        stats = analysis_data.get("data", {}).get("attributes", {}).get("stats", {})
                        results = analysis_data.get("data", {}).get("attributes", {}).get("results", {})
                        
                        vendors = [
                            f"{k}: {v['result']}" 
                            for k, v in results.items()
                            if v.get('category') == 'malicious'
                        ][:10]
                        
                        # Вземаме SHA256 от метаданните на анализа, ако е наличен
                        meta = analysis_data.get("data", {}).get("attributes", {}).get("meta", {})
                        file_info = meta.get("file_info", {})
                        sha256 = file_info.get("sha256", "")
                        
                        result = {
                            "id": analysis_id,
                            "sha256": sha256,
                            "positives": stats.get("malicious", 0),
                            "total": stats.get("harmless", 0) + stats.get("malicious", 0) + 
                                    stats.get("suspicious", 0) + stats.get("undetected", 0),
                            "permalink": f"https://www.virustotal.com/gui/file/{sha256}" if sha256 else f"https://www.virustotal.com/gui/analysis/{analysis_id}",
                            "vendors": vendors,
                            "detections_str": f"{stats.get('malicious', 0)}/{stats.get('harmless', 0) + stats.get('malicious', 0) + stats.get('suspicious', 0) + stats.get('undetected', 0)} детекции",
                            "type": "analysis"
                        }
                        self._cache_result(cache_key, result)
                        return result
                    else:
                        return {"error": f"API грешка при получаване на резултати: {analysis_response.status_code}"}
                elif response.status_code == 429:
                    return {"error": "Прекалено много заявки. Моля, изчакайте."}
                else:
                    return {"error": f"API грешка: {response.status_code}"}
            except Exception as e:
                return {"error": str(e)}
            
    def scan_url_virustotal(self, url):
        """Работещо сканиране на URL с VirusTotal v3 API, с реални детекции и валиден permalink"""
        import urllib.parse
        import time
        import certifi

        cache_key = self._get_cache_key("scan_url_vt", url)
        cached = self._get_cached_result(cache_key)
        if cached:
            return cached

        with self.api_lock:
            api_key = self.config.get_api_key("VIRUSTOTAL_API_KEY")
            if not api_key:
                return {"error": "Липсва VirusTotal API ключ"}

            headers = {
                "x-apikey": api_key,
                "Content-Type": "application/x-www-form-urlencoded",
            }
            data = f"url={urllib.parse.quote(url)}"

            try:
                # Стъпка 1: Изпрати URL за анализ
                response = self.session.post(
                    "https://www.virustotal.com/api/v3/urls",
                    headers=headers,
                    data=data,
                    proxies=self.config.get_proxy_settings(),
                    verify=certifi.where(),
                    timeout=30,
                )

                if response.status_code != 200:
                    return {"error": f"Грешка при подаване към VT: {response.status_code}"}

                json_data = response.json()
                analysis_id = json_data.get("data", {}).get("id", "")
                if not analysis_id:
                    return {"error": "Липсва analysis_id от VT"}

                # Стъпка 2: Изчакай и изтегли резултатите от анализа
                time.sleep(2)  # Важно! Анализът не е готов веднага

                result_response = self.session.get(
                    f"https://www.virustotal.com/api/v3/analyses/{analysis_id}",
                    headers=headers,
                    proxies=self.config.get_proxy_settings(),
                    verify=certifi.where(),
                    timeout=30,
                )

                if result_response.status_code != 200:
                    return {"error": f"Грешка при резултата: {result_response.status_code}"}

                analysis_data = result_response.json()
                stats = analysis_data.get("data", {}).get("attributes", {}).get("stats", {})
                malicious = stats.get("malicious", 0)
                suspicious = stats.get("suspicious", 0)
                total = sum(stats.values())
                detections = malicious + suspicious

                # Вземи URL ID за валиден линк
                url_bytes = url.encode()
                url_id = base64.urlsafe_b64encode(url_bytes).decode().strip("=")

                result = {
                    "id": analysis_id,
                    "positives": detections,
                    "total": total,
                    "permalink": f"https://www.virustotal.com/gui/url/{url_id}",
                    "detections_str": f"{detections}/{total} детекции"
                }

                self._cache_result(cache_key, result)
                return result

            except Exception as e:
                return {"error": f"VirusTotal грешка: {str(e)}"}



    def scan_hybrid_analysis(self, file_path=None, hash_value=None, ip=None, domain=None):
        """Подобрено сканиране с Hybrid Analysis с имейл нотификации"""
        cache_key = self._get_cache_key("scan_ha", file_path or hash_value or ip or domain)
        cached = self._get_cached_result(cache_key)
        if cached:
            return cached
        
        with self.api_lock:
            if not self.config.get_api_key("HYBRIDANALYSIS_API_KEY"):
                return {"error": "Липсва HybridAnalysis API ключ"}

            email_settings = self.config.get_email_settings()
            notification_email = email_settings.get("analysis_notification_email", "")
            
            headers = {
                "api-key": self.config.get_api_key("HYBRIDANALYSIS_API_KEY"),
                "User-Agent": "Falcon",
                "accept": "application/json",
            }

            try:
                if file_path:
                    params = {
                        "environment_id": 160,
                        "hybrid_analysis": "true",
                    }
                    
                    if notification_email:
                        params["email"] = notification_email
                    
                    with open(file_path, "rb") as file:
                        files = {"file": file}
                        response = self.session.post(
                            "https://www.hybrid-analysis.com/api/v2/submit/file",
                            headers=headers,
                            files=files,
                            data=params,
                            proxies=self.config.get_proxy_settings(),
                            verify=False,
                            timeout=60,  # Увеличаваме таймаута за файлове
                        )
                    
                    # Обработка на response за файлове
                    if response.status_code == 201:  # Accepted - processing
                        submission_id = response.json().get("job_id")
                        if submission_id:
                            # Проверяваме статуса на анализа
                            status_url = f"https://www.hybrid-analysis.com/api/v2/report/{submission_id}/state"
                            status_response = self.session.get(
                                status_url,
                                headers=headers,
                                proxies=self.config.get_proxy_settings(),
                                verify=False,
                                timeout=30,
                            )
                            
                            if status_response.status_code == 200:
                                state = status_response.json().get("state")
                                if state == "SUCCESS":
                                    # Взимаме резултатите
                                    result_url = f"https://www.hybrid-analysis.com/api/v2/report/{submission_id}/summary"
                                    result_response = self.session.get(
                                        result_url,
                                        headers=headers,
                                        proxies=self.config.get_proxy_settings(),
                                        verify=False,
                                        timeout=30,
                                    )
                                    
                                    if result_response.status_code == 200:
                                        result = {
                                            "data": result_response.json(),
                                            "notification_sent": bool(notification_email)
                                        }
                                        self._cache_result(cache_key, result)
                                        return result
                                    return {"error": f"API грешка при получаване на резултати: {result_response.status_code}"}
                                return {"error": f"Анализът все още се обработва. Статус: {state}"}
                            return {"error": f"API грешка при проверка на статуса: {status_response.status_code}"}
                        return {"error": "Неуспешно получаване на submission ID"}
                    
                # Оригинална логика за всички случаи (hash, ip, domain)
                if hash_value:
                    response = self.session.post(
                        "https://www.hybrid-analysis.com/api/v2/search/hash",
                        headers=headers,
                        data={"hash": hash_value},
                        proxies=self.config.get_proxy_settings(),
                        verify=False,
                        timeout=30,
                    )
                elif ip:
                    response = self.session.post(
                        "https://www.hybrid-analysis.com/api/v2/search/terms",
                        headers=headers,
                        data={"host": ip},
                        proxies=self.config.get_proxy_settings(),
                        verify=False,
                        timeout=30,
                    )
                elif domain:
                    response = self.session.post(
                        "https://www.hybrid-analysis.com/api/v2/search/terms",
                        headers=headers,
                        data={"domain": domain},
                        proxies=self.config.get_proxy_settings(),
                        verify=False,
                        timeout=30,
                    )
                else:
                    return {"error": "Не е предоставен параметър за търсене"}

                if response.status_code == 200:
                    result = {
                        "data": response.json(),
                        "notification_sent": bool(notification_email)
                    }
                    self._cache_result(cache_key, result)
                    return result
                return {"error": f"Грешка в API: {response.status_code}"}
            except Exception as e:
                return {"error": str(e)}
            
    @lru_cache(maxsize=1000)
    @functools.lru_cache(maxsize=1000)
    
    
    @lru_cache(maxsize=1000)
    @functools.lru_cache(maxsize=1000)
    def analyze_ip(self, ip_address, services=None):
        """Оптимизиран анализ на IP адрес с кеширане"""
        if services is None:
            services=(
                "abuseipdb",
                "virustotal",
                "shodan",
                "geo",
                "geojs",
                "ipapi",
                "ipgeolocation",
                "whois",)

        cache_key = self._get_cache_key("analyze_ip", ip_address, *sorted(services))
        cached = self._get_cached_result(cache_key)
        if cached:
            return cached

        results = {}
        with concurrent.futures.ThreadPoolExecutor(
            max_workers=min(self.config.get_max_threads(), len(services))
        ) as executor:
            future_to_service = {}

            if "abuseipdb" in services:
                future_to_service[executor.submit(self.check_abuseipdb, ip_address)] = "abuseipdb"
            if "virustotal" in services:
                future_to_service[executor.submit(self.check_virustotal_ip, ip_address)] = "virustotal"
            if "shodan" in services:
                future_to_service[executor.submit(self.check_shodan, ip_address)] = "shodan"
            if "geo" in services:
                future_to_service[executor.submit(self.get_geolocation, ip_address)] = "geo"
            if "geojs" in services:
                future_to_service[executor.submit(self.check_geojs, ip_address)] = "geojs"
            if "ipapi" in services:
                future_to_service[executor.submit(self.check_ip_api, ip_address)] = "ipapi"
            if "ipgeolocation" in services:
                future_to_service[executor.submit(self.check_ipgeolocation, ip_address)] = "ipgeolocation"
            if "whois" in services:
                future_to_service[executor.submit(self.get_whois, ip_address)] = "whois"

            for future in concurrent.futures.as_completed(future_to_service):
                service = future_to_service[future]
                try:
                    results[service] = future.result()
                except Exception as e:
                    results[service] = {"error": str(e)}

        self._cache_result(cache_key, results)
        return results

    def check_abuseipdb(self, ip):
        """Проверява IP с AbuseIPDB"""
        with self.api_lock:
            if not self.config.get_api_key("ABUSEIP_API_KEY"):
                return {"error": "Липсва AbuseIPDB API ключ"}

            url = "https://api.abuseipdb.com/api/v2/check"
            headers = {
                "Key": self.config.get_api_key("ABUSEIP_API_KEY"),
                "Accept": "application/json",
            }
            params = {"ipAddress": ip, "maxAgeInDays": "90"}

            try:
                response = requests.get(
                    url,
                    headers=headers,
                    params=params,
                    proxies=self.config.get_proxy_settings(),
                    verify=False,
                    timeout=2,
                )

                if response.status_code == 200:
                    data = response.json()["data"]
                    return {
                        "isPublic": data.get("isPublic", False),
                        "ipVersion": data.get("ipVersion", 4),
                        "isWhitelisted": data.get("isWhitelisted", False),
                        "abuseConfidenceScore": data.get("abuseConfidenceScore", 0),
                        "countryCode": data.get("countryCode", "N/A"),
                        "usageType": data.get("usageType", "N/A"),
                        "isp": data.get("isp", "N/A"),
                        "domain": data.get("domain", "N/A"),
                        "totalReports": data.get("totalReports", 0),
                        "lastReportedAt": data.get("lastReportedAt", "N/A"),
                    }
                return {"error": f"Грешка в API: {response.status_code}"}
            except Exception as e:
                return {"error": str(e)}

    def check_virustotal_ip(self, ip):
        """Проверява IP с VirusTotal"""
        with self.api_lock:
            if not self.config.get_api_key("VIRUSTOTAL_API_KEY"):
                return {"error": "Липсва VirusTotal API ключ"}

            url = f"https://www.virustotal.com/api/v3/ip_addresses/{ip}"
            headers = {"x-apikey": self.config.get_api_key("VIRUSTOTAL_API_KEY")}

            try:
                response = requests.get(
                    url,
                    headers=headers,
                    proxies=self.config.get_proxy_settings(),
                    verify=False,
                    timeout=30,
                )

                if response.status_code == 200:
                    data = response.json()["data"]
                    attributes = data.get("attributes", {})
                    return {
                        "asn": attributes.get("asn", "N/A"),
                        "as_owner": attributes.get("as_owner", "N/A"),
                        "country": attributes.get("country", "N/A"),
                        "reputation": attributes.get("reputation", 0),
                        "last_analysis_stats": attributes.get(
                            "last_analysis_stats", {}
                        ),
                        "total_votes": attributes.get("total_votes", {}),
                        "network": attributes.get("network", "N/A"),
                    }
                return {"error": f"Грешка в API: {response.status_code}"}
            except Exception as e:
                return {"error": str(e)}

    def check_shodan(self, ip):
        """Проверява IP с Shodan"""
        with self.api_lock:
            if not self.config.get_api_key("SHODAN_API_KEY"):
                return {"error": "Липсва Shodan API ключ"}

            try:
                api = shodan.Shodan(self.config.get_api_key("SHODAN_API_KEY"))
                result = api.host(ip)

                return {
                    "ports": result.get("ports", []),
                    "hostnames": result.get("hostnames", []),
                    "org": result.get("org", "N/A"),
                    "os": result.get("os", "N/A"),
                    "vulns": result.get("vulns", []),
                    "tags": result.get("tags", []),
                    "asn": result.get("asn", "N/A"),
                    "isp": result.get("isp", "N/A"),
                    "last_update": result.get("last_update", "N/A"),
                }
            except shodan.APIError as e:
                return {"error": str(e)}
            except Exception as e:
                return {"error": str(e)}

    def get_geolocation(self, ip, force_source=None):
        """
        Връща геолокационни данни за IP адрес
        :param ip: IP адрес за проверка
        :param force_source: опционален параметър за конкретен източник
        :return: резултати от гео проверка
        """
        cache_key = self._get_cache_key("get_geolocation", ip, force_source if force_source else "auto")
        cached = self._get_cached_result(cache_key)
        if cached:
            return cached
        
        result = self.geo_data_source.get_geo_data(ip, force_source)
        self._cache_result(cache_key, result)
        return result


    def get_whois(self, ip_or_domain):
        """Връща WHOIS информация с прокси поддръжка за IP и домейн"""
        try:
            ipaddress.ip_address(ip_or_domain)  # Проверка дали е IP
            try:
                rdap_url = f"https://rdap.arin.net/registry/ip/{ip_or_domain}"
                response = requests.get(
                    rdap_url,
                    proxies=self.config.get_proxy_settings(),
                    timeout=5
                )
                if response.status_code != 200:
                    return {"error": f"RDAP отговор: {response.status_code}"}
                result = response.json()
                return {
                    "asn": result.get("handle", "N/A"),
                    "org": result.get("name", "N/A"),
                    "country": result.get("country", "N/A"),
                    "range": f"{result.get('startAddress', '')} - {result.get('endAddress', '')}",
                    "remarks": [r.get("description", [""])[0] for r in result.get("remarks", [])],
                    "emails": [
                        v for r in result.get("entities", [])
                        for v in r.get("vcardArray", [[], []])[1] if v[0] == "email"
                    ]
                }
            except Exception as e:
                return {"error": f"WHOIS IP грешка: {str(e)}"}

        except ValueError:
            # не е IP → домейн
            try:
                domain_info = whois.whois(ip_or_domain)
                return {
                    "domain_name": domain_info.domain_name,
                    "registrar": domain_info.registrar,
                    "whois_server": domain_info.whois_server,
                    "creation_date": domain_info.creation_date,
                    "expiration_date": domain_info.expiration_date,
                    "updated_date": domain_info.updated_date,
                    "name_servers": domain_info.name_servers,
                    "status": domain_info.status,
                    "emails": domain_info.emails,
                    "dnssec": domain_info.dnssec,
                }
            except Exception as e:
                return {"error": f"WHOIS домейн грешка: {str(e)}"}



    def dns_lookup(self, domain, record_type="A"):
        """Извършва DNS заявка"""
        try:
            answers = dns.resolver.resolve(domain, record_type)
            return [str(r) for r in answers]
        except Exception as e:
            return {"error": str(e)}

    def check_mxtoolbox(self, query, query_type):
        """Проверява с MXToolbox"""
        with self.api_lock:
            if not self.config.get_api_key("MXTOOLBOX_API_KEY"):
                return {"error": "Липсва MXToolbox API ключ"}

            url = f"https://api.mxtoolbox.com/api/v1/lookup/{query_type}/{query}"
            headers = {
                "Authorization": self.config.get_api_key("MXTOOLBOX_API_KEY"),
                "Accept": "application/json",
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
            }      
            proxy_settings = self.config.get_proxy_settings()
            print(f"[DEBUG] MXToolbox заявка: {url}")
            print(f"[DEBUG] Headers: {headers}")
            print(f"[DEBUG] Прокси настройки: {proxy_settings}")

            try:
                response = requests.get(
                    url,
                    headers=headers,
                    proxies=self.config.get_proxy_settings(),
                    verify=False,
                    timeout=10,
                )
                
                print(f"[DEBUG] Статус код: {response.status_code}")
                print(f"[DEBUG] Отговор: {response.text[:200]}...")
                
                if response.status_code == 200:
                    data = response.json()
                    # Обработка на различните типове заявки
                    if query_type == "blacklist":
                        return {
                            "data": {
                                "results": [
                                    {
                                        "name": item.get("Name"),
                                        "status": item.get("Status")
                                    }
                                    for item in data.get("Information", [])
                                ]
                            }
                        }
                    elif query_type in ["a", "mx", "txt"]:
                        return {
                            "data": {
                                "records": [
                                    record.get("Value")
                                    for record in data.get("Information", [])
                                ]
                            }
                        }
                    else:
                        return {"data": data}
                return {"error": f"API грешка: {response.status_code}"}
            except Exception as e:
                print(f"[ERROR] Изключение при заявка към MXToolbox: {str(e)}")
                return {"error": str(e)}

    def report_phishing(self, url, service, log=None):
        """Докладва фишинг URL към различни услуги"""
        email = self.config.get_api_key("analysis_notification_email")

        services = {
            "google": {
                "url": "https://safebrowsing.google.com/safebrowsing/report_phish/",
                "params": {"url": url},
                "method": "POST"
            },
            "phishtank": {
                "url": "https://www.phishtank.com/api_add_url.php",
                "params": {
                    "url": url,
                    "format": "json",
                    "app_key": self.config.get_api_key("PHISHTANK_API_KEY"),
                    "User-Agent": "phishtank/shakalakaboom:DSC SOC Tool"
                },
                "method": "POST"
            },
            "microsoft": {
                "url": "https://www.microsoft.com/en-us/wdsi/support/report-unsafe-site",
                "params": {"url": url},
                "method": "POST"
            },
            "openphish": {
                "url": "https://www.circl.lu/urlabuse/",
                "params": {"url": url},
                "method": "POST"
            },
            "apwg": {
                "url": "https://apwg.org/report-phishing/",
                "params": {"url": url},
                "method": "POST"
            },
            "netcraft": {
                "url": "https://report.netcraft.com/phishing-submit",
                "params": {
                    "url": url,
                    "email": email,
                    "submit": "Submit"
                },
                "method": "POST"
            },
            "brightcloud": {
                "url": "https://www.brightcloud.com/tools/change-request-url-reputation-confirmation.php",
                "params": {
                    "url": url,
                    "category": "Phishing",
                    "reason": "Please inspect, flagged by SOC Tool",
                    "email": email,
                    "confirm": "Submit"
                },
                "method": "POST"
            }
        }

        if service not in services:
            if log:
                log(f"Невалидна услуга за докладване: {service}")
            return {"error": "Невалидна услуга"}

        service_info = services[service]

        try:
            proxies = self.config.get_proxy_settings()
            headers = {"User-Agent": "SOC-Tool/1.0"}

            if service_info["method"] == "POST":
                response = requests.post(service_info["url"], data=service_info["params"], proxies=proxies, headers=headers, timeout=15, verify=True)
            else:
                response = requests.get(service_info["url"], params=service_info["params"], proxies=proxies, headers=headers, timeout=15, verify=False)

            text = response.text.lower()

            if service == "phishtank":
                try:
                    json_resp = response.json()
                    if json_resp.get("meta", {}).get("status") == "success":
                        if log:
                            log(f"PhishTank успешно получи доклада за {url}")
                        return {"success": True, "details": json_resp}
                    else:
                        error_text = json_resp.get("errortext", "PhishTank неизвестна грешка")
                        if log:
                            log(f"PhishTank грешка: {error_text}")
                        return {"error": error_text}
                except Exception:
                    return {"success": True, "message": "PhishTank доклад изпратен (без JSON отговор)"}

            if service == "openphish":
                try:
                    soup = BeautifulSoup(response.text, "html.parser")
                    result = {}
                    for section in soup.find_all("div", class_="panel-body"):
                        title = section.find_previous("div", class_="panel-heading")
                        if title:
                            key = title.get_text(strip=True)
                            val = section.get_text(separator="\n", strip=True)
                            result[key] = val
                    if result:
                        if log:
                            log(f"CIRCL URLAbuse информация за {url} извлечена успешно.")
                        return {"success": True, "data": result}
                    else:
                        return {"error": "CIRCL не върна полезни данни"}
                except Exception as parse_err:
                    return {"error": f"CIRCL парсинг грешка: {parse_err}"}

            if service in ["microsoft", "apwg", "brightcloud", "netcraft"]:
                if any(kw in text for kw in ["thank you", "submitted", "success", "report received"]):
                    if log:
                        log(f"Успешно докладвано към {service} за URL: {url}")
                    return {"success": True}
                else:
                    return {"error": f"{service} не потвърди успешно докладване (възможна captcha или отказ)"}

            if service == "google":
                if response.status_code == 200:
                    if log:
                        log(f"Успешно докладвано към Google за URL: {url}")
                    return {"success": True}
                else:
                    return {"error": f"Google върна статус {response.status_code}"}

            if response.status_code == 200:
                if log:
                    log(f"Успешно докладвано към {service} за URL: {url}")
                return {"success": True, "response": response.text}
            else:
                error_msg = f"Грешка {response.status_code}: {response.text}"
                if log:
                    log(f"{service} грешка при докладване на {url}: {error_msg}")
                return {"error": error_msg}

        except requests.exceptions.RequestException as e:
            if log:
                log(f"{service} изключение при докладване: {str(e)}")
            return {"error": str(e)}


    def check_geojs(self, ip):
        try:
            response = requests.get(
                f"https://get.geojs.io/v1/ip/geo/{ip}.json",
                proxies=self.config.get_proxy_settings(),
                timeout=2
            )
            if response.status_code == 200:
                data = response.json()
                return {
                    "country": data.get("country", "N/A"),
                    "region": data.get("region", "N/A"),
                    "city": data.get("city", "N/A"),
                    "organization": data.get("organization_name", "N/A"),
                    "timezone": data.get("timezone", "N/A"),
                    "accuracy": data.get("accuracy", "N/A"),
                }
            return {"error": f"Грешка от GeoJS: {response.status_code}"}
        except Exception as e:
            return {"error": str(e)}

    def check_ip_api(self, ip):
        try:
            response = requests.get(
                f"http://ip-api.com/json/{ip}",
                proxies=self.config.get_proxy_settings(),
                timeout=2
            )
            if response.status_code == 200:
                data = response.json()
                return {
                    "country": data.get("country", "N/A"),
                    "region": data.get("regionName", "N/A"),
                    "city": data.get("city", "N/A"),
                    "isp": data.get("isp", "N/A"),
                    "org": data.get("org", "N/A"),
                    "as": data.get("as", "N/A"),
                    "lat": data.get("lat", "N/A"),
                    "lon": data.get("lon", "N/A"),
                    "timezone": data.get("timezone", "N/A"),
                }
            return {"error": f"Грешка от ip-api.com: {response.status_code}"}
        except Exception as e:
            return {"error": str(e)}

        
    def check_ipgeolocation(self, ip):
        try:
            api_key = self.config.get_api_key("GEO_API_KEY")
            if not api_key:
                return {"error": "Липсва API ключ за ipgeolocation.io"}
            response = requests.get(
                f"https://api.ipgeolocation.io/ipgeo?apiKey={api_key}&ip={ip}",
                proxies=self.config.get_proxy_settings(),
                timeout=2
            )
            if response.status_code == 200:
                return response.json()
            return {"error": f"Грешка от ipgeolocation.io: {response.status_code}"}
        except Exception as e:

            return {"error": str(e)}
    def report_to_abuseipdb(self, ip, categories, comment=""):
        """Докладва IP към AbuseIPDB"""
        with self.api_lock:
            if not self.config.get_api_key("ABUSEIP_API_KEY"):
                return {"error": "Липсва AbuseIPDB API ключ"}

            url = "https://api.abuseipdb.com/api/v2/report"
            headers = {
                "Key": self.config.get_api_key("ABUSEIP_API_KEY"),
                "Accept": "application/json",
            }
            params = {
                "ip": ip,
                "categories": ",".join(map(str, categories)),
                "comment": comment,
            }

            try:
                response = requests.post(
                    url,
                    headers=headers,
                    params=params,
                    proxies=self.config.get_proxy_settings(),
                    verify=False,
                    timeout=2,
                )

                if response.status_code == 200:
                    return {"success": True}
                return {"error": f"Грешка в API: {response.status_code}"}
            except Exception as e:
                return {"error": str(e)}


    # Continuous Ping с възможност за стоп
    def continuous_ping(self, ips, stop_event, timeout=1, threads=200):
        CREATE_NO_WINDOW = 0x08000000 if os.name == 'nt' else 0
        def ping_ip(ip):
            try:
                CREATE_NO_WINDOW = 0x08000000 if os.name == 'nt' else 0
                subprocess.run(['ping', '-n', '1', '-w', str(timeout*1000), ip],
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, creationflags=CREATE_NO_WINDOW)
            except:
                pass

        while not stop_event.is_set():
            with ThreadPoolExecutor(max_workers=threads) as executor:
                for ip in ips:
                    if stop_event.is_set():
                        break
                    executor.submit(ping_ip, ip)

    # Оптимизиран Traceroute с прокси поддръжка
    def traceroute_with_proxy(self, host, proxy_settings=None, max_hops=30, timeout=1.2):
        results = []
        for ttl in range(1, max_hops + 1):
            try:
                s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                s.settimeout(timeout)
                s.setsockopt(socket.SOL_IP, socket.IP_TTL, ttl)

                if proxy_settings:
                    proxy_host, proxy_port = proxy_settings['http'].replace("http://", "").split(':')
                    s.connect((proxy_host, int(proxy_port)))
                    s.send(f"CONNECT {host}:80 HTTP/1.0\r\n\r\n".encode())
                else:
                    s.connect((host, 80))

                addr = s.getpeername()[0]
                results.append((ttl, addr))
                if addr == host:
                    break
            except:
                results.append((ttl, "*"))
            finally:
                s.close()
        return results

    # Бърз и многопоточен портов скенер с експорт
    def fast_port_scan(self, host, ports, mode='quick', threads=1000, timeout=0.5):
        open_ports = []
        lock = threading.Lock()

        def scan(port):
            s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            s.settimeout(timeout)
            result = s.connect_ex((host, port))
            if result == 0:
                with lock:
                    open_ports.append(port)
            s.close()

        with ThreadPoolExecutor(max_workers=threads) as executor:
            for port in ports:
                executor.submit(scan, port)

        return open_ports

    # Оптимизирана пакетна IP проверка

    #def batch_ip_analysis_fast(self, ips, services, proxy_settings=None, threads=100):
    #    results = []

     #   def analyze(ip):
     #       try:
     #           return (ip, self.analyze_ip(ip, services=services))
     #       except Exception as e:
      #          return (ip, {"error": str(e)})

       # with ThreadPoolExecutor(max_workers=threads) as executor:
        #    futures = {executor.submit(analyze, ip): ip for ip in ips}
         #   for future in as_completed(futures):
          #      results.append(future.result())

        #return results
        #def analyze(ip):
         #   return (ip, self.analyze_ip(ip, services=services))

        #with ThreadPoolExecutor(max_workers=threads) as executor:
         #   futures = {executor.submit(analyze, ip): ip for ip in ips}
          #  for future in as_completed(futures):
           #     results.append(future.result())

        #return results



class EmailManager:
    """Управление на имейл шаблони и изпращане"""
    """Разширен мениджър за имейли с Outlook поддръжка"""

    def __init__(self, config):
        self.config = config
        self.outlook_available = self._check_outlook_availability()
        
    def _check_outlook_availability(self):
        try:
            import win32com.client
            return True
        except ImportError:
            return False
        
    

    def generate_phishing_report(self, url, ip=None, template_type="standard"):
        """Генерира фишинг доклад според избрания шаблон"""
        templates = {
            "standard": self._standard_phishing_template,
            "gdbop": self._gdbop_template,
            "internal": self._internal_alert_template,
        }

        if template_type not in templates:
            template_type = "standard"

        return templates[template_type](url, ip)

    def _standard_phishing_template(self, url, ip=None):
        """Стандартен шаблон за фишинг доклад"""
        return f"""Subject: Phishing Report

Dear All,

We would like to inform you about a detected phishing attempt involving an impersonation of a legitimate banking website.

The phishing site is a replica of a well-known electronic banking portal and aims to deceive users into entering their credentials.

The phishing email contains the following URL: {url}
IP: {ip}, located in country, and belongs to isp

It redirects to the following effective URL: {url}
IP: {ip}, located in country, and belongs to isp

We kindly request that the phishing page be taken down as soon as possible to prevent further harm.

The phishing email includes the following content:

(Insert phishing email content here if applicable)

Best regards,
Cybersecurity Operations Center (SOC)

"""

    def _gdbop_template(self, url, ip=None):
        """Шаблон за доклад до ГДБОП"""
        return f"""Subject: Фишинг кампания към клиенти на Банка ДСК

До
Директора на
Главна дирекция „Борба с организираната престъпност“
Министерство на вътрешните работи

Уважаеми г-н Директор,

Уведомяваме Ви, че на {datetime.now().strftime("%Y-%m-%d")} г. бяха получени сигнали от потребители, информиращи за фишинг електронни съобщения, изпратени от името на българска финансова институция. Съобщенията наподобяват легитимна комуникация и съдържат подвеждащи връзки, насочващи към измамни уебсайтове.

От извършените проверки е установено следното:

Хипервръзката в електронното писмо води към следния URL адрес: {url}
IP адрес: {ip},локация [country], доставчик на интернет услуги: [isp]

Страницата препраща автоматично към друг URL адрес: {url}
IP адрес: {ip}, локализиран в [държава], с доставчик на интернет услуги: [ISP]

В тази връзка, молим за извършване на проверка относно изложените факти и обстоятелства.

При необходимост можем да предоставим допълнителна информация и доказателства.

С уважение,
Екип по информационна сигурност
Email: [ваш служебен имейл]
Телефон: [контакт за връзка]


"""

    def _internal_alert_template(self, url, ip=None):
        """Шаблон за вътрешно съобщение"""
        return f"""Subject: НОВА ФИШИНГ КАМПАНИЯ - 

Здравейте колеги, 

Установен е IP адрес, от който са извършвани логвания и вероятни злонамерени действия, във връзка с phishing атаки. 
Клиенти на организацията, най-вероятно са получили фишинг email и са въвели доброволно данни в измамна страница. 
В тази връзка, моля да блокирате профилите им и активните им банкови карти, след което да ги подадете към Кол центъра за осъществяване на контакт с клиентите.

{ip if ip else 'Неизвестен'}



Best Regards,

SOC
"""

    def send_email(self, recipient, subject, body, attachments=None, use_outlook=None):
        """Разширена функция за изпращане на имейли с Outlook поддръжка"""
        email_settings = self.config.get_email_settings()
        
        if use_outlook is None:
            use_outlook = email_settings.get("use_outlook", False) and self.outlook_available
        
        if use_outlook:
            return self._send_via_outlook(recipient, subject, body, attachments)
        else:
            return self._send_via_smtp(recipient, subject, body, attachments)
        
    def _send_via_outlook(self, recipient, subject, body, attachments=None):
        """Изпращане на имейл чрез Outlook"""
        try:
            import win32com.client
            outlook = win32com.client.Dispatch("Outlook.Application")
            mail = outlook.CreateItem(0)
            
            # Кодиране на кирилица
            mail.Subject = subject
            mail.Body = body
            mail.To = recipient
            
            if attachments:
                for attachment in attachments:
                    if os.path.exists(attachment):
                        mail.Attachments.Add(attachment)
            
            # Изпращане
            mail.Send()
            
            return {"success": True, "method": "outlook"}
        except Exception as e:
            return {"error": str(e), "method": "outlook"}
        
    def _send_via_smtp(self, recipient, subject, body, attachments=None):
        """Изпращане на имейл чрез SMTP"""
        email_settings = self.config.get_email_settings()
        
        if not email_settings["smtp_server"] or not email_settings["email_user"]:
            return {"error": "SMTP настройките не са конфигурирани"}
        
        try:
            msg = MIMEMultipart()
            msg["From"] = email_settings["email_user"]
            msg["To"] = recipient
            msg["Subject"] = subject
            
            # Кодиране на кирилица
            msg.attach(MIMEText(body, "plain", "utf-8"))
            
            if attachments:
                for attachment in attachments:
                    if os.path.exists(attachment):
                        with open(attachment, "rb") as f:
                            part = MIMEBase("application", "octet-stream")
                            part.set_payload(f.read())
                            encoders.encode_base64(part)
                            part.add_header(
                                "Content-Disposition",
                                f'attachment; filename="{os.path.basename(attachment)}"',
                            )
                            msg.attach(part)
            
            with smtplib.SMTP(
                email_settings["smtp_server"], email_settings["smtp_port"]
            ) as server:
                server.starttls()
                server.login(
                    email_settings["email_user"], email_settings["email_password"]
                )
                server.send_message(msg)
            
            return {"success": True, "method": "smtp"}
        except Exception as e:
            return {"error": str(e), "method": "smtp"}   
        
    def send_scan_results(self, recipient, subject, scan_type, results, attachments=None):
        """Изпраща резултати от сканиране по имейл"""
        email_settings = self.config.get_email_settings()
        
        # Генериране на тялото на имейла
        body = f"""Резултати от {scan_type} анализ:

{results}

Това е автоматично генериран имейл. Моля, не отговаряйте на него.
"""
        
        # Изпращане на имейла
        return self.send_email(
            recipient=recipient,
            subject=subject,
            body=body,
            attachments=attachments
        )


class SOCGUI:
    def run_traceroute(self):
        def traceroute_worker():
            ip = self.traceroute_entry.get().strip()
            self.traceroute_output.delete(1.0, tk.END)
            if not ip:
                self.traceroute_output.insert(tk.END, "Моля, въведете IP адрес или хост.")
                return

            system_platform = platform.system().lower()
            if system_platform == "windows":
                cmd = ["tracert", "-w", "6000", ip]
            else:
                cmd = ["traceroute", "-w", "6", "-q", "1", ip]

            try:
                creationflags = subprocess.CREATE_NO_WINDOW if platform.system().lower() == "windows" else 0
                result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, timeout=10, creationflags=creationflags)
                self.traceroute_output.insert(tk.END, result.stdout)
            except subprocess.TimeoutExpired:
                self.traceroute_output.insert(tk.END, "Traceroute командата изтече.")
            except Exception as e:
                self.traceroute_output.insert(tk.END, f"Грешка при traceroute: {e}")

        threading.Thread(target=traceroute_worker, daemon=True).start()


    def check_phishing(self, url):
        """Комбинирана проверка на фишинг URL"""
        results = {
            "url": url,
            "virus_total": {},
            "hybrid_analysis": {},
            "heuristic": {},
        }
    
        # VirusTotal проверка
        vt_result = self.security_tools.scan_url_virustotal(url)
        results["virus_total"] = vt_result
    
        # Hybrid Analysis – домейн от URL
        parsed = urllib.parse.urlparse(url)
        domain = parsed.netloc or url
        ha_result = self.security_tools.scan_hybrid_analysis(domain=domain)
        results["hybrid_analysis"] = ha_result
    
        # Евристичен анализ
        phishing_keywords = ["login", "secure", "verify", "update", "account", "signin", "bank"]
        matches = [kw for kw in phishing_keywords if kw in url.lower()]
        results["heuristic"] = {
            "keyword_matches": matches,
            "suspicious": len(matches) >= 2
        }
    
        return results
    def get_abuseipdb_blacklist(self):
        """Извлича AbuseIPDB blacklist и го показва в отделен прозорец."""
        try:
            result = self.security_tools.get_abuseipdb_blacklist()
            if "error" in result:
                messagebox.showerror("Грешка", f"AbuseIPDB грешка: {result['error']}")
                return
    
            output = "\n".join(result.get("blacklist", []))
            window = tk.Toplevel(self.root)
            window.title("AbuseIPDB Blacklist")
            text = scrolledtext.ScrolledText(window, width=100, height=20, font=("Consolas", 10))
            text.pack(fill=tk.BOTH, expand=True)
            text.insert(tk.END, output)
        except Exception as e:
            messagebox.showerror("Грешка", f"Неуспешно зареждане на blacklist: {str(e)}")
            
    def handle_report_phishing(self):
        url = self.phishing_url_entry.get().strip()
        if not url:
            messagebox.showwarning("Грешка", "Моля, въведете URL за докладване.")
            return
        self.log_activity(f"Докладване на фишинг URL: {url}")
    
        selected_services = ["google", "microsoft", "phishtank", "openphish", "apwg", "netcraft", "brightcloud"]
        for service in selected_services:
            result = self.security_tools.report_phishing(url, service, log=self.log_activity)
            if "success" in result:
                self.update_status(f"{service.capitalize()}: Успешно докладвано!")
            else:
                self.update_status(f"{service.capitalize()} грешка: {result['error']}")
    def handle_check_phishing(self):
        url = self.phishing_url_entry.get().strip()
        if not url:
            messagebox.showwarning("Грешка", "Моля, въведете URL за проверка.")
            return
    
        self.update_status(f"Проверка на фишинг URL: {url}")
        results = self.check_phishing(url)
    
        vt = results["virus_total"]
        ha = results["hybrid_analysis"]
        heur = results["heuristic"]
    
        vt_str = vt.get("detections_str", "Няма данни")
        vt_link = vt.get("permalink", "")
        ha_summary = "Намерени данни" if ha.get("data") else "Няма данни"
        heur_str = ", ".join(heur.get("keyword_matches", []))
        heur_flag = "⚠️ Подозрителен URL" if heur.get("suspicious") else "Няма директни индикации"
    
        output = (
            f"--- VirusTotal ---\n"
            f"Детекции: {vt_str}\n"
            f"Линк: {vt_link}\n\n"
            f"--- Hybrid Analysis ---\n"
            f"{ha_summary}\n\n"
            f"--- Евристика ---\n"
            f"Ключови думи: {heur_str}\n"
            f"{heur_flag}\n"
        )
    
        self.phishing_result_text.delete(1.0, tk.END)
        self.phishing_result_text.insert(tk.END, output)
    
        # Ако има над 5 детекции, предложи автоматично докладване
        try:
            detected, total = map(int, vt_str.split("/"))
            if detected >= 5:
                should_report = messagebox.askyesno(
                    "Открит фишинг!",
                    f"VirusTotal откри {detected} детекции от {total}.\n"
                    f"Искаш ли да го докладваме автоматично към Google?"
                )
                if should_report:
                    self.report_phishing_external("google")
        except:
            pass  # Ако форматът е неочакван – игнорирай
    """Графичен интерфейс на SOC инструмента"""
    """Основен GUI клас с всички подобрения"""

    def __init__(self, root):
        self.intel_result = None
        self.ip_analysis_tab = None
        self.url_analysis_tab = None
        self.file_analysis_tab = None
        self.mxtoolbox_tab = None
        self.batch_ip_tab = None
        self.network_tools_tab = None
        self.root = root
        self.config = ConfigManager()
        self.security_tools = SecurityTools(self.config, log_activity=self.log_activity)
        self.email_manager = EmailManager(self.config)
        self.ps_history = []
        self.ps_history = -1
        #self.is_admin = self.check_admin_privileges()
        self.is_admin = self.check_admin_privileges()
        self.ps_process = None
        self.security_tools = SecurityTools(self.config)
        self.osint_results = {}
        self.last_pastebin_results = []
        self.last_tld_results = []
        self.selected_ips = []
        self.selected_domains = []
        self.selected_hashes = []




        
        # Теми и езици
        self.themes = {
            'dark': {
                'bg': '#333333',
                'fg': '#ffffff',
                'accent': '#0078d7',
                'entry_bg': '#555555',
                'highlight': '#005b9f'
            },
            'light': {
                'bg': '#f5f5f5',
                'fg': '#000000',
                'accent': '#1e90ff',
                'entry_bg': '#ffffff',
                'highlight': '#b3d9ff'
            },
            'blue': {
                'bg': '#002b49',
                'fg': '#ffffff',
                'accent': '#4fc3f7',
                'entry_bg': '#004d80',
                'highlight': '#0086d1'
            },
            'green': {
                'bg': '#1a3a1a',
                'fg': '#ffffff',
                'accent': '#4CAF50',
                'entry_bg': '#2d5d2d',
                'highlight': '#3e8e3e'
            },
            'purple': {
                'bg': '#2a0a3a',
                'fg': '#ffffff',
                'accent': '#9C27B0',
                'entry_bg': '#4a1a5a',
                'highlight': '#7B1FA2'
            }
        }
        
        self.languages = {
            'bg': {
                'file_menu': 'Файл',
                'new': 'Нов',
                'open': 'Отвори',
                'save': 'Запази',
                'exit': 'Изход',
                'tools_menu': 'Инструменти',
                'settings': 'Настройки',
                'api_keys': 'API Ключове',
                'email_settings': 'Имейл Настройки',
                'test_connection': 'Тест на връзката',
                'help_menu': 'Помощ',
                'docs': 'Документация',
                'about': 'Относно'
            },
            'en': {
                'file_menu': 'File',
                'new': 'New',
                'open': 'Open',
                'save': 'Save',
                'exit': 'Exit',
                'tools_menu': 'Tools',
                'settings': 'Settings',
                'api_keys': 'API Keys',
                'email_settings': 'Email Settings',
                'test_connection': 'Test Connection',
                'help_menu': 'Help',
                'docs': 'Documentation',
                'about': 'About'
            }
        }
        
        self.current_theme = self.config.get_ui_settings()["theme"]
        self.current_language = self.config.get_ui_settings()["language"]
        
        self.setup_ui()
        self.setup_menus()
        self.setup_styles()
        self.setup_context_menus()
        # Променливи за състояние
        self.current_file = None
        self.batch_ip_results = []
        self.ip_queue = Queue()
        self.processing = False
        self.active_threads = 0

    def setup_ui(self):
        # """Инициализира потребителския интерфейс"""
        #"""Инициализира потребителския интерфейс с всички подобрения"""
        self.root.title("SOC Tool v2.1")
        self.root.geometry("1400x900")

        # Основен контейнер
        self.main_frame = ttk.Frame(self.root)
        self.main_frame.pack(fill=tk.BOTH, expand=True)

        # Ноутбук с раздели
        self.notebook = ttk.Notebook(self.main_frame)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Създаване на раздели
        self.create_dashboard_tab()
        self.create_ip_analysis_tab()
        self.create_batch_ip_tab()
        self.create_url_analysis_tab()
        self.create_file_analysis_tab()
        self.create_phishing_tab()
        self.create_email_analysis_tab()
        self.create_network_tools_tab()
        self.create_threat_intel_tab()
        self.create_mxtoolbox_tab()
        self.create_powershell_tab()
        self.setup_ad_tab()
        self.create_osint_tab()
        



        # Статус бар
        self.status_frame = ttk.Frame(self.main_frame)
        self.status_frame.pack(fill=tk.X, side=tk.BOTTOM)
        
        self.status_bar = ttk.Label(
            self.status_frame, 
            text="Готов", 
            relief=tk.SUNKEN, 
            anchor=tk.W,
            style='Status.TLabel'
        )
        self.status_bar.pack(side=tk.LEFT, fill=tk.X)
        
        self.thread_count_label = ttk.Label(
            self.status_frame,
            text="Нишки: 0",
            relief=tk.SUNKEN,
            width=10,
            style='Status.TLabel'
        )
        self.thread_count_label.pack(side=tk.RIGHT)

        # Зареждане на икона
        self.load_icon()
        
        # Възстановяване на последно използвания таб
        last_tab = self.config.get_ui_settings()["last_used_tab"]
        if last_tab < len(self.notebook.tabs()):
            self.notebook.select(last_tab)

    def setup_styles(self):
        """Конфигурира стиловете на интерфейса"""
        """Конфигурира стиловете на интерфейса с теми"""
        
        self.style = ttk.Style()
        #ui_settings = self.config.get_ui_settings()
        
        # Прилагане на текущата тема
        theme = self.themes[self.current_theme]
        

        #if ui_settings["theme"] == "dark":
        #   self.style.theme_use("clam")
        #    bg_color = "#333333"
        #    fg_color = "#ffffff"
        #    entry_bg = "#555555"
        #else:
        #    bg_color = "#f0f0f0"
        #    fg_color = "#000000"
        #    entry_bg = "#ffffff"

            # Общи стилове
        self.style.theme_use('clam')
        self.style.configure('.', 
            background=theme['bg'],
            foreground=theme['fg'],
            fieldbackground=theme['entry_bg'],
            selectbackground=theme['accent'],
            selectforeground=theme['fg'],
            font=(self.config.get_ui_settings()["font"], 
                 self.config.get_ui_settings()["font_size"])
        )

        # Стилове за конкретни елементи
        self.style.configure('TFrame', background=theme['bg'])
        self.style.configure('TLabel', background=theme['bg'], foreground=theme['fg'])
        self.style.configure('TButton', padding=5)
        self.style.configure('TEntry', fieldbackground=theme['entry_bg'])
        self.style.configure('TCombobox', fieldbackground=theme['entry_bg'])
        self.style.configure('TNotebook', background=theme['bg'])
        self.style.configure('TNotebook.Tab', 
            padding=[10, 5],
            background=theme['bg'],
            foreground=theme['fg']
        )
        self.style.configure('Header.TLabel', 
            font=(self.config.get_ui_settings()["font"], 14, "bold"),
            foreground=theme['accent']
        )
        self.style.configure('Status.TLabel',
            relief=tk.SUNKEN,
            anchor=tk.W,
            background=theme['highlight'],
            foreground=theme['fg']
        )
        
        self.style.configure("Dark.TFrame", background="#2e2e2e")
        
        # Стилове за Treeview
        self.style.configure('Treeview',
            background=theme['entry_bg'],
            foreground=theme['fg'],
            fieldbackground=theme['entry_bg'],
            rowheight=25
        )
        self.style.map('Treeview',
            background=[('selected', theme['accent'])],
            foreground=[('selected', theme['fg'])]
        )
        # Прилагане на стиловете към root
        self.root.config(bg=theme['bg'])

    def setup_status_bar(self):
        self.status_bar = ttk.Label(self.root, text="Готово", anchor="w")
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
    def update_status(self, message):
        """Актуализира статус бара с допълнителна информация"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.status_bar.config(text=f"{timestamp} - {message}")
        self.root.update_idletasks()
        
    def background_tasks(self):
        """Фонови задачи за актуализиране на GUI"""
        while True:
            self.thread_count_label.config(text=f"Нишки: {threading.active_count()}")
            time.sleep(1)
            
    def on_tab_changed(self, event):
        """Запазва избрания таб при промяна"""
        selected_tab = self.notebook.index(self.notebook.select())
        self.config.config.set("UI", "LAST_USED_TAB", str(selected_tab))
        with open(CONFIG_FILE, "w") as configfile:
            self.config.config.write(configfile)
    
    def load_icon(self):
        """Зарежда икона на приложението"""
        try:
            icon_data = self.config.config.get("UI", "ICON_BASE64", fallback="")
            if icon_data:
                with open(ICON_FILE, "wb") as f:
                    f.write(base64.b64decode(icon_data))
                self.root.iconbitmap(ICON_FILE)
                os.remove(ICON_FILE)
        except Exception as e:
            print(f"Грешка при зареждане на икона: {e}")

    def setup_menus(self):
        """Създава главното меню"""
        menubar = tk.Menu(self.root)
        lang = self.languages[self.current_language]
        
        # Меню Файл
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Нов", command=self.new_analysis)
        file_menu.add_command(label="Отвори", command=self.open_file)
        file_menu.add_command(label="Запази", command=self.save_results)
        file_menu.add_separator()
        file_menu.add_command(label="Изход", command=self.root.quit)
        menubar.add_cascade(label="Файл", menu=file_menu)

        # Меню Инструменти
        tools_menu = tk.Menu(menubar, tearoff=0)
        tools_menu.add_command(label="Настройки", command=self.open_settings)
        tools_menu.add_command(label="API Ключове", command=self.open_api_settings)
        tools_menu.add_command(
            label="Имейл Настройки", command=self.open_email_settings
        )
        tools_menu.add_separator()
        tools_menu.add_command(label="Тест на връзката", command=self.test_connection)
        menubar.add_cascade(label="Инструменти", menu=tools_menu)

        # Меню Помощ
        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="Документация", command=self.show_documentation)
        help_menu.add_command(label="Относно", command=self.show_about)
        menubar.add_cascade(label="Помощ", menu=help_menu)

        self.root.config(menu=menubar)
        
    def export_ip_results(self):
        rows = [self.ip_treeview.item(row)['values'] for row in self.ip_treeview.get_children()]
        if not rows:
            messagebox.showinfo('Експорт', 'Няма данни за експортиране.')
            return
        file = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV файлове", "*.csv"), ("Всички файлове", "*.*")])
        if file:
            with open(file, "w", encoding="utf-8") as f:
                for row in rows:
                    f.write(",".join(map(str, row)) + "\n")
            messagebox.showinfo("Готово", f"Резултатите са записани в: {file}")

    def setup_context_menus(self):
        """Създава контекстни менюта за различни елементи"""
        # Контекстно меню за IP резултати
        self.ip_context_menu = Menu(self.ip_treeview, tearoff=0)
        self.ip_context_menu.add_command(
            label="Копирай", command=self.copy_selected_ip_row
        )
        self.ip_context_menu.add_command(
            label="Изтрий", command=self.delete_selected_ip_row
        )
        self.ip_context_menu.add_command(
            label="Експортирай", command=self.export_ip_results
        )
        self.ip_treeview.bind("<Button-3>", self.show_ip_context_menu)

        # Контекстно меню за текстови резултати
        self.text_context_menu = Menu(self.ip_text_results, tearoff=0)
        self.text_context_menu.add_command(
            label="Копирай", command=self.copy_selected_text
        )
        self.ip_text_results.bind("<Button-3>", self.show_text_context_menu)

        # Контекстно меню за пакетна IP проверка
        self.batch_ip_context_menu = Menu(self.batch_ip_treeview, tearoff=0)
        self.batch_ip_context_menu.add_command(
            label="Копирай", command=self.copy_batch_ip_row
        )
        self.batch_ip_context_menu.add_command(
            label="Копирай избрани", command=self.copy_selected_batch_rows
        )
        self.batch_ip_context_menu.add_command(
            label="Изтрий", command=self.delete_batch_ip_row
        )
        self.batch_ip_context_menu.add_command(
            label="Изтрий избрани", command=self.delete_selected_batch_rows
        )
        self.batch_ip_context_menu.add_command(
            label="Изтрий всички", command=self.delete_all_batch_ip_rows
        )
        self.batch_ip_context_menu.add_command(
            label="Докладвай до AbuseIPDB", command=self.report_selected_to_abuseipdb
        )
        self.batch_ip_treeview.bind("<Button-3>", self.show_batch_ip_context_menu)
        
        self.batch_ip_context_menu.add_command(
            label="Избери всички", command=lambda: self.select_all_batch_rows()
        )
        

        # Добавяне на клавишни комбинации за копиране
        self.root.bind_all("<Control-c>", self.handle_copy_shortcut)
        self.root.bind_all("<Control-C>", self.handle_copy_shortcut)
        
    def handle_copy_shortcut(self, event):
        """Обработва клавишната комбинация за копиране"""
        widget = self.root.focus_get()
        
        if widget == self.ip_treeview:
            self.copy_selected_ip_row()
        elif widget == self.batch_ip_treeview:
            self.copy_batch_ip_row()
        elif isinstance(widget, tk.Text):
            if widget.tag_ranges(tk.SEL):
                selected_text = widget.get(tk.SEL_FIRST, tk.SEL_LAST)
                self.root.clipboard_clear()
                self.root.clipboard_append(selected_text)

    def show_ip_context_menu(self, event):
        """Показва контекстното меню за IP резултати"""
        try:
            self.ip_context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.ip_context_menu.grab_release()

    def show_text_context_menu(self, event):
        """Показва контекстното меню за текстови резултати"""
        try:
            self.text_context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.text_context_menu.grab_release()

    def show_batch_ip_context_menu(self, event):
        """Показва контекстното меню за пакетна IP проверка"""
        try:
            self.batch_ip_context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.batch_ip_context_menu.grab_release()

    def copy_selected_ip_row(self):
        """Копира избрания ред от IP дървото във формат IP - Държава - ISP"""
        selected = self.ip_treeview.selection()
        if selected:
            values = self.ip_treeview.item(selected[0], "values")
            ip = values[0]
            country = values[1] if len(values) > 1 else "N/A"
            isp = values[2] if len(values) > 2 else "N/A"
            country = country.split('(')[0].strip() if '(' in country else country
            formatted = f"{ip} - {country} - {isp}"
            self.root.clipboard_clear()
            self.root.clipboard_append(formatted)
            
    def delete_selected_ip_row(self):
        """Изтрива избрания ред от IP дървото"""
        selected = self.ip_treeview.selection()
        if selected:
            self.ip_treeview.delete(selected[0])
            
    def copy_selected_text(self):
        """Копира избрания текст"""
        if self.ip_text_results.tag_ranges(tk.SEL):
            selected_text = self.ip_text_results.get(tk.SEL_FIRST, tk.SEL_LAST)
            self.root.clipboard_clear()
            self.root.clipboard_append(selected_text)

    def copy_batch_ip_row(self):
        """Копира избрания ред от пакетна IP проверка във формат IP - Държава - ISP"""
        selected = self.batch_ip_treeview.selection()
        if selected:
            values = self.batch_ip_treeview.item(selected[0], "values")
            ip = values[0]
            country = values[1] if len(values) > 1 else "N/A"
            isp = values[2] if len(values) > 2 else "N/A"
            country = country.split('(')[0].strip() if '(' in country else country
            formatted = f"{ip} - {country} - {isp}"
            self.root.clipboard_clear()
            self.root.clipboard_append(formatted)

    def copy_selected_batch_rows(self):
        """Копира избраните редове във формат IP - Държава - ISP"""
        selected = self.batch_ip_treeview.selection()
        if not selected:
            return

        rows = []
        for row in selected:
            values = self.batch_ip_treeview.item(row)["values"]
            ip = values[0]
            country = values[1] if len(values) > 1 else "N/A"
            isp = values[2] if len(values) > 2 else "N/A"
            country = country.split('(')[0].strip() if '(' in country else country
            rows.append(f"{ip} - {country} - {isp}")

        self.root.clipboard_clear()
        self.root.clipboard_append("\n".join(rows))
        self.root.update()

    def delete_selected_batch_rows(self):
        """Изтрива избраните редове от дървото"""
        for row in self.batch_ip_treeview.selection():
            self.batch_ip_treeview.delete(row)

    def select_all_batch_rows(self):
        """Избира всички редове в дървото за пакетна IP проверка"""
        for row in self.batch_ip_treeview.get_children():
            self.batch_ip_treeview.selection_add(row)

    def delete_all_batch_ip_rows(self):
        """Изтрива всички редове от пакетна IP проверка"""
        if messagebox.askyesno("Потвърждение", "Сигурни ли сте, че искате да изтриете всички редове?"):
            self.batch_ip_treeview.delete(*self.batch_ip_treeview.get_children())
            
    def delete_batch_ip_row(self):
        """Изтрива избрания ред от пакетна IP проверка"""
        selected = self.batch_ip_treeview.selection()
        if selected:
            self.batch_ip_treeview.delete(selected[0])
    def report_selected_to_abuseipdb(self):
        """Докладва избрания IP към AbuseIPDB"""
        selected = self.batch_ip_treeview.selection()
        if not selected:
            return
            
        ip = self.batch_ip_treeview.item(selected[0], "values")[0]
        
        # Прозорец за избор на категории
        report_window = tk.Toplevel(self.root)
        report_window.title(f"Докладване на {ip} към AbuseIPDB")
        
        categories = [
            ("DNS Compromise", 1),
            ("DNS Poisoning", 2),
            ("Fraud Orders", 3),
            ("DDoS Attack", 4),
            ("FTP Brute-Force", 5),
            ("Ping of Death", 6),
            ("Phishing", 7),
            ("Fraud VoIP", 8),
            ("Open Proxy", 9),
            ("Web Spam", 10),
            ("Email Spam", 11),
            ("Blog Spam", 12),
            ("VPN IP", 13),
            ("Port Scan", 14),
            ("Hacking", 15),
            ("SQL Injection", 16),
            ("Spoofing", 17),
            ("Brute-Force", 18),
            ("Bad Web Bot", 19),
            ("Exploited Host", 20),
            ("Web App Attack", 21),
            ("SSH", 22),
            ("IoT Targeted", 23)
        ]
        
        selected_categories = []
        
        def toggle_category(cat_num, var):
            if var.get():
                selected_categories.append(cat_num)
            else:
                if cat_num in selected_categories:
                    selected_categories.remove(cat_num)
        
        frame = ttk.Frame(report_window)
        frame.pack(padx=10, pady=10)
        
        ttk.Label(frame, text="Изберете категории:").pack(anchor=tk.W)
        
        for text, cat_num in categories:
            var = tk.BooleanVar()
            cb = ttk.Checkbutton(frame, text=text, variable=var,
                               command=lambda n=cat_num, v=var: toggle_category(n, v))
            cb.pack(anchor=tk.W)
        
        ttk.Label(frame, text="Коментар:").pack(anchor=tk.W)
        comment_entry = ttk.Entry(frame, width=40)
        comment_entry.pack(fill=tk.X)
        
        def submit_report():
            if not selected_categories:
                messagebox.showwarning("Грешка", "Моля, изберете поне една категория")
                return
                
            result = self.security_tools.report_to_abuseipdb(
                ip=ip,
                categories=selected_categories,
                comment=comment_entry.get()
            )
            
            if "error" in result:
                messagebox.showerror("Грешка", f"Грешка при докладване: {result['error']}")
            else:
                messagebox.showinfo("Успех", "IP адресът е докладван успешно")
            
            report_window.destroy()
        
        ttk.Button(frame, text="Докладвай", command=submit_report).pack(pady=10)

    def create_dashboard_tab(self):
        """Създава таб с dashboard"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Dashboard")

        # Заглавен ред
        header = ttk.Frame(tab)
        header.pack(fill=tk.X, pady=10)
        ttk.Label(header, text="SOC Tool Dashboard", style="Header.TLabel").pack()

        # Бързи инструменти
        quick_tools = ttk.Frame(tab)
        quick_tools.pack(fill=tk.X, pady=10)

        tools = [
            ("IP Проверка", self.show_ip_tools),
            ("URL Сканиране", self.show_url_tools),
            ("Файлов Анализ", self.show_file_tools),
            ("Мрежови Инструменти", self.show_network_tools),
            ("Пакетна IP Проверка", self.show_batch_ip_tools),
            ("MX Toolbox", self.show_mxtoolbox),
            ("PowerShell", self.show_powershell_tab),
            ("Email Threat Analysis", self.show_email_analysis_tab),
            ("Active Directory", self.show_ad_tab),
            ("OSINT", self.show_osint_tab),
            ("Threat Inteligence", self.show_threat_inteligence_tab)
        ]

        for i, (text, cmd) in enumerate(tools):
            btn = ttk.Button(quick_tools, text=text, command=cmd)
            btn.grid(row=i // 3, column=i % 3, padx=5, pady=5, sticky=tk.EW)

        # Последни активности
        activity_frame = ttk.LabelFrame(tab, text="Последни активности", padding=10)
        activity_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.activity_log = scrolledtext.ScrolledText(
            activity_frame, wrap=tk.WORD, font=("Consolas", 10)
        )
        self.activity_log.pack(fill=tk.BOTH, expand=True)
    
    def show_ip_tools(self):
        self.notebook.select(self.ip_analysis_tab)
        self.log_activity("Пренасочване към IP Анализ")

    def show_url_tools(self):
        self.notebook.select(self.url_analysis_tab)
        self.log_activity("Пренасочване към URL Анализ")

    def show_file_tools(self):
        self.notebook.select(self.file_analysis_tab)
        self.log_activity("Пренасочване към Файлов Анализ")

    def show_network_tools(self):
        self.notebook.select(self.network_tools_tab)
        self.log_activity("Пренасочване към Мрежови Инструменти")

    def show_batch_ip_tools(self):
        self.notebook.select(self.batch_ip_tab)
        self.log_activity("Пренасочване към Пакетна IP Проверка")

    def show_mxtoolbox(self):
        self.notebook.select(self.mxtoolbox_tab)
        self.log_activity("Пренасочване към MX Toolbox")

    def show_powershell_tab(self):
        self.notebook.select(self.powershell_tab)
        self.log_activity("Пренасочване към PowerShell")

    def show_ad_tab(self):
        self.notebook.select(self.ad_tab)
        self.log_activity("Пренасочване към Active Directory")

    def show_osint_tab(self):
        self.notebook.select(self.osint_tab)
        self.log_activity("Пренасочване към OSINT")

    def show_email_analysis_tab(self):
        self.notebook.select(self.email_analysis_tab)
        self.log_activity("Пренасочване към Email Analysis")

    def show_threat_inteligence_tab(self):
        self.notebook.select(self.threat_intel_tab)
        self.log_activity("Пренасочване към Threat Inelience")

    def show_phishing_report_tab(self):
        self.notebook.select(self.phishing_tab)
        self.log_activity("Пренасочване към Фишинг Доклади")

        
    def log_activity(self, message):
        """Добавя събитие към историята в Dashboard"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.activity_log.insert(tk.END, f"[{timestamp}] {message}\n")
        self.activity_log.see(tk.END)

    def create_ip_analysis_tab(self):
        """Създава таб за IP анализ"""
        tab = ttk.Frame(self.notebook)
        self.ip_analysis_tab = tab
        self.notebook.add(tab, text="IP Анализ")

        # Входни данни
        input_frame = ttk.Frame(tab)
        input_frame.pack(fill=tk.X, padx=10, pady=10)

        ttk.Label(input_frame, text="IP Адрес:").pack(side=tk.LEFT)
        self.ip_entry = ttk.Entry(input_frame)
        self.ip_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)

        # Бутони за анализ
        btn_frame = ttk.Frame(tab)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)

        analysis_services=(
            ("AbuseIPDB", "abuseipdb"),
            ("VirusTotal", "virustotal"),
            ("Shodan", "shodan"),
            ("GeoIP", "geo"),
            ("WHOIS", "whois"),
            ("Всички", "all"),)

        for i, (text, service) in enumerate(analysis_services):
            btn = ttk.Button(
                btn_frame,
                text=text,
                command=lambda s=service: self.analyze_ip_service(s),
            )
            btn.grid(row=0, column=i, padx=2, sticky=tk.EW)

        # Резултати
        results_frame = ttk.Frame(tab)
        results_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.ip_result_notebook = ttk.Notebook(results_frame)
        self.ip_result_notebook.pack(fill=tk.BOTH, expand=True)

        # Текстови резултати
        text_tab = ttk.Frame(self.ip_result_notebook)
        self.ip_result_notebook.add(text_tab, text="Текст")
        self.ip_text_results = scrolledtext.ScrolledText(
            text_tab, wrap=tk.WORD, font=("Consolas", 10)
        )
        self.ip_text_results.pack(fill=tk.BOTH, expand=True)

        # Таблични резултати
        table_tab = ttk.Frame(self.ip_result_notebook)
        self.ip_result_notebook.add(table_tab, text="Таблица")

        self.ip_treeview = ttk.Treeview(
            table_tab, columns=("Свойство", "Стойност"), show="headings"
        )
        self.ip_treeview.heading("Свойство", text="Свойство")
        self.ip_treeview.heading("Стойност", text="Стойност")
        self.ip_treeview.column("Свойство", width=200)
        self.ip_treeview.column("Стойност", width=400)

        scrollbar = ttk.Scrollbar(
            table_tab, orient="vertical", command=self.ip_treeview.yview
        )
        self.ip_treeview.configure(yscrollcommand=scrollbar.set)

        self.ip_treeview.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Докладване
        report_frame = ttk.Frame(tab)
        report_frame.pack(fill=tk.X, pady=5)

        report_sites = [
            ("Spamhaus", "https://check.spamhaus.org/results?query="),
            ("AbuseIPDB", "https://www.abuseipdb.com/report?ip="),
            ("CINSscore", "https://cinsscore.com/"),
            ("IPvoid", "https://www.ipvoid.com/"),
        ]

        for i, (text, url) in enumerate(report_sites):
            btn = ttk.Button(
                report_frame,
                text=text,
                command=lambda u=url: self.report_ip_site(u),
            )
            btn.grid(row=0, column=i, padx=2, sticky=tk.EW)

    def analyze_ip_service(self, service):
        """Анализира IP адрес с избраната услуга"""
        ip = self.ip_entry.get().strip()
        if not ip:
            messagebox.showwarning("Грешка", "Моля, въведете IP адрес")
            return
        self.log_activity(f"Стартиран IP анализ за {ip} чрез {service}")

        if service == "all":
            self.analyze_ip_all_services()
            return

        self.status_bar.config(text=f"Проверява {service}...")
        results = self.security_tools.analyze_ip(ip, services=(service,))

        # Показване на резултати в текстовия изглед
        self.ip_text_results.delete(1.0, tk.END)
        self.ip_text_results.insert(tk.END, f"\n=== {service.upper()} ===\n")
        if "error" in results.get(service, {}):
            self.ip_text_results.insert(
                tk.END, f"Грешка: {results[service]['error']}\n"
            )
        else:
            for key, value in results.get(service, {}).items():
                if isinstance(value, dict):
                    self.ip_text_results.insert(tk.END, f"{key}:\n")
                    for subkey, subvalue in value.items():
                        self.ip_text_results.insert(
                            tk.END, f"  {subkey}: {subvalue}\n"
                        )
                else:
                    self.ip_text_results.insert(tk.END, f"{key}: {value}\n")
                    
        # Показване на резултати в табличния изглед
        self.ip_treeview.delete(*self.ip_treeview.get_children())
        if "error" not in results.get(service, {}):
            for key, value in results[service].items():
                if isinstance(value, dict):
                    for subkey, subvalue in value.items():
                        self.ip_treeview.insert(
                            "",
                            tk.END,
                            values=(f"{service}.{key}.{subkey}", str(subvalue)),
                        )
                else:
                    self.ip_treeview.insert(
                        "", tk.END, values=(f"{service}.{key}", str(value))
                    )

        self.status_bar.config(text=f"Готово - проверка на {ip} с {service}")

    def analyze_ip_all_services(self):
        """Анализира IP адрес с всички услуги"""
        ip = self.ip_entry.get().strip()
        if not ip:
            messagebox.showwarning("Грешка", "Моля, въведете IP адрес")
            return
        self.log_activity(f"Стартиранa IP проверка за {ip}")

        self.status_bar.config(text=f"Проверява всички услуги за {ip}...")
        self.root.update()

        services=(
            "abuseipdb",
            "virustotal",
            "shodan",
            "geo",
            "geojs",
            "ipapi",
            "ipgeolocation",
            "whois",)
        results = self.security_tools.analyze_ip(ip, services=services)

        # Показване на резултати в текстовия изглед
        self.ip_text_results.delete(1.0, tk.END)
        for service, data in results.items():
            self.ip_text_results.insert(tk.END, f"\n=== {service.upper()} ===\n")
            if "error" in data:
                self.ip_text_results.insert(tk.END, f"Грешка: {data['error']}\n")
            else:
                for key, value in data.items():
                    if isinstance(value, dict):
                        self.ip_text_results.insert(tk.END, f"{key}:\n")
                        for subkey, subvalue in value.items():
                            self.ip_text_results.insert(
                                tk.END, f"  {subkey}: {subvalue}\n"
                            )
                    else:
                        self.ip_text_results.insert(tk.END, f"{key}: {value}\n")

        # Показване на резултати в табличния изглед
        self.ip_treeview.delete(*self.ip_treeview.get_children())
        for service, data in results.items():
            if "error" not in data:
                for key, value in data.items():
                    if isinstance(value, dict):
                        for subkey, subvalue in value.items():
                            self.ip_treeview.insert(
                                "",
                                tk.END,
                                values=(f"{service}.{key}.{subkey}", str(subvalue)),
                            )
                    else:
                        self.ip_treeview.insert(
                            "", tk.END, values=(f"{service}.{key}", str(value))
                        )

        self.status_bar.config(text=f"Готово - проверка на {ip} завършена")

    def report_ip_site(self, url):
        """Отваря сайт за докладване на IP"""
        ip = self.ip_entry.get().strip()
        if not ip:
            messagebox.showwarning("Грешка", "Моля, въведете IP адрес")
            return

        try:
            report_url = f"{url}{ip}"
            webbrowser.open(report_url)
            self.status_bar.config(
                text=f"Отворен е сайтът за докладване: {report_url}"
            )
        except Exception as e:
            messagebox.showerror("Грешка", f"Неуспешно отваряне на сайта: {str(e)}")

    def create_batch_ip_tab(self):
        """Създава таб за пакетна проверка на IP адреси"""
        tab = ttk.Frame(self.notebook)
        self.batch_ip_tab = tab
        self.notebook.add(tab, text="Пакетна IP Проверка")

        # Входни данни
        input_frame = ttk.Frame(tab)
        input_frame.pack(fill=tk.X, padx=10, pady=10)

        ttk.Label(
            input_frame,
            text="IP Адреси (по един на ред или разделени със запетая):",
        ).pack(side=tk.LEFT)
        self.batch_ip_text = scrolledtext.ScrolledText(
            input_frame, width=50, height=7, wrap=tk.WORD
        )
        self.batch_ip_text.pack(side=tk.LEFT, padx=5, fill=tk.BOTH, expand=True)

        # Бутони за анализ
        btn_frame = ttk.Frame(tab)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Button(
            btn_frame, text="Провери IP адреси", command=self.analyze_batch_ips
        ).pack(side=tk.LEFT, padx=5)

        ttk.Button(
            btn_frame, text="Копирай резултати", command=self.copy_batch_results
        ).pack(side=tk.LEFT, padx=5)

        ttk.Button(
            btn_frame, text="Експортирай в CSV", command=self.export_batch_results
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            btn_frame, text="Импортирай от файл", command=self.import_ips_from_file
        ).pack(side=tk.LEFT, padx=5)

        ttk.Button(
            btn_frame, text="Изчисти дублирани", command=self.remove_duplicate_ips
        ).pack(side=tk.LEFT, padx=5)

        ttk.Button(
            btn_frame, text="Спри проверка", command=self.stop_batch_processing
        ).pack(side=tk.LEFT, padx=5)
        
        # Допълнителни бутони за специфични проверки
        specific_checks_frame = ttk.Frame(tab)
        specific_checks_frame.pack(fill=tk.X, padx=10, pady=5)

        specific_checks = [
            ("AbuseIPDB", "abuseipdb"),
            ("GeoIP", "geo"),
            ("IP API", "ipapi"),
            ("GeoJS", "geojs"),
            ("IP Geolocation", "ipgeolocation"),
            ("IPinfo", "ipinfo"),
            ("FindIP", "findip"),
            ("IP2Location", "ip2location"),
            ("Maxmind", "maxmind")
            
        ]

        for i, (text, service) in enumerate(specific_checks):
            btn = ttk.Button(
                specific_checks_frame,
                text=text,
                command=lambda s=service: self.analyze_batch_ips(s)
            )
            btn.grid(row=0, column=i, padx=2, sticky=tk.EW)

        # Резултати
        results_frame = ttk.Frame(tab)
        results_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        columns = ("IP", "Държава", "ISP") #, "Abuse Score", "Shodan Порт", "Детайли")
        self.batch_ip_treeview = ttk.Treeview(
            results_frame, 
            columns=columns, 
            show="headings"
            
        )
        self.batch_ip_treeview.heading("IP", text="IP")
        self.batch_ip_treeview.heading("Държава", text="Държава")
        self.batch_ip_treeview.heading("ISP", text="ISP")

        for col in columns:
            self.batch_ip_treeview.heading(col, text=col)
            self.batch_ip_treeview.column(col, width=120, stretch=True)


        scrollbar = ttk.Scrollbar(
            results_frame, orient="vertical", command=self.batch_ip_treeview.yview
        )
        self.batch_ip_treeview.configure(yscrollcommand=scrollbar.set)

        self.batch_ip_treeview.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Прогрес бар
        self.progress = ttk.Progressbar(
            tab, orient=tk.HORIZONTAL, mode="determinate"
        )
        self.progress.pack(fill=tk.X, padx=10, pady=5)

    def remove_duplicate_ips(self):
        """Премахва дублиращи се IP адреси от текстовото поле"""
        ip_text = self.batch_ip_text.get("1.0", tk.END).strip()
        if not ip_text:
            return

        # Разделяме IP адресите и премахваме дубликатите
        ips = []
        for line in ip_text.split("\n"):
            for ip in line.split(","):
                ip = ip.strip()
                if ip and ip not in ips:
                    ips.append(ip)

        # Обновяваме текстовото поле
        self.batch_ip_text.delete("1.0", tk.END)
        self.batch_ip_text.insert(tk.END, "\n".join(ips))
        self.status_bar.config(text=f"Останали са {len(ips)} уникални IP адреса")

    async def async_analyze_batch_ips(self, ips, force_source=None):
        """Асинхронна обработка на IP адреси"""
        self.batch_ip_treeview.delete(*self.batch_ip_treeview.get_children())
        self.progress["maximum"] = len(ips)
        self.progress["value"] = 0

        proxy_settings = self.config.get_proxy_settings()
        results = await self.security_tools.async_batch_ip_analysis(
            ips, 
            [force_source] if force_source else None,
            proxy_settings
        )
        
        # Определяме услугите според източника
        services = []
        if force_source:
            services = [force_source]
        else:
            # Автоматично избираме най-бързите услуги
            if self.config.get_api_key("ABUSEIP_API_KEY"):
                services.append("abuseipdb")
            else:
                services.extend(["ipapi", "ipgeolocation", "geojs"])
        
        # Извличаме прокси настройките
        #proxy_settings = self.config.get_proxy_settings()
        
        # Стартираме асинхронната проверка
        #results = await self.security_tools.async_batch_ip_analysis(ips, services, proxy_settings)
        
        # Показваме резултатите
        for ip, data in results:
            country = data.get('country', 'N/A')
            isp = data.get('isp', 'N/A')
            source = data.get('source', 'N/A')
            
            # По-добра обработка на държавата
            if country == 'N/A' and source != 'N/A':
                country = f"N/A ({source})"
            
            self.batch_ip_treeview.insert("", tk.END, values=(
                ip,
                country,
                isp,
                source
            ))
            self.progress["value"] += 1
            self.root.update()
        return results

    def analyze_batch_ips(self, force_source=None):
        """Стартира асинхронна обработка на IP адреси"""
        if self.processing:
            return
        
        ip_text = self.batch_ip_text.get("1.0", tk.END).strip()
        if not ip_text:
            messagebox.showwarning("Грешка", "Моля, въведете IP адреси")
            return
        
        ips = []
        for line in ip_text.split("\n"):
            for ip in line.split(","):
                ip = ip.strip()
                if ip:
                    try:
                        ipaddress.ip_address(ip)
                        ips.append(ip)
                    except ValueError:
                        continue
        
        if not ips:
            messagebox.showwarning("Грешка", "Не са намерени валидни IP адреси")
            return
        
        if len(ips) > MAX_IP_PER_BATCH:
            ips = ips[:MAX_IP_PER_BATCH]
            messagebox.showwarning("Предупреждение", f"Броят на IP адресите е ограничен до {MAX_IP_PER_BATCH}")
        
        self.processing = True
        self.status_bar.config(text=f"Анализиране на {len(ips)} IP адреси...")
        
        # Стартираме асинхронната задача
        threading.Thread(
            target=self.run_async_batch_ips,
            args=(ips, force_source),
            daemon=True
        ).start()

    def run_async_batch_ips(self, ips, force_source):
        """Стартира асинхронната задача от отделна нишка"""
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        try:
            results = loop.run_until_complete(
                self.async_analyze_batch_ips(ips, force_source)
            )
            self.batch_ip_results = results
            self.status_bar.config(text=f"Готово - анализирани {len(ips)} IP адреси")
        except Exception as e:
            messagebox.showerror("Грешка", f"Грешка при обработка: {str(e)}")
        finally:
            self.processing = False
            loop.close()

    def process_batch_ips(self, ips, force_source=None):
        """Обработва списък с IP адреси в отделна нишка"""
        try:
            with concurrent.futures.ThreadPoolExecutor(
                max_workers=self.config.get_max_threads()
            ) as executor:
                future_to_ip = {
                    executor.submit(self.analyze_single_batch_ip, ip, force_source): ip 
                    for ip in ips
                }
                
                for future in concurrent.futures.as_completed(future_to_ip):
                    ip = future_to_ip[future]
                    try:
                        result = future.result()
                        # Почистване на държавата (премахване на кода в скоби)
                        country = result.get("country", "N/A")
                        if '(' in country:
                            country = country.split('(')[0].strip()
                        
                        self.batch_ip_treeview.insert("", tk.END, values=(
                            ip,
                            country,
                            result.get("isp", "N/A")
                        ))
                    except Exception as e:
                        self.batch_ip_treeview.insert("", tk.END, values=(
                            ip, "Грешка", str(e)
                        ))
                    
                    self.progress["value"] += 1
                    self.root.update()
                    
            self.status_bar.config(text=f"Готово - анализирани {len(ips)} IP адреси")
        finally:
            self.processing = False

    def stop_batch_processing(self):
        """Спира текущата обработка на IP адреси"""
        if self.processing:
            self.processing = False
            self.status_bar.config(text="Обработката на IP адреси е спряна")
        else:
            messagebox.showinfo("Информация", "Няма текуща обработка на IP адреси")

    def analyze_single_batch_ip(self, ip, force_source=None):
        """Анализира единичен IP адрес за пакетна проверка"""
        result = {"ip": ip, "country": "N/A", "isp": "N/A"}
        geo_result = self.security_tools.get_geolocation(ip, force_source)
        # Добавяне на всички услуги
        services=(
            "abuseipdb",
            "virustotal",
            "shodan",
            "geo",
            "geojs",
            "ipapi",
            "ipgeolocation",
            "whois",)
        analysis = self.security_tools.analyze_ip(ip, services=services)
        try:
            # Взимаме гео данни от избрания източник
            geo_result = self.security_tools.get_geolocation(ip, force_source)
            
            if "error" not in geo_result:
                result.update({
                    "country": geo_result.get("country", "N/A"),
                    "isp": geo_result.get("isp", "N/A")
                })
            
            # Извличане на ключова информация
            result.update({
                "ip": geo_result.get("ip", "N/A"),
                "country": geo_result.get("country", "N/A"),
                "isp": geo_result.get("isp", "N/A")
                #"source": geo_result.get("source", "N/A")
            })
        
            # Добавяме AbuseIPDB данни ако не сме форсирали друг източник
            if not force_source and self.config.get_api_key("ABUSEIP_API_KEY"):
                abuse_data = self.security_tools.check_abuseipdb(ip)
                if "error" not in abuse_data:
                    result.update({
                        "country": abuse_data.get("countryCode", result["country"]),
                        "isp": abuse_data.get("isp", result["isp"])
                    })
        
        except Exception as e:
            print(f"Грешка при анализ на {ip}: {str(e)}")
            
        return result

    def copy_batch_results(self):
        """Копира резултатите във формат IP - Държава - ISP"""
        text = ""
        for row in self.batch_ip_treeview.get_children():
            values = self.batch_ip_treeview.item(row)['values']
            ip = values[0]
            country = values[1] if len(values) > 1 else "N/A"
            isp = values[2] if len(values) > 2 else "N/A"
            
            # Форматиране на държавата (премахване на кода в скоби)
            country = country.split('(')[0].strip() if '(' in country else country
            text += f"{ip} - {country} - {isp}\n"
        
        self.root.clipboard_clear()
        self.root.clipboard_append(text.strip())
        messagebox.showinfo("Копиране", "Данните са копирани в клипборда!")

    def export_batch_results(self):
        """Експортира резултатите от пакетната проверка във файл"""
        if not self.batch_ip_results:
            messagebox.showwarning("Грешка", "Няма резултати за експортиране")
            print("[DEBUG] batch_ip_results:", self.batch_ip_results)
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV файлове", "*.csv"), ("Всички файлове", "*.*")],
        )

        if not file_path:
            return

        try:
            with open(file_path, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f)
                writer.writerow(["IP", "Държава", "Доставчик"])

                for result in self.batch_ip_results:
                    ip = result[0]
                    data = result[1] if isinstance(result[1], dict) else {}
                    writer.writerow([
                        ip,
                        data.get("country", "N/A"),
                        data.get("isp", "N/A")
                    ])

            messagebox.showinfo(
                "Успех", f"Резултатите са запазени във файл: {file_path}"
            )
        except Exception as e:
            messagebox.showerror("Грешка", f"Грешка при запис на файл: {str(e)}")
            
    def load_ips_from_file(self):
        file_path = filedialog.askopenfilename(
            title="Избери CSV или TXT файл с IP адреси",
            filetypes=[
                ("CSV файлове", "*.csv"),
                ("Текстови файлове", "*.txt"),
                ("Всички файлове", "*.*")
            ]
        )
        ips = []

        if file_path:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    for line in f:
                        line = line.strip()
                        if not line:
                            continue
                        parts = line.split(",") if "," in line else [line]
                        for part in parts:
                            ip = part.strip()
                            try:
                                ipaddress.ip_address(ip)
                                ips.append(ip)
                            except ValueError:
                                continue  # игнорира невалидни IP адреси
            except Exception as e:
                messagebox.showerror("Грешка при зареждане", str(e))
        return ips
    
    def import_ips_from_file(self):
        ips = self.load_ips_from_file()
        if ips:
            self.batch_ip_text.delete("1.0", tk.END)
            self.batch_ip_text.insert("1.0", "\n".join(ips))
            self.status_bar.config(text=f"Импортирани {len(ips)} IP адреса от файл")

    def create_url_analysis_tab(self):
        """Създава таб за анализ на URL адреси"""
        tab = ttk.Frame(self.notebook)
        self.url_analysis_tab = tab
        self.notebook.add(tab, text="URL Анализ")

        # Входни данни
        input_frame = ttk.Frame(tab)
        input_frame.pack(fill=tk.X, padx=10, pady=10)

        ttk.Label(input_frame, text="URL Адрес:").pack(side=tk.LEFT)
        self.url_entry = ttk.Entry(input_frame)
        self.url_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)

        # Бутони за анализ
        btn_frame = ttk.Frame(tab)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Button(
            btn_frame,
            text="Сканирай с VirusTotal",
            command=self.scan_url_virustotal,
        ).pack(side=tk.LEFT, padx=2)

        ttk.Button(
            btn_frame, text="Провери с URLScan.io", command=self.scan_url_urlscan
        ).pack(side=tk.LEFT, padx=2)

        ttk.Button(
            btn_frame, text="Извличане на DNS записи", command=self.url_dns_lookup
        ).pack(side=tk.LEFT, padx=2)

        # Резултати
        results_frame = ttk.Frame(tab)
        results_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.url_result_notebook = ttk.Notebook(results_frame)
        self.url_result_notebook.pack(fill=tk.BOTH, expand=True)

        # Текстови резултати
        text_tab = ttk.Frame(self.url_result_notebook)
        self.url_result_notebook.add(text_tab, text="Текст")
        self.url_text_results = scrolledtext.ScrolledText(
            text_tab, wrap=tk.WORD, font=("Consolas", 10)
        )
        self.url_text_results.pack(fill=tk.BOTH, expand=True)

        # HTML преглед
        html_tab = ttk.Frame(self.url_result_notebook)
        self.url_result_notebook.add(html_tab, text="HTML")
        self.url_html_view = scrolledtext.ScrolledText(
            html_tab, wrap=tk.WORD, font=("Consolas", 10)
        )
        self.url_html_view.pack(fill=tk.BOTH, expand=True)
        # Таблични резултати
        table_tab = ttk.Frame(self.url_result_notebook)
        self.url_result_notebook.add(table_tab, text="Таблица")

        self.url_treeview = ttk.Treeview(
            table_tab, columns=("Свойство", "Стойност"), show="headings"
        )
        self.url_treeview.heading("Свойство", text="Свойство")
        self.url_treeview.heading("Стойност", text="Стойност")
        self.url_treeview.column("Свойство", width=200)
        self.url_treeview.column("Стойност", width=400)

        scrollbar = ttk.Scrollbar(
            table_tab, orient="vertical", command=self.url_treeview.yview
        )
        self.url_treeview.configure(yscrollcommand=scrollbar.set)

        self.url_treeview.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Докладване
        report_frame = ttk.Frame(tab)
        report_frame.pack(fill=tk.X, pady=5)

        report_services=(
            ("Google Safe Browsing", "google"),
            ("PhishTank", "phishtank"),
            ("Microsoft", "microsoft"),
            ("APWG", "apwg"),)

        for i, (text, service) in enumerate(report_services):
            btn = ttk.Button(
                report_frame,
                text=text,
                command=lambda s=service: self.report_url(s),
            )
            btn.grid(row=0, column=i, padx=2, sticky=tk.EW)

    def scan_url_virustotal(self):
        """Сканира URL с VirusTotal"""
        url = self.url_entry.get().strip()
        if not url:
            messagebox.showwarning("Грешка", "Моля, въведете URL адрес")
            return
        self.log_activity(f"Стартирано VirusTotal сканиране на URL: {url}")

        self.status_bar.config(text=f"Сканиране на {url} с VirusTotal...")
        self.root.update()

        result = self.security_tools.scan_url_virustotal(url)

        self.url_text_results.delete(1.0, tk.END)

        if "error" in result:
            self.url_text_results.insert(tk.END, f"Грешка: {result['error']}")
            self.status_bar.config(text=f"Грешка при сканиране на {url}")
        else:
            self.url_text_results.insert(tk.END, f"Резултати от VirusTotal:\n\n")
            self.url_text_results.insert(tk.END, f"URL: {url}\n")
            self.url_text_results.insert(
                tk.END, f"ID на сканиране: {result['id']}\n"
            )
            self.url_text_results.insert(
                tk.END,
                f"Зловредни отзиви: {result['positives']}/{result['total']}\n",
            )
            self.url_text_results.insert(
                tk.END, f"Пермалинк: {result['permalink']}\n"
            )

            # Добавяне на линк за отваряне
            self.url_text_results.tag_config("link", foreground="blue", underline=1)
            self.url_text_results.tag_bind(
                "link", "<Button-1>", lambda e: webbrowser.open(result["permalink"])
            )
            self.url_text_results.insert(
                tk.END, "\nКликнете тук за отваряне на резултатите", "link"
            )

            self.status_bar.config(text=f"Готово - {url} сканиран с VirusTotal")

    def scan_url_urlscan(self):
        """Сканира URL с URLScan.io"""
        url = self.url_entry.get().strip()
        if not url:
            messagebox.showwarning("Грешка", "Моля, въведете URL адрес")
            return
        self.log_activity(f"Стартирано URLSCAN сканиране на URL: {url}")
        

        if not self.config.get_api_key("URLSCAN_API_KEY"):
            messagebox.showwarning("Грешка", "Липсва URLScan.io API ключ")
            return

        self.status_bar.config(text=f"Сканиране на {url} с URLScan.io...")
        self.root.update()

        headers = {
            "API-Key": self.config.get_api_key("URLSCAN_API_KEY"),
            "Content-Type": "application/json",
        }
        payload = {"url": url, "visibility": "public"}

        try:
            # Изпращане на заявка за сканиране
            response = requests.post(
                "https://urlscan.io/api/v1/scan/",
                headers=headers,
                json=payload,
                proxies=self.config.get_proxy_settings(),
                verify=False,
                timeout=15,
            )

            if response.status_code == 200:
                data = response.json()
                result_url = data.get("result")
                uuid = data.get("uuid")

                self.url_text_results.delete(1.0, tk.END)
                self.url_text_results.insert(tk.END, f"Сканирането е започнато:\n")
                self.url_text_results.insert(tk.END, f"URL: {url}\n")
                self.url_text_results.insert(tk.END, f"UUID: {uuid}\n")
                self.url_text_results.insert(tk.END, f"Резултати: {result_url}\n")

                # Добавяне на линк за отваряне
                self.url_text_results.tag_config(
                    "link", foreground="blue", underline=1
                )
                self.url_text_results.tag_bind(
                    "link", "<Button-1>", lambda e: webbrowser.open(result_url)
                )
                self.url_text_results.insert(
                    tk.END, "\nКликнете тук за отваряне на резултатите", "link"
                )

                # Започване на проверка за резултати
                self.root.after(10000, lambda: self.check_urlscan_result(uuid))

                self.status_bar.config(text=f"Сканирането на {url} е започнато")
            else:
                self.url_text_results.insert(
                    tk.END, f"Грешка: {response.status_code} - {response.text}"
                )
                self.status_bar.config(text=f"Грешка при сканиране на {url}")
        except Exception as e:
            self.url_text_results.insert(tk.END, f"Грешка: {str(e)}")
            self.status_bar.config(text=f"Грешка при сканиране на {url}")

    def check_urlscan_result(self, uuid, attempt=1):
        """Проверява резултатите от URLScan.io"""
        try:
            response = requests.get(
                f"https://urlscan.io/api/v1/result/{uuid}/",
                proxies=self.config.get_proxy_settings(),
                verify=False,
                timeout=15,
            )

            if response.status_code == 200:
                data = response.json()

                self.url_text_results.delete(1.0, tk.END)
                self.url_text_results.insert(tk.END, "URLScan.io резултати:\n\n")

                # Основна информация
                self.url_text_results.insert(
                    tk.END, f"URL: {data.get('page', {}).get('url', 'N/A')}\n"
                )
                self.url_text_results.insert(
                    tk.END, f"Домейн: {data.get('page', {}).get('domain', 'N/A')}\n"
                )
                self.url_text_results.insert(
                    tk.END, f"IP: {data.get('page', {}).get('ip', 'N/A')}\n"
                )
                self.url_text_results.insert(
                    tk.END,
                    f"Държава: {data.get('page', {}).get('country', 'N/A')}\n",
                )

                # Връзки и ресурси
                self.url_text_results.insert(tk.END, "\nВръзки и ресурси:\n")
                for request in data.get("data", {}).get("requests", [])[
                    :10
                ]:  # Показваме първите 10 заявки
                    self.url_text_results.insert(
                        tk.END,
                        f"- {request.get('request', {}).get('url', 'N/A')}\n",
                    )

                # HTML съдържание
                self.url_html_view.delete(1.0, tk.END)
                self.url_html_view.insert(
                    tk.END, data.get("data", {}).get("dom", "")
                )

                self.status_bar.config(text=f"Готово - резултатите са заредени")
            elif response.status_code == 404 and attempt <= 5:
                # Опитваме се отново след 10 секунди (максимум 5 опита)
                self.status_bar.config(
                    text=f"Очакване на резултати (опит {attempt}/5)..."
                )
                self.root.after(
                    10000, lambda: self.check_urlscan_result(uuid, attempt + 1)
                )
            else:
                self.url_text_results.insert(
                    tk.END, f"Грешка: {response.status_code} - {response.text}"
                )
                self.status_bar.config(text=f"Грешка при проверка на резултатите")
        except Exception as e:
            self.url_text_results.insert(tk.END, f"Грешка: {str(e)}")
            self.status_bar.config(text=f"Грешка при проверка на резултатите")

    def url_dns_lookup(self):
        """Извлича DNS записи за домейн от URL"""
        url = self.url_entry.get().strip()
        if not url:
            messagebox.showwarning("Грешка", "Моля, въведете URL адрес")
            return
        self.log_activity(f"Стартирано DNS извличане: {url}")

        try:
            # Извличане на домейн от URL
            domain = urllib.parse.urlparse(url).netloc
            if not domain:
                domain = url

            self.status_bar.config(text=f"Извличане на DNS записи за {domain}...")
            self.root.update()
            self.url_text_results.delete(1.0, tk.END)
            self.url_text_results.insert(tk.END, f"DNS записи за {domain}:\n\n")

            # Проверка на различни типове DNS записи
            record_types = ["A", "AAAA", "MX", "NS", "TXT", "CNAME"]

            for record_type in record_types:
                try:
                    answers = dns.resolver.resolve(domain, record_type)
                    self.url_text_results.insert(tk.END, f"{record_type} записи:\n")
                    for rdata in answers:
                        self.url_text_results.insert(tk.END, f"- {rdata}\n")
                    self.url_text_results.insert(tk.END, "\n")
                except (dns.resolver.NoAnswer, dns.resolver.NXDOMAIN):
                    continue
                except Exception as e:
                    self.url_text_results.insert(
                        tk.END, f"Грешка при проверка на {record_type}: {str(e)}\n"
                    )

            self.status_bar.config(text=f"Готово - DNS записи за {domain}")
        except Exception as e:
            self.url_text_results.insert(tk.END, f"Грешка: {str(e)}")
            self.status_bar.config(text=f"Грешка при извличане на DNS записи")

    def report_url(self, service):
        """Докладва URL към избраната услуга"""
        url = self.url_entry.get().strip()
        if not url:
            messagebox.showwarning("Грешка", "Моля, въведете URL адрес")
            return
        сelf.log_activity(f"Стартиране на докладване на Url: {url}")

        self.status_bar.config(text=f"Докладване на {url} към {service}...")
        self.root.update()

        result = self.security_tools.report_phishing(url, service)

        if "error" in result:
            messagebox.showerror(
                "Грешка", f"Докладването не бе успешно: {result['error']}"
            )
            self.status_bar.config(text=f"Грешка при докладване на {url}")
        else:
            messagebox.showinfo("Успех", "URL-то беше успешно докладвано")
            self.status_bar.config(text=f"Готово - {url} докладван към {service}")

    def create_file_analysis_tab(self):
        """Създава таб за анализ на файлове"""
        tab = ttk.Frame(self.notebook)
        self.file_analysis_tab = tab
        self.notebook.add(tab, text="Файлов Анализ")

        # Входни данни
        input_frame = ttk.Frame(tab)
        input_frame.pack(fill=tk.X, padx=10, pady=10)

        ttk.Button(input_frame, text="Избери Файл", command=self.select_file).pack(
            side=tk.LEFT, padx=5
        )

        self.file_path_label = ttk.Label(input_frame, text="Файл не е избран")
        self.file_path_label.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)

        # Drag and Drop функционалност
        drop_frame = ttk.LabelFrame(tab, text="Или пуснете файл тук", padding=10)
        drop_frame.pack(fill=tk.X, padx=10, pady=5)

        self.drop_label = ttk.Label(
            drop_frame, text="Пуснете файла тук за анализ", relief=tk.SUNKEN, padding=10
        )
        self.drop_label.pack(fill=tk.BOTH, expand=True)

        # Регистриране на drag and drop
        self.drop_label.drop_target_register(DND_FILES)
        self.drop_label.dnd_bind("<<Drop>>", self.handle_file_drop)

        # Бутони за анализ
        btn_frame = ttk.Frame(tab)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Button(
            btn_frame,
            text="Сканирай с VirusTotal",
            command=self.analyze_file_virustotal,
        ).pack(side=tk.LEFT, padx=2)

        ttk.Button(
            btn_frame,
            text="Анализирай с Hybrid Analysis",
            command=self.analyze_file_hybrid,
        ).pack(side=tk.LEFT, padx=2)

        ttk.Button(
            btn_frame, text="Изчисли хешове", command=self.calculate_hashes
        ).pack(side=tk.LEFT, padx=2)

        # Резултати
        results_frame = ttk.Frame(tab)
        results_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.file_result_notebook = ttk.Notebook(results_frame)
        self.file_result_notebook.pack(fill=tk.BOTH, expand=True)

        # Текстови резултати
        text_tab = ttk.Frame(self.file_result_notebook)
        self.file_result_notebook.add(text_tab, text="Текст")
        self.file_text_results = scrolledtext.ScrolledText(
            text_tab, wrap=tk.WORD, font=("Consolas", 10)
        )
        self.file_text_results.pack(fill=tk.BOTH, expand=True)

        # HEX преглед
        hex_tab = ttk.Frame(self.file_result_notebook)
        self.file_result_notebook.add(hex_tab, text="HEX")
        self.file_hex_view = scrolledtext.ScrolledText(
            hex_tab, wrap=tk.WORD, font=("Consolas", 10)
        )
        self.file_hex_view.pack(fill=tk.BOTH, expand=True)

        # Таблични резултати
        table_tab = ttk.Frame(self.file_result_notebook)
        self.file_result_notebook.add(table_tab, text="Таблица")

        self.file_treeview = ttk.Treeview(
            table_tab, columns=("Свойство", "Стойност"), show="headings"
        )
        self.file_treeview.heading("Свойство", text="Свойство")
        self.file_treeview.heading("Стойност", text="Стойност")
        self.file_treeview.column("Свойство", width=200)
        self.file_treeview.column("Стойност", width=400)

        scrollbar = ttk.Scrollbar(
            table_tab, orient="vertical", command=self.file_treeview.yview
        )
        self.file_treeview.configure(yscrollcommand=scrollbar.set)

        self.file_treeview.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    def select_file(self):
        """Избор на файл за анализ"""
        file_path = filedialog.askopenfilename()
        if file_path:
            self.current_file = file_path
            self.file_path_label.config(text=file_path)

    def handle_file_drop(self, event):
        """Обработка на пуснат файл"""
        file_path = event.data.strip("{}")
        if os.path.isfile(file_path):
            self.current_file = file_path
            self.file_path_label.config(text=file_path)
            self.file_text_results.insert(tk.END, f"Избран файл: {file_path}\n")

    def analyze_file_virustotal(self):
        """Анализира файл с VirusTotal"""
        if not self.current_file:
            messagebox.showwarning("Грешка", "Моля, изберете файл")
            return
        self.log_activity(f"Стартирано VirusTotal сканиране на файл: {self.current_file}")
        
        file_size = os.path.getsize(self.current_file)
        if file_size > 650 * 1024 * 1024:  # 650MB в байтове
            messagebox.showerror(
                "Грешка", 
                f"Файлът надвишава максималния размер от 650MB\n"
                f"Размер на файла: {file_size/(1024*1024):.2f}MB"
            )
            return

        self.status_bar.config(text=f"Сканиране на файл с VirusTotal...")
        self.root.update()

 
        result = self.security_tools.scan_file_virustotal(self.current_file)

        self.file_text_results.delete(1.0, tk.END)

        if "error" in result:
            self.file_text_results.insert(tk.END, f"Грешка: {result['error']}")
            self.status_bar.config(text=f"Грешка при сканиране на файл")
        else:
            self.file_text_results.insert(tk.END, f"Резултати от VirusTotal:\n\n")
            self.file_text_results.insert(
                tk.END, f"Име на файл: {os.path.basename(self.current_file)}\n"
            )
            self.file_text_results.insert(tk.END, f"Размер: {file_size/(1024*1024):.2f}MB\n")
            
            if result.get("type") == "instant":
                self.file_text_results.insert(tk.END, "Статус: Вече сканиран преди\n")
            else:
                self.file_text_results.insert(tk.END, "Статус: Нов анализ\n")
                
            self.file_text_results.insert(tk.END, f"SHA-256: {result.get('sha256', 'N/A')}\n")
            self.file_text_results.insert(
                tk.END,
                f"Зловредни отзиви: {result['positives']}/{result['total']}\n",
            )
            self.file_text_results.insert(
                tk.END, f"Пермалинк: {result['permalink']}\n"
            )

            if result.get("vendors"):
                self.file_text_results.insert(tk.END, "\nДетекции:\n")
                for vendor in result["vendors"]:
                    self.file_text_results.insert(tk.END, f"- {vendor}\n")

            # Добавяне на линк за отваряне
            self.file_text_results.tag_config("link", foreground="blue", underline=1)
            self.file_text_results.tag_bind(
                "link", "<Button-1>", lambda e: webbrowser.open(result["permalink"])
            )
            self.file_text_results.insert(
                tk.END, "\nКликнете тук за отваряне на резултатите", "link"
            )

            self.status_bar.config(text=f"Готово - файлът е сканиран с VirusTotal")

    def analyze_file_hybrid(self):
        """Анализира файл с Hybrid Analysis"""
        if not self.current_file:
            messagebox.showwarning("Грешка", "Моля, изберете файл")
            return

        self.status_bar.config(text=f"Анализиране на файл с Hybrid Analysis...")
        self.root.update()

        result = self.security_tools.scan_hybrid_analysis(
            file_path=self.current_file
        )

        self.file_text_results.delete(1.0, tk.END)

        if "error" in result:
            self.file_text_results.insert(tk.END, f"Грешка: {result['error']}")
            self.status_bar.config(text=f"Грешка при анализ на файл")
        else:
            data = result.get("data", {})
            self.file_text_results.insert(
                tk.END, f"Резултати от Hybrid Analysis:\n\n"
            )
            self.file_text_results.insert(
                tk.END, f"Име на файл: {os.path.basename(self.current_file)}\n"
            )
            self.file_text_results.insert(
                tk.END, f"Статус: {data.get('state', 'N/A')}\n"
            )
            self.file_text_results.insert(
                tk.END, f"Оценка на заплаха: {data.get('threat_score', 'N/A')}\n"
            )
            self.file_text_results.insert(
                tk.END, f"Вредни домейни: {len(data.get('domains', []))}\n"
            )
            self.file_text_results.insert(
                tk.END, f"Вредни IP адреси: {len(data.get('hosts', []))}\n"
            )

            # Показване на първите 10 вредни домейна и IP адреси
            if data.get("domains"):
                self.file_text_results.insert(tk.END, "\nВредни домейни:\n")
                for domain in data.get("domains", [])[:10]:
                    self.file_text_results.insert(tk.END, f"- {domain}\n")

            if data.get("hosts"):
                self.file_text_results.insert(tk.END, "\nВредни IP адреси:\n")
                for host in data.get("hosts", [])[:10]:
                    self.file_text_results.insert(tk.END, f"- {host}\n")

            self.status_bar.config(
                text=f"Готово - файлът е анализиран с Hybrid Analysis"
            )

    def calculate_hashes(self):
        """Изчислява хешове на файла"""
        if not self.current_file:
            messagebox.showwarning("Грешка", "Моля, изберете файл")
            return

        self.status_bar.config(text=f"Изчисляване на хешове...")
        self.root.update()

        try:
            with open(self.current_file, "rb") as f:
                file_data = f.read()

                md5 = hashlib.md5(file_data).hexdigest()
                sha1 = hashlib.sha1(file_data).hexdigest()
                sha256 = hashlib.sha256(file_data).hexdigest()

                self.file_text_results.delete(1.0, tk.END)
                self.file_text_results.insert(
                    tk.END, f"Хешове за {os.path.basename(self.current_file)}:\n\n"
                )
                self.file_text_results.insert(tk.END, f"MD5: {md5}\n")
                self.file_text_results.insert(tk.END, f"SHA-1: {sha1}\n")
                self.file_text_results.insert(tk.END, f"SHA-256: {sha256}\n")

                self.status_bar.config(text=f"Готово - хешовете са изчислени")
        except Exception as e:
            self.file_text_results.insert(tk.END, f"Грешка: {str(e)}")
            self.status_bar.config(text=f"Грешка при изчисляване на хешове")

    def create_phishing_tab(self):
        """Създава таб за фишинг доклади"""
        tab = ttk.Frame(self.notebook)
        self.notebook.add(tab, text="Фишинг Доклади")

        frame = ttk.Frame(tab)
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Входни данни
        input_frame = ttk.Frame(tab)
        input_frame.pack(fill=tk.X, padx=10, pady=10)

        ttk.Label(input_frame, text="Фишинг URL:").pack(side=tk.LEFT)
        self.phishing_url_entry = ttk.Entry(input_frame)
        self.phishing_url_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)

        ttk.Label(input_frame, text="IP Адрес:").pack(side=tk.LEFT)
        self.phishing_ip_entry = ttk.Entry(input_frame, width=20)
        self.phishing_ip_entry.pack(side=tk.LEFT, padx=5)
        
        # Бутони за анализ
        btn_frame = ttk.Frame(tab)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)

        analysis_services=(
            ("Провери за фишинг", self.handle_check_phishing),
            ("Докладвай", self.handle_report_phishing),
            ("Генерирай доклад", lambda: self.generate_phishing_template(self.phishing_template_var.get())),)

        for i, (text, cmd) in enumerate(analysis_services):
            btn = ttk.Button(btn_frame, text=text, command=cmd)
            btn.grid(row=0, column=i, padx=2, sticky=tk.EW)

        # Резултати
        results_frame = ttk.Frame(tab)
        results_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.phishing_result_notebook = ttk.Notebook(results_frame)
        self.phishing_result_notebook.pack(fill=tk.BOTH, expand=True)

        # Текстови резултати
        text_tab = ttk.Frame(self.phishing_result_notebook)
        self.phishing_result_notebook.add(text_tab, text="Текст")
        self.phishing_text_results = scrolledtext.ScrolledText(
            text_tab, wrap=tk.WORD, font=("Consolas", 10)
        )
        self.phishing_text_results.pack(fill=tk.BOTH, expand=True)
            
        # Шаблони за доклади
        templates_tab = ttk.Frame(self.phishing_result_notebook)
        self.phishing_result_notebook.add(templates_tab, text="Шаблони")

        self.phishing_template_var = tk.StringVar(value="standard")
        templates = [
            ("Report Phishin доклад", "standard"),
            ("До ГДБОП", "gdbop"),
            ("Нова фишинг кампания", "internal"),
        ]

        for text, value in templates:
            rb = ttk.Radiobutton(
                templates_tab,
                text=text,
                variable=self.phishing_template_var,
                value=value
            )
            rb.pack(anchor=tk.W)

        # Имейл функционалности
        email_frame = ttk.LabelFrame(tab, text="Имейл", padding=10)
        email_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Label(email_frame, text="Получател:").pack(side=tk.LEFT)
        self.email_recipient = ttk.Entry(email_frame)
        self.email_recipient.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        self.email_recipient.insert(
            0, self.config.get_email_settings()["default_recipient"]
        )

        ttk.Button(
            email_frame, text="Изпрати", command=self.send_phishing_report
        ).pack(side=tk.LEFT, padx=5)

        ttk.Button(
            email_frame, text="Отвори в Outlook", command=self.open_in_outlook
        ).pack(side=tk.LEFT, padx=5)

        # Докладване към външни услуги
        report_frame = ttk.LabelFrame(
            tab, text="Докладване към външни услуги", padding=10
        )
        report_frame.pack(fill=tk.X, padx=10, pady=5)

        services=(
            ("Google Safe Browsing", "google"),
            ("PhishTank", "phishtank"),
            ("Microsoft", "microsoft"),
            ("APWG", "apwg"),
            ("OpenPhish", "openphish"),)

        for i, (text, service) in enumerate(services):
            btn = ttk.Button(
                report_frame,
                text=text,
                command=lambda s=service: self.report_phishing_external(s),
            )
            btn.grid(row=i // 3, column=i % 3, padx=2, pady=2, sticky=tk.EW)

        # Преглед на шаблона
        preview_frame = ttk.Frame(tab)
        preview_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.phishing_template_preview = scrolledtext.ScrolledText(
            preview_frame, wrap=tk.WORD, font=("Consolas", 10)
        )
        self.phishing_template_preview.pack(fill=tk.BOTH, expand=True)

    def generate_phishing_template(self, template_type):
        """Генерира фишинг шаблон според избрания тип"""
        url = self.phishing_url_entry.get().strip()
        ip = self.phishing_ip_entry.get().strip()

        if not url:
            messagebox.showwarning("Грешка", "Моля, въведете URL адрес")
            return

        template = self.email_manager.generate_phishing_report(
            url, ip, template_type
        )
        self.phishing_template_preview.delete(1.0, tk.END)
        self.phishing_template_preview.insert(tk.END, template)

    def send_phishing_report(self):
        """Изпраща фишинг доклада по имейл"""
        recipient = self.email_recipient.get().strip()
        if not recipient:
            messagebox.showwarning("Грешка", "Моля, въведете получател")
            return

        email_text = self.phishing_template_preview.get(1.0, tk.END).strip()
        if not email_text:
            messagebox.showwarning("Грешка", "Няма генериран шаблон за изпращане")
            return

        # Извличане на тема от имейла (първи ред)
        subject = email_text.split("\n")[0].replace("Subject: ", "")
        body = "\n".join(email_text.split("\n")[1:])

        result = self.email_manager.send_email(recipient, subject, body)

        if "error" in result:
            messagebox.showerror(
                "Грешка", f"Грешка при изпращане: {result['error']}"
            )
        else:
            messagebox.showinfo("Успех", "Имейлът е изпратен успешно")

    def open_in_outlook(self):
        """Отваря шаблона в Outlook"""
        email_text = self.phishing_template_preview.get(1.0, tk.END).strip()
        if not email_text:
            messagebox.showwarning("Грешка", "Няма генериран шаблон")
            return

        # Запазване на временен файл
        temp_path = os.path.join(os.getenv("TEMP"), "phishing_report.eml")
        try:
            with open(temp_path, "w", encoding="utf-8") as f:
                f.write(email_text)

            # Отваряне на файла с Outlook
            os.startfile(temp_path)
        except Exception as e:
            messagebox.showerror(
                "Грешка", f"Грешка при отваряне в Outlook: {str(e)}"
            )

    def report_phishing_external(self, service):
        """Докладва фишинг URL към външна услуга"""
        url = self.phishing_url_entry.get().strip()
        if not url:
            messagebox.showwarning("Грешка", "Моля, въведете URL адрес")
            return

        result = self.security_tools.report_phishing(url, service.lower(), log=self.log_activity)
        
        
        if result.get("success"):
            # === CIRCL визуализация ===
            data = result.get("data")
            if service.lower() == "openphish" and isinstance(result.get("data"), dict):
                circl_data = result["data"]
                if log := getattr(self, "log_activity", None):
                    log(f"Резултати от CIRCL URLAbuse за {url}:")
                    for section_title, section_text in circl_data.items():
                        log(f"[CIRCL] --- {section_title} ---")
                        for line in section_text.splitlines():
                            if line.strip():
                                log(f"[CIRCL] {line.strip()}")
            else:
                messagebox.showinfo("Успех", f"URL-то беше докладвано успешно към {service}")
        else:
            error_msg = result.get("error", "Неизвестна грешка")
            messagebox.showerror("Грешка", f"{service} грешка: {error_msg}")



    def create_network_tools_tab(self):
        """Създава таб с мрежови инструменти"""
        tab = ttk.Frame(self.notebook)
        self.network_tools_tab = tab
        self.notebook.add(tab, text="Мрежови Инструменти")

        # Създаваме Notebook за разделяне на инструментите
        tools_notebook = ttk.Notebook(tab)
        tools_notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

    # ========== Създаване на раздел за Ping/Traceroute ==========
        basic_frame = ttk.Frame(tools_notebook)
        tools_notebook.add(basic_frame, text="Основни инструменти")
        
        verbose_frame = ttk.Frame(tools_notebook)
        tools_notebook.add(verbose_frame, text="Verbose Порт скенер")
        self.add_port_scanner_ui(verbose_frame)

        # Ping инструмент
        ping_frame = ttk.LabelFrame(tab, text="Ping", padding=10)
        ping_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Label(ping_frame, text="Хост:").pack(side=tk.LEFT)
        self.ping_host = ttk.Entry(ping_frame)
        self.ping_host.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        # Entry за няколко IP адреса (ping_entry)
        ping_label = ttk.Label(ping_frame, text="IP адреси (разделени с интервал или запетая):")
        ping_label.pack(side=tk.LEFT, padx=(10, 0))
        self.ping_entry = ttk.Entry(ping_frame, width=50)
        self.ping_entry.pack(side=tk.LEFT, padx=5)

        ttk.Button(ping_frame, text="Ping", command=self.run_ping).pack(
            side=tk.LEFT, padx=5
        )

        self.stop_ping_btn = ttk.Button(
            ping_frame, text="Спри", command=self.stop_ping, state=tk.DISABLED
        )
        self.stop_ping_btn.pack(side=tk.LEFT, padx=5)

        # Traceroute инструмент
        trace_frame = ttk.LabelFrame(tab, text="Traceroute", padding=10)
        trace_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Label(trace_frame, text="Хост:").pack(side=tk.LEFT)
        self.trace_host = ttk.Entry(trace_frame)
        self.trace_host.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)

        ttk.Button(
            trace_frame, text="Traceroute", command=self.run_traceroute
        ).pack(side=tk.LEFT, padx=5)

    # ========== Създаване на раздел за Port Scanner ==========
        portscan_frame = ttk.Frame(tools_notebook)
        tools_notebook.add(portscan_frame, text="Порт Скенер")

        fast_btn_frame = ttk.Frame(portscan_frame)  
        fast_btn_frame.pack(fill=tk.X, padx=5, pady=2)

        ttk.Button(fast_btn_frame, text="Бързо сканиране (5 сек)", command=self.fast_scan).pack(side=tk.LEFT, padx=2)
        ttk.Button(fast_btn_frame, text="Масово сканиране", command=self.mass_scan_ips).pack(side=tk.LEFT, padx=2)
        
        
        # Порт скенер контроли
        scan_control_frame = ttk.Frame(portscan_frame)
        scan_control_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(scan_control_frame, text="Цели:").grid(row=0, column=0, sticky="e")
        self.ps_target_entry = ttk.Entry(scan_control_frame, width=30)
        self.ps_target_entry.grid(row=0, column=1, padx=5)
        self.ps_target_entry.insert(0, "192.168.1.1, 192.168.1.0/24")

        ttk.Label(scan_control_frame, text="Портове:").grid(row=1, column=0, sticky="e")
        self.ps_ports_entry = ttk.Entry(scan_control_frame)
        self.ps_ports_entry.grid(row=1, column=1, padx=5, sticky="ew")
        self.ps_ports_entry.insert(0, "1-1024,3306,3389")

        # Тип сканиране
        self.ps_type_var = tk.StringVar(value="tcp")
        ttk.Radiobutton(scan_control_frame, text="TCP", variable=self.ps_type_var, value="tcp").grid(row=2, column=0)
        ttk.Radiobutton(scan_control_frame, text="UDP", variable=self.ps_type_var, value="udp").grid(row=2, column=1)

        # Бутони
        btn_frame = ttk.Frame(scan_control_frame)
        btn_frame.grid(row=3, column=0, columnspan=2, pady=5)
        
        self.ps_start_btn = ttk.Button(btn_frame, text="Стартирай", command=self.start_port_scan)
        self.ps_start_btn.pack(side=tk.LEFT, padx=2)
        
        self.ps_stop_btn = ttk.Button(btn_frame, text="Спри", command=self.stop_port_scan, state=tk.DISABLED)
        self.ps_stop_btn.pack(side=tk.LEFT, padx=2)
        
        self.ps_export_btn = ttk.Button(btn_frame, text="Експорт CSV", command=self.export_scan_results)
        self.ps_export_btn.pack(side=tk.LEFT, padx=2)

        # Прогрес
        self.ps_progress = ttk.Progressbar(scan_control_frame, orient=tk.HORIZONTAL, mode='determinate')
        self.ps_progress.grid(row=4, column=0, columnspan=2, sticky="ew", pady=5)



        # Резултати
        results_frame = ttk.Frame(tab)
        results_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.network_results = scrolledtext.ScrolledText(
            results_frame, wrap=tk.WORD, font=("Consolas", 10)
        )
        self.network_results.pack(fill=tk.BOTH, expand=True)

        
        
        # Резултати
        result_frame = ttk.Frame(portscan_frame)
        result_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        columns = ("IP", "Порт", "Тип", "Статус", "Услуга", "Банер")
        self.ps_result_tree = ttk.Treeview(result_frame, columns=columns, show="headings", selectmode="browse")
        
        for col in columns:
            self.ps_result_tree.heading(col, text=col)
            self.ps_result_tree.column(col, width=100, stretch=True)
        
        scrollbar = ttk.Scrollbar(result_frame, orient="vertical", command=self.ps_result_tree.yview)
        self.ps_result_tree.configure(yscrollcommand=scrollbar.set)
        
        self.ps_result_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            # Semaphore за контрол на броя едновременни сканирания
        self.scan_semaphore = threading.Semaphore(100)  # Макс 100 едновременни сканирания

        # Lock за синхронизация на изхода
        self.output_lock = threading.Lock()
        
#r


    async def fast_scan(self):
        """Бързо сканиране на ключови портове за 5-10 секунди"""
        target = self.ps_target_entry.get().strip()
        if not target:
            messagebox.showwarning("Грешка", "Моля, въведете целеви IP адрес")
            return
        
        COMMON_PORTS = [21, 22, 23, 25, 53, 80, 110, 139, 143, 443, 445, 
                        3389, 3306, 8080, 8443, 5900, 6379, 8000, 27017]
        
        self.log_activity(f"Стартирано бързо сканиране на {target}")
        self.ps_result_tree.delete(*self.ps_result_tree.get_children())
        
        try:
            start_time = time.time()
            open_ports = await self.async_scan_ports(target, COMMON_PORTS, timeout=1)
            elapsed = time.time() - start_time
            
            self.log_activity(f"Бързо сканиране завършено за {elapsed:.2f} сек. Открити {len(open_ports)} порта")
            
            for port, banner in open_ports:
                try:
                    service = socket.getservbyport(port)
                except OSError:
                    service = "unknown"
                except Exception as e:
                    service = f"err: {e}"

                self._add_scan_result(target, port, "TCP", banner)

        except Exception as e:
            self.log_activity(f"Грешка при бързо сканиране: {str(e)}")
        finally:
            loop.close()

    async def async_scan_ports(self, ip, ports, timeout=0.5):
        """Асинхронно сканиране на портове"""
        open_ports = []
    
        async def check_port(port):
            try:
                reader, writer = await asyncio.wait_for(
                    asyncio.open_connection(ip, port),
                    timeout=timeout
                )
                
                # Опит за получаване на банер
                banner = ""
                try:
                    writer.write(b"GET / HTTP/1.1\r\n\r\n")
                    await writer.drain()
                    banner = await asyncio.wait_for(reader.read(1024), timeout=0.3)
                    banner = banner.decode(errors='ignore').strip()
                except:
                    pass
                    
                writer.close()
                await writer.wait_closed()
                
                return port, banner
            except:
                return None
        
        # Създаване на задачи за всички портове
        tasks = [check_port(port) for port in ports]
        
        # Изпълнение на задачите паралелно
        results = await asyncio.gather(*tasks)
        
        # Филтриране на отворените портове
        for result in results:
            if result:
                open_ports.append(result)
        
        return open_ports
        
        
      
    
    def mass_scan_ips(self):
        """Многопоточено сканиране на множество IP адреси"""
        targets = self.ps_target_entry.get().strip()
        ports_str = self.ps_ports_entry.get().strip()
        
        if not targets or not ports_str:
            messagebox.showwarning("Грешка", "Моля, въведете цели и портове")
            return
        
        try:
            # Парсване на целите
            target_ips = self.parse_targets(targets)
            ports = self.parse_ports(ports_str)
            
            # Ограничаване на броя нишки за големи сканирания
            max_threads = min(500, len(target_ips) * len(ports))
            
            self.log_activity(f"Стартирано масово сканиране на {len(target_ips)} IP адреса")
            self.ps_result_tree.delete(*self.ps_result_tree.get_children())
            
            # Създаване на ThreadPoolExecutor
            with ThreadPoolExecutor(max_workers=max_threads) as executor:
                # Създаване на futures за всяка комбинация IP:порт
                futures = []
                for ip in target_ips:
                    for port in ports:
                        futures.append(
                            executor.submit(
                                self.check_port_with_timeout,
                                ip, port, self.ps_type_var.get(), 1
                            )
                        )
                
                # Проследяване на прогреса
                completed = 0
                total = len(futures)
                self.ps_progress['maximum'] = total
                
                for future in concurrent.futures.as_completed(futures):
                    completed += 1
                    self.ps_progress['value'] = completed
                    
                    ip, port, is_open, banner = future.result()
                    if is_open:
                        service = socket.getservbyport(port) if port <= 65535 else "unknown"
                        self.add_scan_result(ip, port, self.ps_type_var.get().upper(), 
                                        "OPEN", service, banner)
                    
                    # Обновяване на GUI всяки 50 проверки
                    if completed % 50 == 0:
                        self.root.update()
            
            self.log_activity("Масовото сканиране завършено успешно")
        
        except Exception as e:
            self.log_activity(f"Грешка при масово сканиране: {str(e)}")

    def check_port_with_timeout(self, ip, port, scan_type, timeout):
        """Проверка на порт с таймаут за използване с ThreadPool"""
        try:
            is_open, banner = self.check_port(ip, port, scan_type, timeout)
            return ip, port, is_open, banner
        except:
            return ip, port, False, ""
        

    def start_port_scan(self):
        """Стартира асинхронно сканиране на портове"""
        if hasattr(self, 'ps_scan_thread') and self.ps_scan_thread.is_alive():
            return
        
        targets = self.ps_target_entry.get().strip()
        ports_str = self.ps_ports_entry.get().strip()
        scan_type = self.ps_type_var.get()
        
        if not targets or not ports_str:
            messagebox.showwarning("Грешка", "Моля, въведете цели и портове")
            return
        
        # Парсване на цели
        try:
            target_ips = self.parse_targets(targets)
        except ValueError as e:
            messagebox.showerror("Грешка", f"Невалидни цели: {str(e)}")
            return
        
        # Парсване на портове
        try:
            ports = self.parse_ports(ports_str)
        except ValueError as e:
            messagebox.showerror("Грешка", f"Невалидни портове: {str(e)}")
            return
        
        # Подготовка за сканиране
        self.ps_result_tree.delete(*self.ps_result_tree.get_children())
        self.ps_progress['maximum'] = len(target_ips) * len(ports)
        self.ps_progress['value'] = 0
        
        self.ps_scan_running = True
        self.ps_scan_stop = False
        self.ps_start_btn.config(state=tk.DISABLED)
        self.ps_stop_btn.config(state=tk.NORMAL)
        
        # Логване на активността
        self.log_activity(f"Стартирано {scan_type.upper()} сканиране на {len(target_ips)} цели и {len(ports)} портове")
        
        # Стартиране на сканирането в отделна нишка
        self.ps_scan_thread = threading.Thread(
            target=self.run_port_scan,
            args=(target_ips, ports, scan_type),
            daemon=True
        )
        self.ps_scan_thread.start()

    def stop_port_scan(self):
        """Спира текущото сканиране на портове"""
        if hasattr(self, 'ps_scan_running') and self.ps_scan_running:
            self.ps_scan_stop = True
            self.ps_stop_btn.config(state=tk.DISABLED)
            self.log_activity("Спиране на сканирането на портове...")

    def run_port_scan(self, targets, ports, scan_type):
        """Изпълнява сканирането на портове с синхронизация"""
        try:
            with ThreadPoolExecutor(max_workers=100) as executor:
                futures = []
                for ip in targets:
                    if self.ps_scan_stop:
                        break
                        
                    for port in ports:
                        if self.ps_scan_stop:
                            break
                        
                        # Ограничаваме броя едновременни сканирания
                        with self.scan_semaphore:
                            future = executor.submit(
                                self._safe_check_port,
                                ip, port, scan_type
                            )
                            futures.append(future)
                        
                        # Проверка на прогреса
                        self.root.after(0, self._update_scan_progress)
                
                # Обработка на резултатите
                for future in concurrent.futures.as_completed(futures):
                    ip, port, is_open, banner = future.result()
                    if is_open:
                        self.root.after(0, self._add_scan_result, ip, port, scan_type, banner)
        
        except Exception as e:
            self.root.after(0, self.log_activity, f"Грешка при сканиране: {str(e)}")
        
        finally:
            self.root.after(0, self._finish_scan)
        
            self.ps_scan_running = False
            self.ps_start_btn.config(state=tk.NORMAL)
            self.ps_stop_btn.config(state=tk.DISABLED)

    def _safe_check_port(self, ip, port, scan_type):
        """Проверка на порт със синхронизация"""
        try:
            is_open, banner = self.check_port(ip, port, scan_type)
            return ip, port, is_open, banner
        except Exception as e:
            return ip, port, False, str(e)

    def check_port(self, ip, port, scan_type, timeout=0.5):
        """Оптимизирана проверка на порт с по-нисък таймаут"""
        try:
            if scan_type == "tcp":
                # Използваме socket.create_connection за по-бързо TCP сканиране
                s = socket.create_connection((ip, port), timeout=timeout)
                s.settimeout(0.3)  # По-кратък таймаут за банер
                
                banner = ""
                try:
                    s.send(b"GET / HTTP/1.1\r\nHost: %b\r\n\r\n" % ip.encode())
                    banner = s.recv(1024).decode(errors='ignore').strip()
                except:
                    pass
                    
                s.close()
                return True, banner
            else:
                # UDP сканиране с по-кратък таймаут
                with socket.socket(socket.AF_INET, socket.SOCK_DGRAM) as s:
                    s.settimeout(timeout)
                    s.sendto(b'\x00', (ip, port))
                    data, _ = s.recvfrom(1024)
                    return True, data.decode(errors='ignore').strip()
        
        except socket.timeout:
            return False, ""
        except ConnectionRefusedError:
            return False, ""
        except:
            return False, ""
            
    def _update_scan_progress(self):
        """Обновяване на прогреса от главния thread"""
        self.ps_progress['value'] += 1
        self.status_bar.config(text=f"Сканирани {self.ps_progress['value']}/{self.ps_progress['maximum']}")

    def _add_scan_result(self, ip, port, scan_type, banner):
        """Добавяне на резултат от главния thread"""
        with self.output_lock:
            try:
                service = socket.getservbyport(port)
            except OSError:
                service = "unknown"
            except Exception as e:
                service = f"err: {e}"
            
            self.ps_result_tree.insert("", tk.END, 
                                    values=(ip, port, scan_type.upper(), 
                                            "OPEN", service, banner))
            
    def _finish_scan(self):
        """Финализиране на сканирането от главния thread"""
        self.ps_scan_running = False
        self.ps_start_btn.config(state=tk.NORMAL)
        self.ps_stop_btn.config(state=tk.DISABLED)
        self.log_activity(f"Сканирането приключи {'(прекъснато)' if self.ps_scan_stop else ''}")
            
    async def async_network_scan(self, cidr, ports):
        """Паралелно сканиране на цяла мрежа"""
        network = ipaddress.ip_network(cidr)
        tasks = []
        
        for ip in network.hosts():
            ip_str = str(ip)
            tasks.append(self.async_scan_ports(ip_str, ports))
        
        return await asyncio.gather(*tasks)
        

    def add_scan_result(self, ip, port, port_type, status, service, banner):
        """Добавя резултат в дървото с резултати"""
        self.ps_result_tree.insert("", tk.END, values=(ip, port, port_type, status, service, banner))
        self.root.update_idletasks()

    def parse_targets(self, targets_str):
        """Парсва списък с цели (IP адреси или CIDR)"""
        targets = []
        for target in targets_str.split(','):
            target = target.strip()
            if '/' in target:  # CIDR нотация
                network = ipaddress.ip_network(target, strict=False)
                targets.extend([str(host) for host in network.hosts()])
            else:  # Единичен IP
                ipaddress.ip_address(target)  # Валидация
                targets.append(target)
        return targets

    def parse_ports(self, ports_str):
        """Парсва списък с портове (диапазони или единични)"""
        ports = set()
        for part in ports_str.split(','):
            part = part.strip()
            if '-' in part:
                start, end = map(int, part.split('-'))
                ports.update(range(start, end+1))
            else:
                ports.add(int(part))
        return sorted(ports)

    def export_scan_results(self):
        """Експортира резултатите от сканирането във CSV файл"""
        if not self.ps_result_tree.get_children():
            messagebox.showwarning("Грешка", "Няма резултати за експортиране")
            return
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV файлове", "*.csv"), ("Всички файлове", "*.*")]
        )
        
        if file_path:
            try:
                with open(file_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerow(["IP", "Порт", "Тип", "Статус", "Услуга", "Банер"])
                    
                    for item in self.ps_result_tree.get_children():
                        writer.writerow(self.ps_result_tree.item(item)['values'])
                
                self.log_activity(f"Резултатите са експортирани във {file_path}")
                messagebox.showinfo("Успех", "Резултатите са експортирани успешно")
            
            except Exception as e:
                messagebox.showerror("Грешка", f"Грешка при експортиране: {str(e)}")

#r

        
        
    def add_port_scanner_ui(self, parent):
        port_frame = ttk.LabelFrame(parent, text="Скенер на портове", padding=10)
        port_frame.pack(fill=tk.X, padx=10, pady=10)

        ttk.Label(port_frame, text="Хост:").grid(row=0, column=0, sticky="w")
        self.portscan_ip_entry = ttk.Entry(port_frame, width=20)
        self.portscan_ip_entry.insert(0, "192.168.1.1")
        self.portscan_ip_entry.grid(row=0, column=1, padx=5, pady=2)

        ttk.Label(port_frame, text="Начален порт:").grid(row=1, column=0, sticky="w")
        self.portscan_start = ttk.Entry(port_frame, width=10)
        self.portscan_start.insert(0, "1")
        self.portscan_start.grid(row=1, column=1, padx=5, pady=2, sticky="w")

        ttk.Label(port_frame, text="Краен порт:").grid(row=2, column=0, sticky="w")
        self.portscan_end = ttk.Entry(port_frame, width=10)
        self.portscan_end.insert(0, "1024")
        self.portscan_end.grid(row=2, column=1, padx=5, pady=2, sticky="w")

        # Чекбокс за подробен режим (verbose)
        self.verbose_var = tk.BooleanVar(value=False)
        verbose_check = ttk.Checkbutton(
            port_frame,
            text="Verbose Scan",
            variable=self.verbose_var,
            command=lambda: setattr(self, 'enable_verbose_scan', self.verbose_var.get())
        )
        verbose_check.grid(row=3, column=0, columnspan=2, sticky="w")

        ttk.Button(port_frame, text="Сканирай портове", command=lambda: self.scan_ports_threaded(
            self.portscan_ip_entry.get().strip(),
            int(self.portscan_start.get()),
            int(self.portscan_end.get())
        )).grid(row=4, column=0, columnspan=2, pady=5)
        
        # DNS Lookup инструмент
        #dns_frame = ttk.LabelFrame(tab, text="DNS Lookup", padding=10)
        #dns_frame.pack(fill=tk.X, padx=10, pady=5)

        #ttk.Label(dns_frame, text="Домейн:").pack(side=tk.LEFT)
        #self.dns_domain_entry = ttk.Entry(dns_frame)
        #self.dns_domain_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)

        #ttk.Label(dns_frame, text="Тип:").pack(side=tk.LEFT, padx=5)
        #self.dns_type_combobox = ttk.Combobox(
         #   dns_frame,
          #  values=["A", "AAAA", "MX", "NS", "SOA", "TXT"],
           # width=5
        #)
        #self.dns_type_combobox.current(0)
        #self.dns_type_combobox.pack(side=tk.LEFT)

        #ttk.Button(
         #   dns_frame,
          #  text="Търси",
           # command=self.run_dns_lookup
       # ).pack(side=tk.LEFT)

        #self.dns_result = scrolledtext.ScrolledText(
         #   dns_frame, height=5, wrap=tk.WORD, font=("Consolas", 9)
        #)
        #self.dns_result.pack(fill=tk.X, pady=5)
        
# Добавяне на бърз TCP порт скенер с енумерация
    COMMON_PORTS = {
        21: "FTP",
        22: "SSH",
        23: "Telnet",
        25: "SMTP",
        53: "DNS",
        80: "HTTP",
        110: "POP3",
        139: "NetBIOS",
        161: "SNMP",
        162: "SNMP",
        143: "IMAP",
        443: "HTTPS",
        445: "SMB",
        993: "IMAP SSL",
        3306: "MySQL",
        3389: "RDP"
    }

    def scan_single_port(self, host, port, results):
        COMMON_PORTS = {
            21: "FTP",
            22: "SSH",
            23: "Telnet",
            25: "SMTP",
            53: "DNS",
            80: "HTTP",
            110: "POP3",
            135: "RPC",
            139: "NetBIOS",
            161: "SNMP",
            162: "SNMP",
            143: "IMAP",
            443: "HTTPS",
            445: "SMB",
            993: "IMAP SSL",
            2179: "VMRDP",
            2869: "ICSLAP",
            3240: "USBIPD/TRIOMOTIO/CITRIX",
            3306: "MySQL",
            3389: "RDP"
        }
        
        try:
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.settimeout(1)
                result = s.connect_ex((host, port))

                if result == 0:
                    banner = ""
                    try:
                        s.sendall(b"\r\n")
                        banner = s.recv(1024).decode(errors='ignore').strip()
                    except:
                        pass
                    service = COMMON_PORTS.get(port, "Неизвестен")
                    results.append((port, "ОТВОРЕН", service, banner))
                elif result == 10060:
                    results.append((port, "ФИЛТРИРАН", "", ""))
                elif result == 10061:
                    if getattr(self, 'enable_verbose_scan', False):
                        results.append((port, "ЗАТВОРЕН", "", ""))
        except Exception as e:
            if getattr(self, 'enable_verbose_scan', False):
                results.append((port, f"ГРЕШКА: {e}", "", ""))

    def scan_ports_threaded(self, host, start_port, end_port):
        self.network_results.delete(1.0, tk.END)
        self.network_results.insert(tk.END, f"Сканиране на портове от {start_port} до {end_port} на {host}...\n")
        self.log_activity(f"Стартирано сканиране на портове: {host} ({start_port}-{end_port})")
        results = []
        threads = []
        for port in range(start_port, end_port + 1):
            t = threading.Thread(target=self.scan_single_port, args=(host, port, results))
            threads.append(t)
            t.start()

        for t in threads:
            t.join()

        if results:
            for port, status, service, banner in sorted(results):
                self.network_results.insert(tk.END, f"{status}: порт {port}")
                if service:
                    self.network_results.insert(tk.END, f" ({service})")
                self.network_results.insert(tk.END, "\n")
                if banner:
                    self.network_results.insert(tk.END, f"  Банер: {banner}\n")
        else:
            self.network_results.insert(tk.END, "Няма открити отворени портове.\n")

        self.status_var.set(f"Сканирани {completed}/{total} комбинации | Отворени: {open_count}")
        self.status_bar.config(text=f"Сканиране приключи за {host}")
        self.log_activity(f"Сканиране на портове приключи: {host}")
    



    def run_ping(self):
        try:
            ips_input = self.ping_entry.get().strip()
        except Exception:
            messagebox.showerror("Грешка", "Липсва поле за IP адреси.")
            return

        if ',' in ips_input:
            ips = [ip.strip() for ip in ips_input.split(',')]
        else:
            ips = ips_input.split()

        output = self.ping_output if hasattr(self, 'ping_output') else self.ps_output
        output.delete(1.0, tk.END)
        output.insert(tk.END, "Стартирам ping към:\n" + ", ".join(ips) + "\n\n")

        self.ping_stop = False
        self.stop_ping_btn.config(state=tk.NORMAL)

        for ip in ips:
            if not self.ping_stop:
                threading.Thread(target=self._async_ping, args=(ip,), daemon=True).start()





    def _async_ping(self, ip):
        import platform
        CREATE_NO_WINDOW = 0x08000000 if platform.system() == 'Windows' else 0, time
        
        try:
            system_platform = platform.system()
            base_cmd = ["ping", "-n", "1"] if system_platform == "Windows" else ["ping", "-c", "1"]
            if not hasattr(self, 'stop_ping_event'):
                self.stop_ping_event = threading.Event()
            while not self.stop_ping_event.is_set():
                cmd = base_cmd + [ip]
                print(f"📡 Ping към {ip} с: {' '.join(cmd)}")
                creationflags = subprocess.CREATE_NO_WINDOW if platform.system() == "Windows" else 0
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=5, creationflags=creationflags)
                output_text = result.stdout.strip() or result.stderr.strip()
                timestamp = time.strftime('%H:%M:%S')
                if not output_text:
                    output_text = f"[{timestamp}] [{ip}] ⚠️ Няма отговор"
                else:
                    output_text = f"[{timestamp}] [{ip}]\n{output_text}"

                self.root.after(0, lambda output_text=output_text: self._safe_insert_ping(output_text))
                time.sleep(1)
        except Exception as e:
            err = f"[{ip}] ❌ Ping грешка: {e}\n"
            print(err)
            self.root.after(0, lambda: self.ps_output.insert(tk.END, err))

    def _safe_insert_ping(self, output_text):
        try:
            output = getattr(self, 'ps_output', None) or getattr(self, 'ping_output', None) or getattr(self, 'network_results', None)
            if output:
                output.insert(tk.END, output_text + "\n\n")
                output.see(tk.END)
            else:
                print("❌ Няма валиден output widget")
        except Exception as e:
            print(f"Insert грешка: {e}")
            def safe_insert():
                try:
                    output.insert(tk.END, output_text + "\n\n")
                    output.see(tk.END)
                except Exception as e:
                    print(f"❌ GUI Insert грешка: {e}")

            self.root.after(0, safe_insert)

        except Exception as e:
            err = f"[{ip}] ❌ Грешка при ping: {e}\n"
            print(err)
            def safe_error():
                try:
                    output.insert(tk.END, err)
                except Exception as gui_err:
                    print(f"❌ GUI Insert грешка при грешка: {gui_err}")
            self.root.after(0, safe_error)
    def _update_ping_output(self, text):
        """Безопасно обновяване на ping изхода от главния thread"""
        self.ps_output.insert(tk.END, text)
        self.ps_output.see(tk.END)


    def stop_ping(self):
        if hasattr(self, 'stop_ping_event'):
            self.stop_ping_event.set()
            print("⛔ Stop Ping активиран (event).")
    def read_ping_output(self, process, host):
        """Чете изхода от ping процеса в реално време"""
        while True:
            output = process.stdout.readline()
            if output == "" and process.poll() is not None:
                break
            if output:
                self.network_results.insert(tk.END, f"{host}: {output}")
                self.network_results.see(tk.END)
                self.root.update()

            
    def perform_traceroute(self, host):
        self.update_status(f"Стартиране на traceroute към {host}...")
        self.log_activity(f"Traceroute стартиран за {host}")
        self.status_bar.config(text=f"Traceroute към {host}...")
        self.root.update()

        def run():
            try:
                if os.name == 'nt':
                    cmd = ["tracert", "-d", host]
                else:
                    cmd = ["traceroute", host]

                proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                output, error = proc.communicate(timeout=60)

                self.network_results.delete(1.0, tk.END)
                if proc.returncode == 0:
                    self.network_results.insert(tk.END, output)
                    self.log_activity(f"Traceroute успешно приключи за {host}")
                else:
                    self.network_results.insert(tk.END, f"Грешка при traceroute: {error}")
                    self.log_activity(f"Traceroute неуспешен за {host}: {error}")

                self.status_bar.config(text=f"Traceroute към {host} завършен")
            except subprocess.TimeoutExpired:
                self.network_results.insert(tk.END, "Traceroute таймаут.\n")
                self.status_bar.config(text=f"Traceroute таймаут към {host}")
                self.log_activity(f"Traceroute таймаут за {host}")
            except Exception as e:
                self.network_results.insert(tk.END, f"Грешка: {str(e)}\n")
                self.status_bar.config(text=f"Грешка при traceroute към {host}")
                self.log_activity(f"Traceroute изключение за {host}: {e}")

        threading.Thread(target=run).start()
            
    def run_traceroute(self):
        """Обвивка за изпълнение на traceroute с извличане на хост от полето"""
        host = self.trace_host.get().strip()
        if not host:
            messagebox.showwarning("Грешка", "Моля, въведете хост")
            return
        self.perform_traceroute(host)



    def create_threat_intel_tab(self):
        """Създава таб с Threat Intelligence инструменти"""
        tab = ttk.Frame(self.notebook)
        self.threat_intel_tab = tab
        self.notebook.add(tab, text="Threat Intelligence")

        # Входни данни
        input_frame = ttk.Frame(tab)
        input_frame.pack(fill=tk.X, padx=10, pady=10)

        ttk.Label(input_frame, text="IOC (Indicators of Compromise):").pack(
            side=tk.LEFT
        )
        self.ioc_entry = ttk.Entry(input_frame)
        self.ioc_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)

        # Бутони за проверка
        btn_frame = ttk.Frame(tab)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)

        checks = [
            ("Провери IP", "ip"),
            ("Провери Домейн", "domain"),
            ("Провери Хеш", "hash"),
            ("Провери URL", "url"),
        ]

        for i, (text, ioc_type) in enumerate(checks):
            btn = ttk.Button(
                btn_frame, text=text, command=lambda t=ioc_type: self.check_ioc(t)
            )
            btn.grid(row=0, column=i, padx=2, sticky=tk.EW)
            
        # Threat Intelligence източници
        #intel_frame = ttk.LabelFrame(tab, text="Източници", padding=10)
        #intel_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        #intel_sources = [
        #    ("AbuseIPDB Blacklist", self.get_abuseipdb_blacklist),
        #    ("AlienVault OTX", self.get_alienvault_otx),
        #    ("FireHOL IP Lists", self.get_firehol_iplists),
        #    ("CINSscore", self.get_cinsscore),
        #]

        #for i, (text, cmd) in enumerate(intel_sources):
        #    btn = ttk.Button(
        #        intel_frame,
        #        text=text,
        #        command=cmd
        #    )
        #    btn.grid(row=i//2, column=i%2, padx=5, pady=5, sticky=tk.EW)
        ttk.Button(btn_frame, text="Импортирай IOC от файл", command=self.import_iocs_from_file).grid(row=1, column=0, columnspan=4, pady=(10, 5), sticky=tk.EW)

        btn_select = ttk.Frame(tab)
        btn_select.pack(pady=5)
        ttk.Button(btn_select, text="☑️ Избери IP адреси", command=self.select_imported_ips).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_select, text="☑️ Избери домейни", command=self.select_imported_domains).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_select, text="☑️ Избери хешове", command=self.select_imported_hashes).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_select, text="☑️ Избери всички", command=self.select_all_iocs).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_select, text="📡 Сканирай избраните", command=self.scan_selected_iocs).pack(side=tk.LEFT, padx=5)


        # Резултати
        results_frame = ttk.Frame(tab)
        results_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.threat_intel_results = scrolledtext.ScrolledText(
            results_frame, wrap=tk.WORD, font=("Consolas", 10)
        )
        self.threat_intel_results.pack(fill=tk.BOTH, expand=True)
        self.threat_intel_results.drop_target_register(DND_FILES)
        self.threat_intel_results.dnd_bind('<<Drop>>', self.handle_dropped_file)

        #
        #intel_frame.grid_rowconfigure(2, weight=1)
        #intel_frame.grid_columnconfigure(0, weight=1)
        #intel_frame.grid_columnconfigure(1, weight=1)

    def import_iocs_from_file(self):
        file_path = filedialog.askopenfilename(filetypes=[
            ("Всички формати", "*.txt *.csv *.xlsx *.docx *.pdf"),
            ("Текстови файлове", "*.txt"),
            ("CSV файлове", "*.csv"),
            ("Excel файлове", "*.xlsx"),
            ("Word документи", "*.docx"),
            ("PDF документи", "*.pdf"),
        ])
        if file_path:
            self.import_iocs_from_path(file_path)

        try:
            text = ""
            if file_path.endswith(".txt"):
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            elif file_path.endswith(".csv"):
                df = pd.read_csv(file_path)
                text = "\n".join(df.astype(str).stack().tolist())
            elif file_path.endswith(".xlsx"):
                df = pd.read_excel(file_path)
                text = "\n".join(df.astype(str).stack().tolist())
            elif file_path.endswith(".docx"):
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif file_path.endswith(".pdf"):
                with fitz.open(file_path) as pdf:
                    for page in pdf:
                        text += page.get_text()
            else:
                messagebox.showwarning("Грешка", "Неподдържан файлов формат.")
                return

            # Извличане на IOC с regex
            ips = re.findall(r"\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b", text)
            domains = re.findall(r"\b(?:[a-zA-Z0-9.-]+\.(?:[a-z]{2,}))\b", text)
            hashes = re.findall(r"\b[a-fA-F0-9]{32,64}\b", text)

            # Премахване на дубликати
            ips = list(set(ips))
            domains = list(set(domains))
            hashes = list(set(hashes))

            # Визуализация
            self.threat_intel_results.insert(tk.END, f"[IOC Import] Файл: {file_path}\n")
            self.threat_intel_results.insert(tk.END, f"Открити IP адреси ({len(ips)}):\n" + "\n".join(ips) + "\n\n")
            self.threat_intel_results.insert(tk.END, f"Открити домейни ({len(domains)}):\n" + "\n".join(domains) + "\n\n")
            self.threat_intel_results.insert(tk.END, f"Открити хешове ({len(hashes)}):\n" + "\n".join(hashes) + "\n\n")

            self.status_bar.config(text="Импортирани IOC от файл успешно.")
        except Exception as e:
            messagebox.showerror("Грешка", f"Неуспешен импорт: {str(e)}")
            self.status_bar.config(text="Грешка при импорт на IOC.")
            
    def import_iocs_from_path(self, file_path):
        """Импортира IOC от файл с деобфускиране и показва и фалшиви домейни (.exe, .dll) отделно"""

        try:
            text = ""
            if file_path.endswith(".txt"):
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
            elif file_path.endswith(".csv"):
                df = pd.read_csv(file_path)
                text = "\n".join(df.astype(str).stack().tolist())
            elif file_path.endswith(".xlsx"):
                df = pd.read_excel(file_path)
                text = "\n".join(df.astype(str).stack().tolist())
            elif file_path.endswith(".docx"):
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            elif file_path.endswith(".pdf"):
                with fitz.open(file_path) as pdf:
                    for page in pdf:
                        text += page.get_text()
            else:
                messagebox.showwarning("Грешка", "Неподдържан файлов формат.")
                return

            # --- Деобфускиране ---
            text = re.sub(r"\[\.\]", ".", text)
            text = re.sub(r"\(\.\)", ".", text)
            text = re.sub(r"\{\.\}", ".", text)
            text = text.replace("hxxp://", "http://").replace("hxxps://", "https://")

            # --- Regex-и ---
            ip_regex = r"\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b"
            domain_regex = r"\b[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b"
            hash_regex = r"\b[a-fA-F0-9]{32,64}\b"

            ips = sorted(list(set(re.findall(ip_regex, text))))
            hashes = sorted(list(set(re.findall(hash_regex, text))))

            # --- Домейни и фалшиви домейни ---
            raw_domains = re.findall(domain_regex, text)
            bad_exts = (".exe", ".dll", ".bat", ".sys", ".cmd", ".log", ".ini", ".json", ".xml", ".txt")

            domains = []
            suspicious_domains = []

            for d in raw_domains:
                lower_d = d.lower()
                if any(lower_d.endswith(ext) for ext in bad_exts) or "\\" in lower_d or "/" in lower_d:
                    suspicious_domains.append(d)
                else:
                    domains.append(d)

            domains = sorted(list(set(domains)))
            suspicious_domains = sorted(list(set(suspicious_domains)))

            # --- Показване в резултати ---
            self.threat_intel_results.insert(tk.END, f"[IOC Import] Файл: {file_path}\n")
            self.threat_intel_results.insert(tk.END, f"✅ Открити IP адреси ({len(ips)}):\n" + "\n".join(ips) + "\n\n")
            self.threat_intel_results.insert(tk.END, f"✅ Открити домейни ({len(domains)}):\n" + "\n".join(domains) + "\n\n")
            self.threat_intel_results.insert(tk.END, f"✅ Открити хешове ({len(hashes)}):\n" + "\n".join(hashes) + "\n\n")

            if suspicious_domains:
                self.threat_intel_results.insert(tk.END, f"⚠️ Подозрителни 'домейни' (вероятно файлове, {len(suspicious_domains)}):\n" +
                                                "\n".join(suspicious_domains) + "\n\n")

            # --- Запазване за сканиране ---
            self.imported_ips = ips
            self.imported_domains = domains
            self.imported_hashes = hashes

            self.status_bar.config(text="Импортирани IOC успешно.")

        except Exception as e:
            messagebox.showerror("Грешка", f"Неуспешен импорт на файл: {str(e)}")
            self.status_bar.config(text="Грешка при импорт.")


    def select_imported_ips(self):
        self.selected_ips = self.imported_ips.copy()
        self.status_bar.config(text=f"Избрани IP адреси: {len(self.selected_ips)}")

    def select_imported_domains(self):
        self.selected_domains = self.imported_domains.copy()
        self.status_bar.config(text=f"Избрани домейни: {len(self.selected_domains)}")

    def select_imported_hashes(self):
        self.selected_hashes = self.imported_hashes.copy()
        self.status_bar.config(text=f"Избрани хешове: {len(self.selected_hashes)}")

    def scan_selected_iocs(self):
        self.threat_intel_results.insert(tk.END, "[Scan] Стартирано сканиране на избраните IOC...\n")

        for ip in self.selected_ips:
            result = self.lookup_ip_abuseipdb(ip)
            self.threat_intel_results.insert(tk.END, f"[IP] {ip} ➤ AbuseIPDB: {result}\n")

        for domain in self.selected_domains:
            result = self.security_tools.scan_url_virustotal(f"http://{domain}")
            self.threat_intel_results.insert(tk.END, f"[Domain] {domain} ➤ VT URL: {result}\n")

        for h in self.selected_hashes:
            result = self.scan_hash_virustotal(h)
            self.threat_intel_results.insert(tk.END, f"[Hash] {h} ➤ VT File: {result}\n")

        self.status_bar.config(text="Сканиране на избраните IOC завършено.")
    
    def scan_hash_virustotal(self, hash_value):
        try:
            api_key = self.config.get_api_key("VIRUSTOTAL_API_KEY")
            if not api_key:
                return "Грешка: липсва VirusTotal API ключ"

            headers = {"x-apikey": api_key}
            proxies = self.config.get_proxy_settings()
            url = f"https://www.virustotal.com/api/v3/files/{hash_value}"

            response = requests.get(
                url,
                headers=headers,
                proxies=proxies,
                timeout=15,
            )

            if response.status_code == 200:
                data = response.json()
                if "data" in data:
                    attributes = data["data"].get("attributes", {})
                    stats = attributes.get("last_analysis_stats", {})
                    malicious = stats.get("malicious", 0)
                    total = sum(stats.values())
                    detections = f"{malicious}/{total} детекции"
                    permalink = f"https://www.virustotal.com/gui/file/{hash_value}"
                    return detections + f" | [🔍 VT Link]({permalink})"
                else:
                    return "Грешка: липсва ключ 'data' в отговора"
            elif response.status_code == 404:
                return "Файлът не е намерен в базата на VirusTotal"
            elif response.status_code == 403:
                return "Грешка: Нямате достъп до този ресурс (проверете ключа)"
            else:
                return f"Грешка: {response.status_code} {response.reason}"

        except requests.exceptions.ProxyError as e:
            return f"Proxy грешка: {str(e)}"
        except requests.exceptions.RequestException as e:
            return f"Грешка при заявка: {str(e)}"
        except Exception as e:
            return f"Изключение: {str(e)}"


    def select_all_iocs(self):
        self.selected_ips = self.imported_ips.copy()
        self.selected_domains = self.imported_domains.copy()
        self.selected_hashes = self.imported_hashes.copy()
        self.status_bar.config(text="Избрани всички IOC-и.")

    def lookup_ip_abuseipdb(self, ip):
        try:
            api_key = self.config.get_api_key("ABUSEIP_API_KEY")
            headers = {
                "Key": api_key,
                "Accept": "application/json"
            }
            params = {
                "ipAddress": ip,
                "maxAgeInDays": 90  # или 30, според нуждата
            }

            proxies = self.config.get_proxy_settings()
            url = "https://api.abuseipdb.com/api/v2/check"

            response = requests.get(url, headers=headers, params=params, proxies=proxies, timeout=15)
            data = response.json()

            # === Различни варианти на отговор ===
            if response.status_code == 200 and "data" in data:
                score = data["data"].get("abuseConfidenceScore", "N/A")
                country = data["data"].get("countryCode", "??")
                isp = data["data"].get("isp", "N/A")
                return f"{score}% | {country} | {isp}"
            elif "errors" in data:
                detail = data["errors"][0].get("detail", "Грешка без обяснение")
                return f"Грешка: {detail}"
            elif "message" in data:
                return f"Грешка: {data['message']}"
            else:
                return f"Неочакван отговор: {data}"

        except requests.exceptions.RequestException as e:
            return f"Грешка при връзка: {e}"
        except Exception as e:
            return f"Изключение: {str(e)}"




    def check_ioc(self, ioc_type):
        """Проверява IOC според типа"""
        ioc = self.ioc_entry.get().strip()
        if not ioc:
            messagebox.showwarning("Грешка", "Моля, въведете IOC")
            return

        self.status_bar.config(text=f"Проверка на {ioc_type}: {ioc}...")
        self.root.update()

        self.threat_intel_results.delete(1.0, tk.END)
        
        self.threat_intel_results.drop_target_register(DND_FILES)
        self.threat_intel_results.dnd_bind('<<Drop>>', self.handle_dropped_file)

        try:
            if ioc_type == "ip":
                if not self._validate_ip(ioc):
                    self.threat_intel_results.insert(tk.END, "Невалиден IP адрес")
                    return

                result = self.security_tools.analyze_ip(ioc)
                self._display_ioc_results(result)

            elif ioc_type == "domain":
                if not self._validate_domain(ioc):
                    self.threat_intel_results.insert(tk.END, "Невалиден домейн")
                    return

                # Проверка с VirusTotal
                if self.config.get_api_key("VIRUSTOTAL_API_KEY"):
                    url = f"https://www.virustotal.com/api/v3/domains/{ioc}"
                    headers = {
                        "x-apikey": self.config.get_api_key("VIRUSTOTAL_API_KEY")
                    }

                    response = requests.get(
                        url,
                        headers=headers,
                        proxies=self.config.get_proxy_settings(),
                        verify=False,
                        timeout=15,
                    )

                    if response.status_code == 200:
                        data = response.json()
                        self.threat_intel_results.insert(
                            tk.END, "VirusTotal резултати:\n"
                        )
                        self.threat_intel_results.insert(
                            tk.END, f"Домейн: {data.get('data', {}).get('id')}\n"
                        )

                        last_analysis = (
                            data.get("data", {})
                            .get("attributes", {})
                            .get("last_analysis_stats", {})
                        )
                        self.threat_intel_results.insert(
                            tk.END,
                            f"Зловредни отзиви: {last_analysis.get('malicious', 0)}\n",
                        )

                        categories = (
                            data.get("data", {})
                            .get("attributes", {})
                            .get("categories", {})
                        )
                        if categories:
                            self.threat_intel_results.insert(
                                tk.END, "\nКатегории:\n"
                            )
                            for vendor, category in categories.items():
                                self.threat_intel_results.insert(
                                    tk.END, f"- {vendor}: {category}\n"
                                )
                    else:
                        self.threat_intel_results.insert(
                            tk.END,
                            f"Грешка при проверка на домейн: {response.status_code}\n",
                        )

                # DNS записи
                self.threat_intel_results.insert(tk.END, "\nDNS записи:\n")
                record_types = ["A", "AAAA", "MX", "NS", "TXT"]
                for record_type in record_types:
                    try:
                        answers = dns.resolver.resolve(ioc, record_type)
                        self.threat_intel_results.insert(
                            tk.END, f"{record_type}:\n"
                        )
                        for rdata in answers:
                            self.threat_intel_results.insert(tk.END, f"- {rdata}\n")
                    except (dns.resolver.NoAnswer, dns.resolver.NXDOMAIN):
                        continue
                    except Exception as e:
                        self.threat_intel_results.insert(
                            tk.END, f"Грешка при {record_type}: {str(e)}\n"
                        )

            elif ioc_type == "hash":
                hash_type = self._get_hash_type(ioc)   
                if not hash_type:
                    self.threat_intel_results.insert(
                        tk.END, "Невалиден хеш (трябва да е MD5, SHA1 или SHA256)"
                    )
                    return

                result = self.security_tools.scan_hybrid_analysis(hash_value=ioc)
                if "error" in result:
                    self.threat_intel_results.insert(
                        tk.END, f"Грешка: {result['error']}\n"
                    )
                else:
                    data = result.get("data", {})
                    self.threat_intel_results.insert(
                        tk.END, f"Хеш: {ioc} ({hash_type})\n"
                    )
                    self.threat_intel_results.insert(
                        tk.END,
                        f"Оценка на заплаха: {data.get('threat_score', 'N/A')}\n",
                    )

                    if data.get("domains"):
                        self.threat_intel_results.insert(
                            tk.END, "\nВредни домейни:\n"
                        )
                        for domain in data.get("domains", [])[:10]:
                            self.threat_intel_results.insert(
                                tk.END, f"- {domain}\n"
                            )

                    if data.get("hosts"):
                        self.threat_intel_results.insert(
                            tk.END, "\nВредни IP адреси:\n"
                        )
                        for host in data.get("hosts", [])[:10]:
                            self.threat_intel_results.insert(tk.END, f"- {host}\n")

            elif ioc_type == "url":
                if not self._validate_url(ioc):
                    self.threat_intel_results.insert(tk.END, "Невалиден URL")
                    return

                result = self.security_tools.scan_url_virustotal(ioc)
                if "error" in result:
                    self.threat_intel_results.insert(
                        tk.END, f"Грешка: {result['error']}\n"
                    )
                else:
                    self.threat_intel_results.insert(tk.END, f"URL: {ioc}\n")
                    self.threat_intel_results.insert(
                        tk.END,
                        f"Зловредни отзиви: {result['positives']}/{result['total']}\n",
                    )
                    self.threat_intel_results.insert(
                        tk.END, f"Пермалинк: {result['permalink']}\n"
                    )

            self.status_bar.config(text=f"Проверката на {ioc_type} завършена")
        except Exception as e:
            self.threat_intel_results.insert(tk.END, f"Грешка: {str(e)}")
            self.status_bar.config(text=f"Грешка при проверка на IOC")

    def handle_dropped_file(self, event):
        file_path = event.data.strip().strip("{").strip("}")
        if os.path.isfile(file_path):
            self.import_iocs_from_path(file_path)
        else:
            messagebox.showerror("Грешка", f"Файлът не е валиден: {file_path}")


    def _validate_ip(self, ip):
        """Валидира IP адрес"""
        try:
            ipaddress.ip_address(ip)
            return True
        except ValueError:
            return False

    def _validate_domain(self, domain):
        """Валидира домейн"""
        pattern = re.compile(r"^((?!-)[A-Za-z0-9-]{1,63}(?<!-)\.)+[A-Za-z]{2,6}$")
        return bool(pattern.match(domain))

    def _validate_url(self, url):
        """Валидира URL"""
        pattern = re.compile(
            r"^(https?://)?"  # протокол
            r"(([A-Za-z0-9-]+\.)+[A-Za-z]{2,6})"  # домейн
            r"(/[A-Za-z0-9-._~:/?#[\]@!$&\'()*+,;%=]*)?$"
        )  # път
        return bool(pattern.match(url))

    def _get_hash_type(self, hash_value):
        """Определя типа на хеша"""
        if len(hash_value) == 32 and all(
            c in "0123456789abcdef" for c in hash_value.lower()
        ):
            return "md5"
        elif len(hash_value) == 40 and all(
            c in "0123456789abcdef" for c in hash_value.lower()
        ):
            return "sha1"
        elif len(hash_value) == 64 and all(
            c in "0123456789abcdef" for c in hash_value.lower()
        ):
            return "sha256"
        return None

    def _display_ioc_results(self, results):
        """Показва резултати от IOC проверка"""
        for service, data in results.items():
            self.threat_intel_results.insert(
                tk.END, f"\n=== {service.upper()} ===\n\n"
            )

            if "error" in data:
                self.threat_intel_results.insert(
                    tk.END, f"Грешка: {data['error']}\n"
                )
            else:
                for key, value in data.items():
                    if isinstance(value, dict):
                        self.threat_intel_results.insert(tk.END, f"{key}:\n")
                        for subkey, subvalue in value.items():
                            self.threat_intel_results.insert(
                                tk.END, f"  {subkey}: {subvalue}\n"
                            )
                    else:
                        self.threat_intel_results.insert(
                            tk.END, f"{key}: {value}\n"
                        )

    def get_alienvault_otx(self):
        """Взима данни от AlienVault OTX"""
        self.intel_result.delete(1.0, tk.END)
        self.intel_result.insert(tk.END, "Взима данни от AlienVault OTX...\n")
        self.root.update()

        try:
            result = self.security_tools.check_alienvault_otx()
            if "error" in result:
                self.intel_result.insert(tk.END, f"Грешка: {result['error']}")
            else:
                for indicator in result.get("indicators", []):
                    self.intel_result.insert(
                        tk.END, 
                        f"{indicator.get('indicator')}: {indicator.get('type')}\n"
                    )
        except Exception as e:
            self.intel_result.insert(tk.END, f"Грешка: {str(e)}")

    def get_firehol_iplists(self):
        """Взима IP листи от FireHOL"""
        self.intel_result.delete(1.0, tk.END)
        self.intel_result.insert(tk.END, "Взима FireHOL IP листи...\n")
        self.root.update()

        try:
            result = self.security_tools.check_firehol_iplists()
            if "error" in result:
                self.intel_result.insert(tk.END, f"Грешка: {result['error']}")
            else:
                for list_name, ips in result.items():
                    self.intel_result.insert(
                        tk.END, 
                        f"{list_name}: {len(ips)} IP адреса\n"
                    )
        except Exception as e:
            self.intel_result.insert(tk.END, f"Грешка: {str(e)}")
            
    def get_cinsscore(self):
        """Взима CINSscore репутация"""
        self.intel_result.delete(1.0, tk.END)
        self.intel_result.insert(tk.END, "Взима CINSscore данни...\n")
        self.root.update()

        try:
            result = self.security_tools.check_cinsscore()
            if "error" in result:
                self.intel_result.insert(tk.END, f"Грешка: {result['error']}")
            else:
                for ip, score in result.items():
                    self.intel_result.insert(
                        tk.END, 
                        f"{ip}: {score} репутация\n"
                    )
        except Exception as e:
            self.intel_result.insert(tk.END, f"Грешка: {str(e)}")

    def handle_check_phishing(self):
        url = self.phishing_url_entry.get().strip()
        if not url:
            messagebox.showwarning("Грешка", "Моля, въведете URL за проверка.")
            return
        self.log_activity(f"Стартирана фишинг проверка за {url}")

        self.update_status(f"Проверка на фишинг URL: {url}")
        results = self.check_phishing(url)

        vt = results["virus_total"]
        ha = results["hybrid_analysis"]
        heur = results["heuristic"]

        vt_str = vt.get("detections_str", "Няма данни")
        vt_link = vt.get("permalink", "")
        ha_summary = "Намерени данни" if ha.get("data") else "Няма данни"
        heur_str = ", ".join(heur.get("keyword_matches", []))
        heur_flag = "⚠️ Подозрителен URL" if heur.get("suspicious") else "Няма директни индикации"

        output = (
            f"--- VirusTotal ---\n"
            f"Детекции: {vt_str}\n"
            f"Линк: {vt_link}\n\n"
            f"--- Hybrid Analysis ---\n"
            f"{ha_summary}\n\n"
            f"--- Евристика ---\n"
            f"Ключови думи: {heur_str}\n"
            f"{heur_flag}\n"
        )

        self.phishing_text_results.delete(1.0, tk.END)
        self.phishing_text_results.insert(tk.END, output)

        # Ако има над 5 детекции, предложи автоматично докладване
        try:
            detected, total = map(int, vt_str.split("/"))
            if detected >= 5:
                should_report = messagebox.askyesno(
                    "Открит фишинг!",
                    f"VirusTotal откри {detected} детекции от {total}.\n"
                    f"Искаш ли да го докладваме автоматично към Google?"
                )
                if should_report:
                    self.report_phishing_external("google")
        except:
            pass  # Ако форматът е неочакван – игнорирай
    

    def create_mxtoolbox_tab(self):
        """Създава таб с MX Toolbox инструменти"""
        tab = ttk.Frame(self.notebook)
        self.mxtoolbox_tab = tab
        self.notebook.add(tab, text="MX Toolbox")

        # Входни данни
        input_frame = ttk.Frame(tab)
        input_frame.pack(fill=tk.X, padx=10, pady=10)

        ttk.Label(input_frame, text="Домейн/IP:").pack(side=tk.LEFT)
        self.mxtoolbox_entry = ttk.Entry(input_frame)
        self.mxtoolbox_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)

        # Проверки
        check_frame = ttk.LabelFrame(tab, text="Проверки", padding=10)
        check_frame.pack(fill=tk.X, padx=10, pady=5)

        checks = [
            ("DNS (A)", "a"),
            ("DNS (MX)", "mx"),
            ("DNS (TXT)", "txt"),
            ("Blacklist", "blacklist"),
            ("SPF", "spf"),
            ("DKIM", "dkim"),
            ("DMARC", "dmarc"),
            ("Whois", "whois"),
        ]

        for i, (text, check_type) in enumerate(checks):
            btn = ttk.Button(
                check_frame,
                text=text,
                command=lambda ct=check_type: self.run_mxtoolbox_check(ct),
                width=10,
            )
            btn.grid(row=i // 4, column=i % 4, padx=2, pady=2)

        # Резултати
        results_frame = ttk.Frame(tab)
        results_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.mxtoolbox_results = scrolledtext.ScrolledText(
            results_frame, wrap=tk.WORD, font=("Consolas", 10)
        )
        self.mxtoolbox_results.pack(fill=tk.BOTH, expand=True)

    def run_mxtoolbox_check(self, check_type):
        """Изпълнява проверка с MX Toolbox"""
        query = self.mxtoolbox_entry.get().strip()
        if not query:
            messagebox.showwarning("Грешка", "Моля, въведете домейн или IP")
            return
        self.log_activity(f"MXToolbox заявка: {query} [{check_type}]")
        self.status_bar.config(text=f"Проверка {check_type} за {query}...")
        self.root.update()

        result = self.security_tools.check_mxtoolbox(query, check_type)
        self.mxtoolbox_results.delete(1.0, tk.END)

        if "error" in result:
            self.mxtoolbox_results.insert(tk.END, f"Грешка: {result['error']}")
            return

        data = result.get("data", {})
        output = f"Резултати от {check_type} проверка за {query}:\n\n"

        if isinstance(data, dict):
            if check_type == "blacklist" and "Results" in data:
                for item in data["Results"]:
                    output += f"- {item.get('Name', '')}: {item.get('Status', '')}\n"

            elif "Records" in data and isinstance(data["Records"], list):
                output += "--- Records ---\n"
                for record in data["Records"]:
                    output += f"- {record}\n"

            for key in ["Information", "Passed", "Warnings", "Errors", "Transcript"]:
                if key in data and data[key]:
                    output += f"\n--- {key} ---\n"
                    if isinstance(data[key], list):
                        for entry in data[key]:
                            output += json.dumps(entry, indent=2, ensure_ascii=False) + "\n"
                    else:
                        output += str(data[key]) + "\n"

        if output.strip() == f"Резултати от {check_type} проверка за {query}:":
            output += "\nНяма открити данни."

        self.mxtoolbox_results.insert(tk.END, output)
        self.status_bar.config(text=f"Проверката {check_type} завършена")

    def create_powershell_tab(self):
        
        
         # Drag & Drop .ps1 файлове
        def handle_drop(event):
            path = event.data.strip("{}")
            if path.endswith(".ps1"):
                self.run_ps_script(path)
            else:
                self.update_status("Неподдържан файл: само .ps1")

        
        ps_tab = ttk.Frame(self.notebook)
        self.powershell_tab = ps_tab
        self.notebook.add(ps_tab, text="PowerShell")

        self.ps_command_history = []
        self.ps_history_index = -1


        self.ps_output = scrolledtext.ScrolledText(ps_tab, height=22, bg="black", fg="lime", insertbackground="white", undo=True, wrap=tk.WORD)
        self.ps_output.drop_target_register(DND_FILES)
        self.ps_output.dnd_bind('<<Drop>>', handle_drop)
        self.ps_output.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.ps_output.bind("<Control-a>", lambda e: (self.ps_output.tag_add(tk.SEL, "1.0", tk.END), self.ps_output.mark_set(tk.INSERT, "1.0"), self.ps_output.see(tk.INSERT), "break"))
        self.ps_output.bind("<Control-c>", lambda e: self.root.clipboard_append(self.ps_output.get(tk.SEL_FIRST, tk.SEL_LAST)) if self.ps_output.tag_ranges(tk.SEL) else None)
        self.ps_output.insert(tk.END, "PowerShell терминал готов.\n")
        self.ps_output.tag_config("output", foreground="lime")
        self.ps_output.tag_config("error", foreground="red")
        self.ps_output.config(state=tk.DISABLED)

        input_frame = ttk.Frame(ps_tab)
        input_frame.pack(fill=tk.X, padx=5, pady=5)

        self.powershell_commands = ["Get-Process", "Get-Service", "whoami", "cd", "dir", "type", "echo", "mkdir", "ren", "move", "copy", "xcopy", "del", "rd", "IWR", "IEX", "Start-Process", "Set-ExecutionPolicy", "Get-ChildItem", "Get-ItemProperty", "Set-ItemProperty", "Test-Connection", "Resolve-DnsName", "Get-Command", "Get-Help", "Format-Table", "Format-List", "Sort-Object", "Measure", "ForEach-Object", "Where-Object", "net user", "net localgroup", "reg query", "icacls", "schtasks", "tasklist", "taskkill", "wmic", "Get-CimInstance", "Set-NetFirewallProfile", "Get-NetFirewallRule"]

        self.ps_autocomplete = tk.Listbox(ps_tab, height=6, bg="black", fg="white")
        self.ps_autocomplete.pack_forget()

        self.ps_entry = ttk.Entry(input_frame)
        self.ps_entry.drop_target_register(DND_FILES)
        self.ps_entry.dnd_bind('<<Drop>>', handle_drop)
        self.ps_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        self.ps_entry.bind("<Return>", self.execute_ps_command)
        self.ps_entry.bind("<KeyRelease>", self.update_autocomplete)
        self.ps_autocomplete.bind("<ButtonRelease-1>", self.fill_autocomplete)
        self.ps_entry.bind("<Down>", lambda e: self.ps_autocomplete.focus_set())
        self.ps_autocomplete.bind("<Return>", self.fill_autocomplete)
        self.ps_autocomplete.bind("<Right>", self.fill_autocomplete)
        self.ps_entry.bind("<Control-v>", lambda e: self.ps_entry.insert(tk.INSERT, self.root.clipboard_get()))
        self.ps_entry.bind("<Control-c>", lambda e: self.root.clipboard_append(self.ps_entry.get()))
        
        self.ps_entry.bind("<Up>", self.navigate_command_history_up)
        self.ps_entry.bind("<Down>", self.navigate_command_history_down)

        ttk.Button(input_frame, text="Изпрати", command=self.execute_ps_command).pack(side=tk.LEFT, padx=5)
        ttk.Button(input_frame, text="Изпълни .ps1", command=self.run_ps_script_dialog).pack(side=tk.LEFT, padx=5)
        ttk.Button(input_frame, text="Изчисти", command=self.clear_ps_output).pack(side=tk.LEFT, padx=5)
        ttk.Button(input_frame, text="Запази лог", command=self.export_ps_output).pack(side=tk.LEFT, padx=5)
        tk.Button(ps_tab, text="Спри изпълнението", command=self.stop_ps_execution).pack()

        guide_frame = ttk.LabelFrame(ps_tab, text="Чести команди")
        guide_frame.pack(fill=tk.X, padx=5, pady=5)

        common_cmds = ["Get-Process", "Get-Service", "whoami", "Get-NetTCPConnection", "Get-WinEvent", "Resolve-DnsName google.com", "Start-Process powershell -Verb runAs", "net user B00*** /domain", "Get-ADUser -Filter * > ADUsersInfo.txt"]

        for i, cmd in enumerate(common_cmds):
            btn = ttk.Button(guide_frame, text=cmd, command=lambda c=cmd: self.insert_command(c))
            btn.grid(row=i//3, column=i%3, padx=3, pady=3, sticky="ew")

        for i in range(3):
            guide_frame.grid_columnconfigure(i, weight=1)

        ttk.Separator(ps_tab, orient="horizontal").pack(fill=tk.X, padx=5, pady=5)
        bottom_frame = ttk.Frame(ps_tab)
        bottom_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(bottom_frame, text="Clear", command=self.clear_ps_output).pack(side=tk.LEFT, padx=5)
        ttk.Button(bottom_frame, text="Export", command=self.export_ps_output).pack(side=tk.LEFT, padx=5)
        ttk.Button(bottom_frame, text="Admin", command=lambda: self.insert_command("Start-Process powershell -Verb runAs")).pack(side=tk.LEFT, padx=5)

        self.ps_entry.bind('<Up>', self.navigate_ps_history_up)
        self.ps_entry.bind('<Down>', self.navigate_ps_history_down)

    def insert_command(self, cmd):
        self.ps_entry.delete(0, tk.END)
        self.ps_entry.insert(0, cmd)

    def execute_ps_command(self, event=None):
        command = self.ps_entry.get().strip()
        self.log_activity(f"PowerShell команда: {command}")
        if not command:
            return
        self.ps_command_history.append(command)
        self.ps_history_index = len(self.ps_command_history)

        self.ps_output.config(state=tk.NORMAL)
        self.ps_output.insert(tk.END, f"> {command}\n")
        self.ps_output.config(state=tk.DISABLED)
        self.ps_output.see(tk.END)
        self.ps_entry.delete(0, tk.END)
        
        def run_command():
            try:
                CREATE_NO_WINDOW = 0x08000000
                process = subprocess.Popen(
                    ["powershell", "-Command", command],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.STDOUT,
                    text=True,
                    creationflags=CREATE_NO_WINDOW
                )

                for line in process.stdout:
                    self.root.after(0, lambda line=line: self._append_ps_output(line))

                process.stdout.close()
                process.wait()
            except Exception as e:
                self.root.after(0, lambda: self._append_ps_output(f"❌ Грешка: {str(e)}\n"))

        threading.Thread(target=run_command, daemon=True).start()
        
    def _append_ps_output(self, line):
        self.ps_output.config(state=tk.NORMAL)
        tag = "error" if "error" in line.lower() else "output"
        self.ps_output.insert(tk.END, line, tag)
        self.ps_output.see(tk.END)
        self.ps_output.config(state=tk.DISABLED)

    def run_ps_script(self, filepath):
        if not filepath.endswith(".ps1") or not os.path.exists(filepath):
            self._append_ps_output("❌ Невалиден PowerShell скрипт.\n")
            return
        self.log_activity(f"Стартиран Powershell scrypt")
        self.ps_output.config(state=tk.NORMAL)
        self.ps_output.insert(tk.END, f"> Изпълнение на: {filepath}\n")
        self.ps_output.config(state=tk.DISABLED)
        self.ps_output.see(tk.END)

    def run_ps_script_dialog(self):
        filepath = filedialog.askopenfilename(
            filetypes=[("PowerShell Scripts", "*.ps1")],
            title="Избери PowerShell скрипт"
        )
        if filepath:
            self.run_ps_script(filepath)

    def run_ps_script(self, filepath):
        if not filepath.endswith(".ps1") or not os.path.exists(filepath):
            self._append_ps_output("❌ Невалиден PowerShell скрипт.\n")
            return

        self.ps_output.config(state=tk.NORMAL)
        self.ps_output.insert(tk.END, f"> Изпълнение на: {filepath}\n")
        self.ps_output.config(state=tk.DISABLED)
        self.ps_output.see(tk.END)

        def run_script():
            try:
                CREATE_NO_WINDOW = 0x08000000
                cmd = ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", filepath]
                print(f"[DEBUG] CMD: {' '.join(cmd)}")

                self.ps_process = subprocess.Popen(
                    cmd,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.STDOUT,
                    text=True,
                    creationflags=CREATE_NO_WINDOW
                )

                for line in self.ps_process.stdout:
                    print(f"[DEBUG] {line.strip()}")
                    self.root.after(0, lambda line=line: self._append_ps_output(line))

                self.ps_process.stdout.close()
                self.ps_process.wait()
            except Exception as e:
                self.root.after(0, lambda: self._append_ps_output(f"❌ Грешка: {str(e)}\n"))
            finally:
                self.ps_process = None

        threading.Thread(target=run_script, daemon=True).start()

    def navigate_command_history_up(self, event=None):
        if self.ps_command_history and self.ps_history_index > 0:
            self.ps_history_index -= 1
            self.ps_entry.delete("1.0", tk.END)
            self.ps_entry.insert(tk.END, self.ps_command_history[self.ps_history_index])
        return "break"

    def navigate_command_history_down(self, event=None):
        if self.ps_command_history and self.ps_history_index < len(self.ps_command_history) - 1:
            self.ps_history_index += 1
            self.ps_entry.delete("1.0", tk.END)
            self.ps_entry.insert(tk.END, self.ps_command_history[self.ps_history_index])
        else:
            self.ps_entry.delete("1.0", tk.END)
        return "break"

    def clear_ps_output(self):
        self.ps_output.config(state=tk.NORMAL)
        self.ps_output.delete("1.0", tk.END)
        self.ps_output.insert(tk.END, "PowerShell терминал готов.")
        self.ps_output.config(state=tk.DISABLED)
        
    def stop_ps_execution(self):
        if self.ps_process and self.ps_process.poll() is None:
            self.ps_process.terminate()
            self._append_ps_output("⛔ PowerShell процесът беше прекратен.\n")
            self.ps_process = None
        else:
            self._append_ps_output("⚠️ Няма активен PowerShell процес.\n")

    def update_autocomplete(self, event=None):
        typed = self.ps_entry.get()
        self.ps_autocomplete.select_set(0)
        self.ps_entry.bind("<Down>", lambda e: self.ps_autocomplete.focus_set())
        self.ps_autocomplete.bind("<Return>", self.fill_autocomplete)
        self.ps_autocomplete.bind("<Right>", self.fill_autocomplete)
        self.ps_autocomplete.bind("<ButtonRelease-1>", self.fill_autocomplete)
        if not typed:
            self.ps_autocomplete.pack_forget()
            return

        matches = [cmd for cmd in self.powershell_commands if cmd.lower().startswith(typed.lower())]
        if matches:
            self.ps_autocomplete.delete(0, tk.END)
            for match in matches:
                self.ps_autocomplete.insert(tk.END, match)
            self.ps_autocomplete.place(x=self.ps_entry.winfo_rootx() - self.root.winfo_rootx(),
                                       y=self.ps_entry.winfo_rooty() - self.root.winfo_rooty() + 25,
                                       width=self.ps_entry.winfo_width())
        else:
            self.ps_autocomplete.pack_forget()

    def fill_autocomplete(self, event=None):
        try:
            selected = self.ps_autocomplete.get(tk.ACTIVE)
            self.ps_entry.delete(0, tk.END)
            self.ps_entry.insert(0, selected)
            self.ps_autocomplete.pack_forget()
        except tk.TclError:
            pass


        self.ps_output.config(state=tk.NORMAL)
        self.ps_output.delete("1.0", tk.END)
        self.ps_output.insert(tk.END, "PowerShell терминал готов.\n")
        self.ps_output.config(state=tk.DISABLED)

    def export_ps_output(self):
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text Files", "*.txt")])
        if path:
            with open(path, "w", encoding="utf-8") as f:
                f.write(self.ps_output.get("1.0", tk.END))

        if self.ps_history_index < len(self.ps_command_history) - 1:
            self.ps_history_index += 1
            self.ps_entry.delete(0, tk.END)
            self.ps_entry.insert(0, self.ps_command_history[self.ps_history_index])
        else:
            self.ps_entry.delete(0, tk.END)



    #def __init__(self, root):
        # ... останалия init код ...
        #self.setup_ad_tab()
        #self.is_admin = self.check_admin_privileges()
            


    def setup_ad_tab(self):
        """Създава раздел за Active Directory инструменти"""
        self.ad_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.ad_tab, text='Active Directory')

        # Header
        ttk.Label(self.ad_tab, text="Active Directory Инструменти", style='Header.TLabel').pack(pady=10)

        # Input Frame
        input_frame = ttk.Frame(self.ad_tab)
        input_frame.pack(fill=tk.X, padx=10, pady=5)

        # Domain input
        ttk.Label(input_frame, text="Домейн:").grid(row=0, column=0, sticky='e', padx=5)
        self.ad_domain = ttk.Entry(input_frame)
        self.ad_domain.grid(row=0, column=1, sticky='ew', padx=5)
        self.ad_domain.insert(0, socket.getfqdn().split('.', 1)[1] if '.' in socket.getfqdn() else "")

        # Username input
        ttk.Label(input_frame, text="Потребител:").grid(row=1, column=0, sticky='e', padx=5)
        self.ad_username = ttk.Entry(input_frame)
        self.ad_username.grid(row=1, column=1, sticky='ew', padx=5)
        self.ad_username.insert(0, os.getenv('USERNAME') or "")

        # Admin credentials
        ttk.Label(input_frame, text="Админ потребител:").grid(row=2, column=0, sticky='e', padx=5)
        self.ad_admin_user = ttk.Entry(input_frame)
        self.ad_admin_user.grid(row=2, column=1, sticky='ew', padx=5)

        ttk.Label(input_frame, text="Админ парола:").grid(row=3, column=0, sticky='e', padx=5)
        self.ad_admin_pass = ttk.Entry(input_frame, show="*")
        self.ad_admin_pass.grid(row=3, column=1, sticky='ew', padx=5)

        # Buttons
        btn_frame = ttk.Frame(self.ad_tab)
        btn_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Button(btn_frame, text="Търси потребител", command=self.query_ad_user).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Провери парола", command=self.check_password_expiry).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Отключи", command=self.unlock_ad_user, style='Red.TButton').pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Смени парола", command=self.reset_ad_password, style='Green.TButton').pack(side=tk.LEFT, padx=5)

        # Results Treeview with all original columns
        tree_frame = ttk.Frame(self.ad_tab)
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        columns = (
            'DC Name',               # Име на контролер
            'Site',                  # Сайт
            'User State',            # Състояние на акаунт
            'Bad Pwd Count',         # Грешни пароли
            'Last Bad Pwd',          # Последна грешна парола
            'Pwd Last Set',          # Парола сменена на
            'Lockout Time',          # Време на заключване
            'Orig Lock',             # Оригинално заключване
            'Attribute',             # Атрибут
            'Value'                  # Стойност
        )

        self.ad_tree = ttk.Treeview(tree_frame, columns=columns, show='headings')

        # Дефиниране на колоните с български етикети
        column_labels = {
            'DC Name': 'DC Име',
            'Site': 'Сайт',
            'User State': 'Състояние',
            'Bad Pwd Count': 'Грешни пароли',
            'Last Bad Pwd': 'Последна грешна парола',
            'Pwd Last Set': 'Парола сменена на',
            'Lockout Time': 'Заключен до',
            'Orig Lock': 'Оригинално заключване',
            'Attribute': 'Атрибут',
            'Value': 'Стойност'
        }

        for col in columns:
            self.ad_tree.heading(col, text=column_labels.get(col, col))
            self.ad_tree.column(col, width=120, stretch=True)

        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=self.ad_tree.yview)
        self.ad_tree.configure(yscrollcommand=scrollbar.set)

        self.ad_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Status Bar
        self.ad_status = ttk.Label(self.ad_tab, text="Готов", relief=tk.SUNKEN)
        self.ad_status.pack(fill=tk.X, padx=10, pady=5)

    def query_ad_user(self):
        """Търси потребител в Active Directory"""
        username = self.ad_username.get().strip()
        self.log_activity(f"Проверка на AD Lockout за: {username}")
        if not username:
            messagebox.showwarning("Грешка", "Моля въведете потребителско име!")
            return

        domain = self.ad_domain.get().strip()
        if not domain:
            messagebox.showwarning("Грешка", "Моля въведете домейн!")
            return

        try:
            # Initialize AD query
            q = adquery.ADQuery()
            
            # Execute query with all needed attributes
            q.execute_query(
                attributes=[
                    'cn', 'displayName', 'givenName', 'sn', 'mail',
                    'telephoneNumber', 'title', 'department',
                    'userPrincipalName', 'sAMAccountName',
                    'whenCreated', 'whenChanged', 'lastLogon',
                    'badPwdCount', 'pwdLastSet', 'accountExpires',
                    'userAccountControl', 'memberOf', 'lockoutTime',
                    'lastLogonTimestamp', 'badPasswordTime',
                    'distinguishedName', 'objectSid'
                ],
                where_clause=f"sAMAccountName='{username}'"
            )

            # Clear previous results
            self.ad_tree.delete(*self.ad_tree.get_children())

            # Process results and populate all columns
            for row in q.get_results():
                # Get basic user info
                cn = row.get('cn', [''])[0] if isinstance(row.get('cn'), list) else row.get('cn', '')

                bad_pwd_count_raw = row.get('badPwdCount', 0)
                bad_pwd_count = bad_pwd_count_raw[0] if isinstance(bad_pwd_count_raw, list) else bad_pwd_count_raw

                last_bad_pwd_raw = row.get('badPasswordTime', [0])
                last_bad_pwd = self.convert_ad_timestamp(last_bad_pwd_raw[0] if isinstance(last_bad_pwd_raw, list) else last_bad_pwd_raw)

                pwd_last_set_raw = row.get('pwdLastSet', [0])
                pwd_last_set = self.convert_ad_timestamp(pwd_last_set_raw[0] if isinstance(pwd_last_set_raw, list) else pwd_last_set_raw)

                lockout_time_raw = row.get('lockoutTime', [0])
                lockout_time = self.convert_ad_timestamp(lockout_time_raw[0] if isinstance(lockout_time_raw, list) else lockout_time_raw)

                user_state_raw = row.get('userAccountControl', [0])
                user_state = self.decode_user_account_control(user_state_raw[0] if isinstance(user_state_raw, list) else user_state_raw)

                
                # Get additional info for other columns
                dc_name = os.environ.get('LOGONSERVER', '').lstrip('\\')
                if not dc_name:
                    dc_name = socket.getfqdn()  # fallback ако няма LOGONSERVER
                site = "Default-First-Site-Name"  # Default site (can be enhanced)
                orig_lock = lockout_time  # Same as lockout time in this implementation
                
                # Insert main row with DC info
                self.ad_tree.insert('', 'end', values=(
                    dc_name,
                    site,
                    user_state,
                    bad_pwd_count,
                    last_bad_pwd,
                    pwd_last_set,
                    lockout_time,
                    orig_lock,
                    'sAMAccountName',
                    cn
                ))
                
                # Insert additional attributes as separate rows
                for attr, value in row.items():
                    if isinstance(value, list):
                        value = value[0] if value else ""
                    
                    # Skip attributes already shown in main columns
                    if attr in ['badPwdCount', 'badPasswordTime', 'pwdLastSet', 
                               'lockoutTime', 'userAccountControl', 'cn', 'sAMAccountName']:
                        continue
                    
                    # Format special attributes
                    if attr in ['lastLogon', 'lastLogonTimestamp', 'accountExpires', 'whenCreated', 'whenChanged']:
                        value = self.convert_ad_timestamp(value)
                    elif attr == 'memberOf' and isinstance(value, list):
                        value = '\n'.join(value)
                    
                    self.ad_tree.insert('', 'end', values=(
                        '',  # DC Name
                        '',  # Site
                        '',  # User State
                        '',  # Bad Pwd Count
                        '',  # Last Bad Pwd
                        '',  # Pwd Last Set
                        '',  # Lockout Time
                        '',  # Orig Lock
                        attr,
                        str(value)
                    ))

            self.ad_status.config(text=f"Намерен потребител: {username}")
            self.update_status(f"Намерен потребител: {username}")

        except Exception as e:
            messagebox.showerror("Грешка", f"Грешка при търсене: {str(e)}")
            self.update_status(f"Грешка при търсене: {str(e)}")

    def unlock_ad_user(self):
        """Отключва заключен AD акаунт"""
        if not self.is_admin:
            messagebox.showerror("Грешка", "Необходими са администраторски права!")
            return

        username = self.ad_username.get().strip()
        self.log_activity(f"Отключване на: {username}")
        if not username:
            messagebox.showwarning("Грешка", "Моля въведете потребителско име!")
            return

        domain = self.ad_domain.get().strip()
        if not domain:
            messagebox.showwarning("Грешка", "Моля въведете домейн!")
            return

        try:
            # Get user object
            user = aduser.ADUser.from_cn(username)
            
            # Check if user is locked
            if not self.is_user_locked(user):
                messagebox.showinfo("Информация", f"Потребител {username} не е заключен")
                return
            
            # Unlock user
            user.unlock()
            
            messagebox.showinfo("Успех", f"Потребител {username} беше отключен успешно")
            self.query_ad_user()  # Refresh user data
            self.update_status(f"Потребител {username} отключен успешно")

        except Exception as e:
            messagebox.showerror("Грешка", f"Грешка при отключване: {str(e)}")
            self.update_status(f"Грешка при отключване: {str(e)}")

    def is_user_locked(self, user):
        """Проверява дали потребителят е заключен"""
        try:
            # Get user attributes
            user.update_attributes(['lockoutTime'])
            lockout_time = user.get_attribute('lockoutTime', 0)
            
            return lockout_time > 0
        except:
            return False

    def reset_ad_password(self):
        """Ресетва паролата на потребител"""
        if not self.is_admin:
            messagebox.showerror("Грешка", "Необходими са администраторски права!")
            return

        username = self.ad_username.get().strip()
        self.log_activity(f"Резетване на AD Парола за: {username}")
        if not username:
            messagebox.showwarning("Грешка", "Моля въведете потребителско име!")
            return

        new_password = simpledialog.askstring(
            "Смяна на парола", 
            f"Въведете нова парола за {username}:",
            show='*'
        )
        
        if not new_password:
            return

        try:
            # Get user object
            user = aduser.ADUser.from_cn(username)
            
            # Set new password
            user.set_password(new_password)
            
            # Optionally force password change at next login
            user.update_attributes({'pwdLastSet': 0})
            
            messagebox.showinfo("Успех", f"Паролата за {username} беше сменена успешно")
            self.query_ad_user()  # Refresh user data
            self.update_status(f"Паролата за {username} беше сменена")

        except Exception as e:
            messagebox.showerror("Грешка", f"Грешка при смяна на парола: {str(e)}")
            self.update_status(f"Грешка при смяна на парола: {str(e)}")

    def check_password_expiry(self):
        """Проверява кога изтича паролата на потребителя"""
        username = self.ad_username.get().strip()
        self.log_activity(f"Проверка на парола за: {username}")
        if not username:
            messagebox.showwarning("Грешка", "Моля въведете потребителско име!")
            return

        try:
            # Get user object
            user = aduser.ADUser.from_cn(username)
            
            # Get password expiry
            expiry_date = user.get_password_expiry()
            
            if expiry_date:
                now = datetime.now()
                days_left = (expiry_date - now).days
                
                messagebox.showinfo(
                    "Изтичане на парола",
                    f"Паролата изтича на: {expiry_date.strftime('%Y-%m-%d %H:%M:%S')}\n"
                    f"Оставащи дни: {days_left}"
                )
                self.update_status(f"Парола изтича на: {expiry_date.strftime('%Y-%m-%d')} ({days_left} дни остават)")
            else:
                messagebox.showinfo("Изтичане на парола", "Паролата не изтича")
                self.update_status("Паролата не изтича")

        except Exception as e:
            messagebox.showerror("Грешка", f"Грешка при проверка: {str(e)}")
            self.update_status(f"Грешка при проверка на парола: {str(e)}")

    def convert_ad_timestamp(self, timestamp):
        """Конвертира AD timestamp към четим формат, независимо от тип"""
        import datetime

        try:
            original_timestamp = timestamp

            # Ако е списък – вземи първия
            if isinstance(timestamp, list):
                timestamp = timestamp[0] if timestamp else 0

            # Ако вече е datetime обект
            if isinstance(timestamp, datetime.datetime):
                return timestamp.replace(tzinfo=None).strftime("%Y-%m-%d %H:%M:%S")

            # Ако е COM обект с .Value
            if hasattr(timestamp, 'Value'):
                timestamp = timestamp.Value

            # Ако е COM обект без стойност
            if "COMObject" in str(type(timestamp)):
                try:
                    timestamp = int(str(timestamp))
                except Exception as e:
                    if 'Member not found' in str(e):
                        return "Няма данни"
                    print(f"[DEBUG] Неуспешно COM конвертиране: {timestamp} ({e})")
                    return "Невалидна дата"

            # Преобразуване в int
            timestamp = int(timestamp)
            if timestamp <= 0:
                return "Няма дата"

            # Конвертиране от FILETIME
            dt = datetime.datetime(1601, 1, 1) + datetime.timedelta(microseconds=timestamp // 10)

            if dt.year < 1980 or dt.year > datetime.datetime.now().year + 5:
                print(f"[DEBUG] Timestamp извън диапазон: {original_timestamp} -> {dt}")
                return "Невалидна дата"

            return dt.strftime("%Y-%m-%d %H:%M:%S")

        except Exception as e:
            if 'Member not found' in str(e):
                return "Няма данни"
            print(f"[DEBUG] Грешка при конвертиране на timestamp: {original_timestamp} ({e})")
            return "Невалидна дата"



    def decode_user_account_control(self, uac_value):
        """Декодира userAccountControl флагове"""
        flags = {
            0x0001: "SCRIPT",
            0x0002: "Акаунт деактивиран",
            0x0008: "Изисква се home directory",
            0x0010: "Заключен",
            0x0020: "Парола не е задължителна",
            0x0040: "Парола не може да бъде сменена",
            0x0080: "Разрешени криптирани пароли",
            0x0100: "Временен дублиран акаунт",
            0x0200: "Нормален акаунт",
            0x0800: "Междомейнен акаунт",
            0x1000: "Workstation акаунт",
            0x2000: "Сървър акаунт",
            0x10000: "Паролата не изтича",
            0x20000: "Изисква се smartcard",
            0x40000: "Доверен за делегиране",
            0x80000: "Не се делегира",
            0x100000: "Използва само DES ключ",
            0x200000: "Не изисква предварителна автентикация",
            0x400000: "Паролата е изтекла",
            0x800000: "Доверен за автентикация при делегиране"
        }
        
        active_flags = []
        for flag, name in flags.items():
            if uac_value & flag:
                active_flags.append(name)
                
        return ', '.join(active_flags) if active_flags else "Нормален акаунт"

    def check_admin_privileges(self):
        """Проверява дали приложението се изпълнява с администраторски права"""
        try:
            return win32security.IsUserAnAdmin()
        except:
            return False

    def new_analysis(self):
        """Нов анализ - изчиства всички полета"""
        self.current_file = None
        self.file_path_label.config(text="Файл не е избран")
        self.file_text_results.delete(1.0, tk.END)
        self.file_hex_view.delete(1.0, tk.END)
        self.ip_entry.delete(0, tk.END)
        self.ip_text_results.delete(1.0, tk.END)
        self.url_entry.delete(0, tk.END)
        self.url_text_results.delete(1.0, tk.END)
        self.url_html_view.delete(1.0, tk.END)
        self.batch_ip_text.delete(1.0, tk.END)
        self.batch_ip_treeview.delete(*self.batch_ip_treeview.get_children())
        self.phishing_url_entry.delete(0, tk.END)
        self.phishing_ip_entry.delete(0, tk.END)
        self.phishing_template_preview.delete(1.0, tk.END)
        self.network_results.delete(1.0, tk.END)
        self.threat_intel_results.delete(1.0, tk.END)
        self.mxtoolbox_results.delete(1.0, tk.END)
        self.status_bar.config(text="Готов за нов анализ")




    def create_osint_tab(self):
        
        osint_tab = ttk.Frame(self.notebook)
        self.osint_tab = osint_tab
        self.notebook.add(osint_tab, text="OSINT")
        
        
        # Създаваме scrollable frame
        canvas = tk.Canvas(osint_tab)
        scrollbar = ttk.Scrollbar(osint_tab, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        canvas.configure(background="#2e2e2e", highlightthickness=0)
        scrollable_frame.configure(style="Dark.TFrame")
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
            )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.grid(row=0, column=0, sticky="nsew")
        scrollbar.grid(row=0, column=1, sticky="ns")
        
        osint_tab.grid_rowconfigure(0, weight=1)
        osint_tab.grid_columnconfigure(0, weight=1)
        
        # 1. Username
        ttk.Label(scrollable_frame, text="Потребителско име:").grid(row=0, column=0, sticky="e", padx=5, pady=2)
        self.username_entry = ttk.Entry(scrollable_frame, width=30)
        self.username_entry.grid(row=0, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="🔍 Социални мрежи", command=self.run_username_check).grid(row=0, column=2, padx=5, pady=2)

        # 2. Парола
        ttk.Label(scrollable_frame, text="Парола:").grid(row=1, column=0, sticky="e", padx=5, pady=2)
        self.password_entry = ttk.Entry(scrollable_frame, width=30, show="*")
        self.password_entry.grid(row=1, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="🔐 Проверка на парола", command=self.run_password_check).grid(row=1, column=2, padx=5, pady=2)

        # 3. Email HIBP
        ttk.Label(scrollable_frame, text="Имейл (Dehashed):").grid(row=2, column=0, sticky="e", padx=5, pady=2)
        self.hibp_email_entry = ttk.Entry(scrollable_frame, width=30)
        self.hibp_email_entry.grid(row=2, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="📧 Провери имейл", command=self.run_hibp_email_check).grid(row=2, column=2, padx=5, pady=2)

        # 4. CRT.sh
        ttk.Label(scrollable_frame, text="Домейн (CRT.sh):").grid(row=3, column=0, sticky="e", padx=5, pady=2)
        self.crtsh_entry = ttk.Entry(scrollable_frame, width=30)
        self.crtsh_entry.grid(row=3, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="📄 Търси CRT", command=self.run_crtsh_lookup).grid(row=3, column=2, padx=5, pady=2)

        # 5. TLD Checker
        ttk.Label(scrollable_frame, text="Домейн (TLD check):").grid(row=4, column=0, sticky="e", padx=5, pady=2)
        self.tld_domain_entry = ttk.Entry(scrollable_frame, width=30)
        self.tld_domain_entry.grid(row=4, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="🌐 Провери TLD", command=self.run_tld_check).grid(row=4, column=2, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="💾 Запази TLD", command=self.save_tld_results).grid(row=5, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="🛑 Докладвай IP-та", command=self.report_tld_to_abuse).grid(row=5, column=2, padx=5, pady=2)

        # 6. Pastebin
        ttk.Label(scrollable_frame, text="Търсене в Pastebin:").grid(row=6, column=0, sticky="e", padx=5, pady=2)
        self.pastebin_query_entry = ttk.Entry(scrollable_frame, width=30)
        self.pastebin_query_entry.grid(row=6, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="📄 Търси", command=self.run_pastebin_search).grid(row=6, column=2, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="🌐 Отвори линкове", command=self.open_pastebin_links).grid(row=7, column=2, padx=5, pady=2)

        # 7. Hunter.io
        ttk.Label(scrollable_frame, text="Домейн (Hunter):").grid(row=8, column=0, sticky="e", padx=5, pady=2)
        self.hunter_entry = ttk.Entry(scrollable_frame, width=30)
        self.hunter_entry.grid(row=8, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="🔎 Имейли", command=self.run_hunter_lookup).grid(row=8, column=2, padx=5, pady=2)

        # 8. GitHub
        ttk.Label(scrollable_frame, text="GitHub дума:").grid(row=9, column=0, sticky="e", padx=5, pady=2)
        self.github_entry = ttk.Entry(scrollable_frame, width=30)
        self.github_entry.grid(row=9, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="💻 GitHub", command=self.run_github_search).grid(row=9, column=2, padx=5, pady=2)

        # 9. ThreatCrowd
        ttk.Label(scrollable_frame, text="IP/домейн/хеш:").grid(row=10, column=0, sticky="e", padx=5, pady=2)
        self.threatcrowd_entry = ttk.Entry(scrollable_frame, width=30)
        self.threatcrowd_entry.grid(row=10, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="🧪 ThreatCrowd", command=self.run_threatcrowd_lookup).grid(row=10, column=2, padx=5, pady=2)

        # 10. ASN
        ttk.Label(scrollable_frame, text="ASN/IP:").grid(row=11, column=0, sticky="e", padx=5, pady=2)
        self.asn_entry = ttk.Entry(scrollable_frame, width=30)
        self.asn_entry.grid(row=11, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="🌍 ASN", command=self.run_asn_lookup).grid(row=11, column=2, padx=5, pady=2)

        # 11. Shodan
        ttk.Label(scrollable_frame, text="IP/домейн (Shodan):").grid(row=12, column=0, sticky="e", padx=5, pady=2)
        self.shodan_entry = ttk.Entry(scrollable_frame, width=30)
        self.shodan_entry.grid(row=12, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="🔎 Shodan", command=self.run_shodan_lookup).grid(row=12, column=2, padx=5, pady=2)

        # 12. Censys
        ttk.Label(scrollable_frame, text="IP/домейн (Censys):").grid(row=13, column=0, sticky="e", padx=5, pady=2)
        self.censys_entry = ttk.Entry(scrollable_frame, width=30)
        self.censys_entry.grid(row=13, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="🌐 Censys", command=self.run_censys_lookup).grid(row=13, column=2, padx=5, pady=2)

        # 13. Onyphe
        ttk.Label(scrollable_frame, text="Търсене (Onyphe):").grid(row=14, column=0, sticky="e", padx=5, pady=2)
        self.onyphe_entry = ttk.Entry(scrollable_frame, width=30)
        self.onyphe_entry.grid(row=14, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="🧠 Onyphe", command=self.run_onyphe_search).grid(row=14, column=2, padx=5, pady=2)

        # 14. PublicWWW
        ttk.Label(scrollable_frame, text="Ключова дума (PublicWWW):").grid(row=15, column=0, sticky="e", padx=5, pady=2)
        self.publicwww_entry = ttk.Entry(scrollable_frame, width=30)
        self.publicwww_entry.grid(row=15, column=1, padx=5, pady=2)
        ttk.Button(scrollable_frame, text="🕸️ PublicWWW", command=self.run_publicwww_search).grid(row=15, column=2, padx=5, pady=2)

        # Summary & Export
        ttk.Button(scrollable_frame, text="📊 Обобщен OSINT доклад", command=self.show_osint_summary).grid(row=16, column=1, pady=10)
        ttk.Button(scrollable_frame, text="💾 Запази доклад", command=self.save_osint_summary).grid(row=16, column=2, pady=10)

    def show_text_window(self, title, content):
        """Показва прозорец с текст и възможност за копиране"""
        window = tk.Toplevel(self.root)
        window.title(title)
        window.geometry("800x600")
        
        text_frame = ttk.Frame(window)
        text_frame.pack(fill="both", expand=True, padx=5, pady=5)
        
        text_widget = tk.Text(text_frame, wrap="word", font=("Consolas", 10))
        scrollbar = ttk.Scrollbar(text_frame, command=text_widget.yview)
        text_widget.configure(yscrollcommand=scrollbar.set)
        
        text_widget.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        text_widget.insert("1.0", content)
        text_widget.config(state="disabled")
        
        # Добавяме бутон за копиране
        button_frame = ttk.Frame(window)
        button_frame.pack(fill="x", padx=5, pady=5)
        
        ttk.Button(button_frame, text="Копирай в клипборда", 
                  command=lambda: self.root.clipboard_clear() or self.root.clipboard_append(content)).pack(side="left")
        ttk.Button(button_frame, text="Затвори", command=window.destroy).pack(side="right")

    def show_osint_summary(self):
        """Показва обобщение на всички OSINT резултати"""
        if not hasattr(self, 'osint_results'):
            self.osint_results = {}
            
        parts = []
        for key, label in {
            "username": "👤 Профили",
            "password": "🔑 Парола",
            "hibp": "📧 HIBP Email",
            "crtsh": "📄 CRT.sh",
            "tld": "🌐 TLD/WHOIS",
            "pastebin": "📄 Pastebin",
            "hunter": "📧 Hunter.io",
            "github": "💻 GitHub",
            "threatcrowd": "🧪 ThreatCrowd",
            "asn": "🌍 ASN",
            "shodan": "🔎 Shodan",
            "censys": "🌐 Censys",
            "onyphe": "🧠 Onyphe",
            "publicwww": "🕸️ PublicWWW"
        }.items():
            data = self.osint_results.get(key)
            if data:
                if isinstance(data, list):
                    parts.append(f"=== {label} ===\n" + "\n".join(data))
                elif isinstance(data, dict):
                    try:
                        parts.append(f"=== {label} ===\n" + "\n".join(f"{k}: {v}" for k, v in data.items()))
                    except:
                        parts.append(f"=== {label} ===\n{str(data)}")
                else:
                    parts.append(f"=== {label} ===\n{data}")
        
        summary = "\n\n".join(parts) if parts else "❗ Няма налични OSINT резултати."
        self.show_text_window("📊 OSINT Обобщение", summary)

    def save_osint_summary(self):
        """Запазва OSINT резултатите във файл"""
        from tkinter import filedialog, messagebox
        
        if not hasattr(self, 'osint_results') or not self.osint_results:
            messagebox.showinfo("Грешка", "Няма данни за запазване")
            return
            
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
            title="Запази OSINT доклад"
        )
        
        if not file_path:
            return
            
        try:
            with open(file_path, "w", encoding="utf-8") as f:
                for key, data in self.osint_results.items():
                    f.write(f"=== {key.upper()} ===\n")
                    if isinstance(data, (list, tuple)):
                        f.write("\n".join(str(item) for item in data) + "\n\n")
                    elif isinstance(data, dict):
                        f.write("\n".join(f"{k}: {v}" for k, v in data.items()) + "\n\n")
                    else:
                        f.write(str(data) + "\n\n")
            messagebox.showinfo("Успех", f"Докладът е запазен в:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Грешка", f"Грешка при запазване:\n{str(e)}")

    def open_pastebin_links(self):
        """Отваря линковете от Pastebin в браузър"""
        if not hasattr(self, 'last_pastebin_results') or not self.last_pastebin_results:
            messagebox.showinfo("Информация", "Няма налични Pastebin линкове.")
            return
            
        for url in self.last_pastebin_results:
            if isinstance(url, str) and url.startswith(('http://', 'https://')):
                webbrowser.open_new_tab(url)

    def save_tld_results(self):
        """Запазва резултатите от TLD проверката"""
        from tkinter import filedialog, messagebox
        
        if not hasattr(self, 'last_tld_results') or not self.last_tld_results:
            messagebox.showinfo("Информация", "Няма резултати за записване.")
            return
            
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("CSV files", "*.csv"), ("All files", "*.*")],
            title="Запази TLD резултати"
        )
        
        if not file_path:
            return
            
        try:
            if file_path.endswith('.csv'):
                import csv
                with open(file_path, 'w', newline='', encoding='utf-8') as csvfile:
                    fieldnames = ['domain', 'ip', 'registrar', 'country', 'emails']
                    writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                    writer.writeheader()
                    for item in self.last_tld_results:
                        writer.writerow(item)
            else:
                with open(file_path, "w", encoding="utf-8") as f:
                    for r in self.last_tld_results:
                        f.write(f"{r.get('domain', 'N/A')} ({r.get('ip', 'N/A')})\n")
                        f.write(f"  ⤷ {r.get('registrar', 'N/A')} / {r.get('country', 'N/A')}\n")
                        f.write(f"  ⤷ {r.get('emails', 'N/A')}\n\n")
                        
            messagebox.showinfo("Успех", f"Резултатите са запазени в:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Грешка", f"Грешка при запазване:\n{str(e)}")

    def report_tld_to_abuse(self):
        """Докладва IP адреси от TLD проверката"""
        if not hasattr(self, 'last_tld_results') or not self.last_tld_results:
            messagebox.showinfo("Информация", "Няма резултати за докладване.")
            return
            
        reported = []
        for r in self.last_tld_results:
            ip = r.get("ip")
            if ip and ip != "N/A":
                try:
                    res = self.security_tools.report_to_abuseipdb(
                        ip, 
                        categories=["6"], 
                        comment=f"Suspicious domain: {r.get('domain', '')}"
                    )
                    if res.get("success"):
                        reported.append(ip)
                except Exception as e:
                    messagebox.showerror("Грешка", f"Грешка при докладване на {ip}:\n{str(e)}")
                    continue
                    
        if reported:
            messagebox.showinfo("Успех", f"Успешно докладвани IP адреси:\n" + "\n".join(reported))
        else:
            messagebox.showinfo("Информация", "Няма успешно докладвани IP адреси.")

    def run_username_check(self):
        """Многонитова проверка на социални мрежи с прокси поддръжка"""
        import requests
        from concurrent.futures import ThreadPoolExecutor, as_completed

        if not hasattr(self, 'osint_results'):
            self.osint_results = {}

        username = self.username_entry.get().strip()
        self.log_activity(f"OSINT търсене на потребителско име: {username}")
        if not username:
            messagebox.showwarning("Грешка", "Моля, въведете потребителско име.")
            return

        sites = {
                "1337x": f"https://www.1337x.to/user/{username}/",
                "2Dimensions": f"https://2Dimensions.com/a/{username}",
                "7Cups": f"https://www.7cups.com/@{username}",
                "9GAG": f"https://www.9gag.com/u/{username}",
                "APClips": f"https://apclips.com/{username}",
                "About.me": f"https://about.me/{username}",
                "Academia.edu": f"https://independent.academia.edu/{username}",
                "AdmireMe.Vip": f"https://admireme.vip/{username}",
                "Airbit": f"https://airbit.com/{username}",
                "Airliners": f"https://www.airliners.net/user/{username}/profile/photos",
                "All Things Worn": f"https://www.allthingsworn.com/profile/{username}",
                "AllMyLinks": f"https://allmylinks.com/{username}",
                "AniWorld": f"https://aniworld.to/user/profil/{username}",
                "Anilist": f"https://anilist.co/user/{username}/",
                "Apple Developer": f"https://developer.apple.com/forums/profile/{username}",
                "Apple Discussions": f"https://discussions.apple.com/profile/{username}",
                "Archive of Our Own": f"https://archiveofourown.org/users/{username}",
                "Archive.org": f"https://archive.org/details/@{username}",
                "ArtStation": f"https://www.artstation.com/{username}",
                "Asciinema": f"https://asciinema.org/~{username}",
                "Ask Fedora": f"https://ask.fedoraproject.org/u/{username}",
                "Atcoder": f"https://atcoder.jp/users/{username}",
                "Audiojungle": f"https://audiojungle.net/user/{username}",
                "Autofrage": f"https://www.autofrage.net/nutzer/{username}",
                "Avizo": f"https://www.avizo.cz/{username}/",
                "BOOTH": f"https://{username}.booth.pm/",
                "Bandcamp": f"https://www.bandcamp.com/{username}",
                "Bazar.cz": f"https://www.bazar.cz/{username}/",
                "Behance": f"https://www.behance.net/{username}",
                "Bezuzyteczna": f"https://bezuzyteczna.pl/uzytkownicy/{username}",
                "BiggerPockets": f"https://www.biggerpockets.com/users/{username}",
                "BioHacking": f"https://forum.dangerousthings.com/u/{username}",
                "BitBucket": f"https://bitbucket.org/{username}/",
                "Bitwarden Forum": f"https://community.bitwarden.com/u/{username}/summary",
                "Blipfoto": f"https://www.blipfoto.com/{username}",
                "Blogger": f"https://{username}.blogspot.com",
                "Bluesky": f"https://bsky.app/profile/{username}.bsky.social",
                "BoardGameGeek": f"https://boardgamegeek.com/user/{username}",
                "BongaCams": f"https://pt.bongacams.com/profile/{username}",
                "Bookcrossing": f"https://www.bookcrossing.com/mybookshelf/{username}/",
                "BraveCommunity": f"https://community.brave.com/u/{username}/",
                "BugCrowd": f"https://bugcrowd.com/{username}",
                "BuyMeACoffee": f"https://buymeacoff.ee/{username}",
                "BuzzFeed": f"https://buzzfeed.com/{username}",
                "CGTrader": f"https://www.cgtrader.com/{username}",
                "CNET": f"https://www.cnet.com/profiles/{username}/",
                "CSSBattle": f"https://cssbattle.dev/player/{username}",
                "CTAN": f"https://ctan.org/author/{username}",
                "Caddy Community": f"https://caddy.community/u/{username}/summary",
                "Car Talk Community": f"https://community.cartalk.com/u/{username}/summary",
                "CashMe": f"https://www.cash.me/{username}",
                "Cashapp": f"https://cash.app/${username}",
                "Chess.com": f"https://www.chess.com/member/{username}",
                "ClickASnap": f"https://www.clickasnap.com/{username}",
                "Codecademy": f"https://www.codecademy.com/profiles/{username}",
                "Codeforces": f"https://codeforces.com/profile/{username}",
                "Codepen": f"https://codepen.io/{username}",
                "CodersRank": f"https://profile.codersrank.io/user/{username}",
                "Codingame": f"https://www.codingame.com/profile/{username}",
                "ColorLovers": f"https://www.colourlovers.com/lover/{username}",
                "Contently": f"https://{username}.contently.com",
                "Coroflot": f"https://www.coroflot.com/{username}",
                "Crevado": f"https://{username}.crevado.com",
                "Criminal IP": f"https://www.criminalip.io/en/profile/{username}",
                "Crunchyroll": f"https://www.crunchyroll.com/user/{username}",
                "DailyMotion": f"https://www.dailymotion.com/{username}",
                "Deezer": f"https://www.deezer.com/en/profile/{username}",
                "Designspiration": f"https://www.designspiration.com/{username}/",
                "DeviantArt": f"https://{username}.deviantart.com",
                "Docker Hub": f"https://hub.docker.com/u/{username}",
                "DonateKo-fi": f"https://ko-fi.com/{username}",
                "Dribbble": f"https://dribbble.com/{username}",
                "Ebay": f"https://www.ebay.com/usr/{username}",
                "Ello": f"https://ello.co/{username}",
                "Etsy": f"https://www.etsy.com/shop/{username}",
                "EyeEm": f"https://www.eyeem.com/u/{username}",
                "Facebook": f"https://www.facebook.com/{username}",
                "Fandom": f"https://community.fandom.com/wiki/User:{username}",
                "Fiverr": f"https://www.fiverr.com/{username}",
                "Flickr": f"https://www.flickr.com/people/{username}",
                "Flipboard": f"https://flipboard.com/@{username}",
                "Fotolog": f"https://www.fotolog.com/{username}",
                "Freelancer": f"https://www.freelancer.com/u/{username}",
                "Freepik": f"https://www.freepik.com/profile/{username}",
                "FundRazr": f"https://fundrazr.com/profiles/{username}",
                "Gab": f"https://gab.com/{username}",
                "GameSpot": f"https://www.gamespot.com/profile/{username}",
                "Giphy": f"https://giphy.com/{username}",
                "GitHub": f"https://github.com/{username}",
                "GitLab": f"https://gitlab.com/{username}",
                "Gitea": f"https://try.gitea.io/{username}",
                "GoodReads": f"https://www.goodreads.com/user/show/{username}",
                "Gravatar": f"https://en.gravatar.com/{username}",
                "GreasyFork": f"https://greasyfork.org/en/users/{username}",
                "Gumroad": f"https://{username}.gumroad.com",
                "Hackaday": f"https://hackaday.io/{username}",
                "HackTheBox": f"https://app.hackthebox.com/profile/{username}",
                "HackerEarth": f"https://www.hackerearth.com/@{username}",
                "HackerNews": f"https://news.ycombinator.com/user?id={username}",
                "HackerOne": f"https://hackerone.com/{username}",
                "Hashnode": f"https://hashnode.com/@{username}",
                "Heroku": f"https://dashboard.heroku.com/accounts/{username}",
                "Houzz": f"https://www.houzz.com/user/{username}",
                "HumbleBundle": f"https://www.humblebundle.com/user/{username}",
                "IMDB": f"https://www.imdb.com/user/{username}",
                "ImageShack": f"https://imageshack.com/user/{username}",
                "Imgur": f"https://imgur.com/user/{username}",
                "Instagram": f"https://www.instagram.com/{username}",
                "Instructables": f"https://www.instructables.com/member/{username}/",
                "Investing": f"https://www.investing.com/members/{username}",
                "Itch.io": f"https://{username}.itch.io",
                "JSFiddle": f"https://jsfiddle.net/user/{username}",
                "Kaggle": f"https://www.kaggle.com/{username}",
                "Keybase": f"https://keybase.io/{username}",
                "Kiva": f"https://www.kiva.org/lender/{username}",
                "Kongregate": f"https://www.kongregate.com/accounts/{username}",
                "Lichess": f"https://lichess.org/@/{username}",
                "LinkedIn": f"https://www.linkedin.com/in/{username}",
                "LiveJournal": f"https://{username}.livejournal.com",
                "Medium": f"https://medium.com/@{username}",
                "Metacritic": f"https://www.metacritic.com/user/{username}",
                "Minds": f"https://www.minds.com/{username}",
                "MyAnimeList": f"https://myanimelist.net/profile/{username}",
                "NameMC": f"https://namemc.com/profile/{username}",
                "NexusMods": f"https://www.nexusmods.com/users/{username}",
                "OpenSea": f"https://opensea.io/{username}",
                "Patreon": f"https://www.patreon.com/{username}",
                "PayPal.Me": f"https://www.paypal.me/{username}",
                "Periscope": f"https://www.pscp.tv/{username}",
                "Pinterest": f"https://www.pinterest.com/{username}",
                "Quora": f"https://www.quora.com/profile/{username}",
                "Redbubble": f"https://www.redbubble.com/people/{username}",
                "Reddit": f"https://www.reddit.com/user/{username}",
                "Repl.it": f"https://replit.com/@{username}",
                "Roblox": f"https://www.roblox.com/user.aspx?username={username}",
                "Scratch": f"https://scratch.mit.edu/users/{username}",
                "Signal": f"https://signal.art/{username}",
                "Slack": f"https://{username}.slack.com",
                "Snapchat": f"https://www.snapchat.com/add/{username}",
                "SoundCloud": f"https://soundcloud.com/{username}",
                "Spotify": f"https://open.spotify.com/user/{username}",
                "StackOverflow": f"https://stackoverflow.com/users/{username}",
                "Steam": f"https://steamcommunity.com/id/{username}",
                "Strava": f"https://www.strava.com/athletes/{username}",
                "Telegram": f"https://t.me/{username}",
                "TikTok": f"https://www.tiktok.com/@{username}",
                "Tinder": f"https://tinder.com/@{username}",
                "Trello": f"https://trello.com/{username}",
                "Twitch": f"https://www.twitch.tv/{username}",
                "Twitter": f"https://twitter.com/{username}",
                "Unsplash": f"https://unsplash.com/@{username}",
                "Vimeo": f"https://vimeo.com/{username}",
                "VSCO": f"https://vsco.co/{username}",
                "Wattpad": f"https://www.wattpad.com/user/{username}",
                "Wikipedia": f"https://www.wikipedia.org/wiki/User:{username}",
                "WordPress": f"https://{username}.wordpress.com",
                "YouTube": f"https://www.youtube.com/@{username}",
                "ZoomInfo": f"https://www.zoominfo.com/p/{username}",

            
            }

        headers = {"User-Agent": "Mozilla/5.0 (SOC-Tool)"}
        proxy = self.config.get_proxy_settings()
        results = []

        def check_site(name, url):
            try:
                r = requests.head(url, headers=headers, proxies=proxy, timeout=5, allow_redirects=True)
                if r.status_code == 200:
                    return f"✅ {name}: {url}"
                elif r.status_code == 429:
                    return f"⚠️ {name}: Rate limited"
                else:
                    return f"❌ {name}: Не е намерен ({r.status_code})"
            except requests.RequestException as e:
                return f"⚠️ {name}: Грешка ({str(e)})"

        with ThreadPoolExecutor(max_workers=40) as executor:
            future_to_site = {
                executor.submit(check_site, name, url): name for name, url in sites.items()
            }
            for future in as_completed(future_to_site):
                results.append(future.result())

        self.osint_results["username"] = results
        self.show_text_window("🔍 Реални социални профили", "\n".join(sorted(results)))


    def run_password_check(self):
        """Проверява дали паролата е компрометирана"""
        if not hasattr(self, 'osint_results'):
            self.osint_results = {}
            
        password = self.password_entry.get().strip()
        if not password:
            messagebox.showwarning("Грешка", "Моля, въведете парола.")
            return
            
        try:
            result = self.check_password_pwned(password)
            if result.get("error"):
                messagebox.showerror("Грешка", result["error"])
            elif result.get("pwned"):
                self.osint_results["password"] = f"⚠️ Паролата е открита в {result['count']} инцидента с изтичане на данни!"
                messagebox.showwarning("Предупреждение", self.osint_results["password"])
            else:
                self.osint_results["password"] = "✅ Паролата не е намерена в публични инцидента с изтичане на данни."
                messagebox.showinfo("Информация", self.osint_results["password"])
        except Exception as e:
            messagebox.showerror("Грешка", f"Възникна грешка: {str(e)}")


    def check_email_dehashed(email, username, api_key):
        url = f"https://api.dehashed.com/search?query={email}"
        print(f"[DEBUG] Изпраща се заявка към: {url}")
        response = requests.get(url, auth=HTTPBasicAuth(username, api_key), timeout=10)
        print(f"[DEBUG] Статус код: {response.status_code}")
        print(f"[DEBUG] Тяло на отговора:\n{response.text}")
        try:
            response = requests.get(url, auth=HTTPBasicAuth(username, api_key), timeout=10)
            if response.status_code == 200:
                data = response.json()
                results = []
                for row in data.get("entries", []):
                    result = f"🔎 Източник: {row.get('source', 'N/A')}\n"
                    result += f"👤 Потребител: {row.get('username', 'N/A')}\n"
                    result += f"📧 Имейл: {row.get('email', 'N/A')}\n"
                    result += f"🔑 Парола: {row.get('password', 'N/A')}\n"
                    results.append(result)
                return {"found": True, "results": results}
            elif response.status_code == 404:
                return {"found": False}
            else:
                return {"error": f"DeHashed: код {response.status_code}"}
        except Exception as e:
            return {"error": str(e)}

    def run_crtsh_lookup(self):
        """Търси поддомейни в CRT.sh"""
        if not hasattr(self, 'osint_results'):
            self.osint_results = {}
            
        domain = self.crtsh_entry.get().strip()
        if not domain:
            messagebox.showwarning("Грешка", "Моля, въведете домейн.")
            return
            
        try:
            result = self.get_crtsh_domains(domain)
            if result.get("error"):
                messagebox.showerror("Грешка", result["error"])
            else:
                self.osint_results["crtsh"] = result["results"]
                self.show_text_window("CRT.sh Резултати", "\n".join(result["results"]))
        except Exception as e:
            messagebox.showerror("Грешка", f"Възникна грешка: {str(e)}")

    def run_tld_check(self):
        """Проверява TLD за даден домейн"""
        if not hasattr(self, 'osint_results'):
            self.osint_results = {}
            
        domain = self.tld_domain_entry.get().strip()
        if not domain:
            messagebox.showwarning("Грешка", "Моля, въведете домейн.")
            return
            
        try:
            results = self.tld_check(domain)
            self.last_tld_results = results
            self.osint_results["tld"] = results
            
            # Форматиране на резултатите за показване
            formatted_results = []
            for r in results:
                formatted_results.append(
                    f"Домейн: {r.get('domain', 'N/A')}\n"
                    f"IP: {r.get('ip', 'N/A')}\n"
                    f"Регистратор: {r.get('registrar', 'N/A')}\n"
                    f"Държава: {r.get('country', 'N/A')}\n"
                    f"Имейли: {r.get('emails', 'N/A')}\n"
                )
            
            self.show_text_window("TLD Резултати", "\n".join(formatted_results))
        except Exception as e:
            messagebox.showerror("Грешка", f"Възникна грешка: {str(e)}")

    def run_pastebin_search(self):
        """Търси в Pastebin"""
        if not hasattr(self, 'osint_results'):
            self.osint_results = {}
            
        query = self.pastebin_query_entry.get().strip()
        if not query:
            messagebox.showwarning("Грешка", "Моля, въведете дума за търсене.")
            return
            
        try:
            use_api = self.config.config.getboolean("OSINT_API_KEYS", "USE_INTELLIGENCE_X", fallback=True)
            api_key = self.config.config.get("OSINT_API_KEYS", "INTELLIGENCE_X", fallback="")
            proxies = self.config.get_proxy_settings()

            result = self.smart_pastebin_lookup(query, intelx_api_key=api_key, proxies=proxies, use_intelx=use_api)
            if result.get("error"):
                messagebox.showerror("Грешка", result["error"])
            else:
                self.last_pastebin_results = result["results"]
                self.osint_results["pastebin"] = result["results"]
                self.show_text_window("Pastebin Резултати", "\n".join(result["results"]))
        except Exception as e:
            messagebox.showerror("Грешка", f"Възникна грешка: {str(e)}")

    def run_hunter_lookup(self):
        """Търси имейли в Hunter.io"""
        if not hasattr(self, 'osint_results'):
            self.osint_results = {}
            
        domain = self.hunter_entry.get().strip()
        self.log_activity(f"OSINT Hunter.io проверка за домейн: {domain}")
        if not domain:
            messagebox.showwarning("Грешка", "Моля, въведете домейн.")
            return
            
        try:
            key = self.config.config.get("OSINT_API_KEYS", "HUNTER_IO", fallback="")
            if not key:
                messagebox.showerror("Грешка", "Няма API ключ за Hunter.io в config.ini")
                return
            
            result = self.search_hunter_io(domain, key)
            if result.get("error"):
                messagebox.showerror("Грешка", result["error"])
            else:
                self.osint_results["hunter"] = result["results"]
                self.show_text_window("Hunter.io Резултати", "\n".join(result["results"]))
        except Exception as e:
            messagebox.showerror("Грешка", f"Възникна грешка: {str(e)}")

    def run_github_search(self):
        """Търси в GitHub код"""
        if not hasattr(self, 'osint_results'):
            self.osint_results = {}
            
        query = self.github_entry.get().strip()
        if not query:
            messagebox.showwarning("Грешка", "Моля, въведете дума за търсене.")
            return
            
        try:
            token = self.config.config.get("OSINT_API_KEYS", "GITHUB_TOKEN", fallback="")
            proxies = self.config.get_proxy_settings()
            result = self.search_github_code(query, token, proxies=proxies)
            if result.get("error"):
                messagebox.showerror("Грешка", result["error"])
            else:
                self.osint_results["github"] = result["results"]
                self.show_text_window("GitHub Резултати", "\n".join(result["results"]))
        except Exception as e:
            messagebox.showerror("Грешка", f"Възникна грешка: {str(e)}")

    def run_threatcrowd_lookup(self):
        """Търси в ThreatCrowd"""
        if not hasattr(self, 'osint_results'):
            self.osint_results = {}
            
        q = self.threatcrowd_entry.get().strip()
        if not q:
            messagebox.showwarning("Грешка", "Моля, въведете IP/домейн/хеш.")
            return
            
        try:
            result = self.lookup_threatcrowd(q)
            if result.get("error"):
                messagebox.showerror("Грешка", result["error"])
            elif not result.get("found"):
                messagebox.showinfo("Информация", "Няма намерени резултати.")
            else:
                data = result["data"]
                parts = []
                if "resolutions" in data:
                    parts.append("=== Хостове ===\n" + "\n".join([f"{r['domain']} → {r['last_resolved']}" for r in data["resolutions"]]))
                if "emails" in data:
                    parts.append("=== Имейли ===\n" + "\n".join(data["emails"]))
                if "hashes" in data:
                    parts.append("=== Хешове ===\n" + "\n".join(data["hashes"]))
                
                self.osint_results["threatcrowd"] = parts
                self.show_text_window("ThreatCrowd Резултати", "\n\n".join(parts))
        except Exception as e:
            messagebox.showerror("Грешка", f"Възникна грешка: {str(e)}")

    def run_asn_lookup(self):
        """Търси ASN информация"""
        if not hasattr(self, 'osint_results'):
            self.osint_results = {}
            
        target = self.asn_entry.get().strip()
        if not target:
            messagebox.showwarning("Грешка", "Моля, въведете IP или ASN.")
            return
            
        try:
            result = self.lookup_asn_ip_info(target)
            if result.get("error"):
                messagebox.showerror("Грешка", result["error"])
            else:
                d = result["data"]
                out = [
                    f"Организация: {d.get('name', 'N/A')}",
                    f"Държава: {d.get('country_code', 'N/A')}",
                    f"ASN: {d.get('asn', 'N/A')}"
                ]
                if "prefixes" in d:
                    out.append("=== IP диапазони ===\n" + "\n".join(p["prefix"] for p in d["prefixes"][:10]))
                
                self.osint_results["asn"] = out
                self.show_text_window("ASN Резултати", "\n".join(out))
        except Exception as e:
            messagebox.showerror("Грешка", f"Възникна грешка: {str(e)}")

    def run_shodan_lookup(self):
        """Търси в Shodan"""
        proxies = self.config.get_proxy_settings()
        if not hasattr(self, 'osint_results'):
            self.osint_results = {}
            
        target = self.shodan_entry.get().strip()
        self.log_activity(f"OSINT Shodan търсене за IP: {ip}")
        if not target:
            messagebox.showwarning("Грешка", "Моля, въведете IP или домейн.")
            return
            
        try:
            key = self.config.get_api_key("SHODAN_API_KEY", fallback="")
            
            if not key:
                messagebox.showerror("Грешка", "Няма API ключ за Shodan в config.ini")
                return
                
            result = self.search_shodan_host(target, key)
            if result.get("error"):
                messagebox.showerror("Грешка", result["error"])
            else:
                self.osint_results["shodan"] = result["results"]
                self.show_text_window("Shodan Резултати", "\n".join(result["results"]))
        except Exception as e:
            messagebox.showerror("Грешка", f"Възникна грешка: {str(e)}")

    def run_hibp_email_check(self):
        email = self.hibp_email_entry.get().strip()
        self.log_activity(f"OSINT проверка за изтекъл имейл: {email}")
        if not email:
            messagebox.showwarning("Имейл", "Моля, въведи имейл адрес.")
            return

        api_key = self.config.get_osint_key("BREACHDIRECTORY_KEY", fallback="")
        if not api_key:
            messagebox.showerror("BreachDirectory", "Липсва API ключ за BreachDirectory!")
            return
        print(f"[DEBUG] Вика се self.check_email_breachdirectory от клас SOCGUI")
        result = self.check_email_breachdirectory(email, api_key, proxies=self.config.get_proxy_settings())  # self. !!!

        if result.get("error"):
            messagebox.showerror("Грешка", result["error"])
        elif result.get("found", 0) > 0:
            formatted = []
            for entry in result["results"]:
                formatted.append(
                    f"📧 Имейл: {entry.get('email')}\n"
                    f"🔐 Парола: {entry.get('password')}\n"
                    f"🔑 SHA1: {entry.get('sha1')}\n"
                    f"🧬 Hash: {entry.get('hash')}\n"
                    f"🗂️ Източници: {', '.join(entry.get('sources', []))}\n"
                    f"{'-'*50}"
                )
            self.osint_results["hibp"] = formatted
            self.show_text_window("🔐 Dehashed (BreachDirectory)", "\n\n".join(formatted))
        else:
            self.osint_results["hibp"] = "✅ Имейлът не е намерен в инцидента с изтичане на данни."
            messagebox.showinfo("Dehashed", "✅ Имейлът не е намерен в инцидента с изтичане на данни (BreachDirectory)")


    def check_email_breachdirectory(self, email, api_key):
        """Извлича информация за инцидента с изтичане на данни чрез BreachDirectory"""
        url = f"https://breachdirectory.p.rapidapi.com/"
        headers = {
            "X-RapidAPI-Key": api_key,
            "X-RapidAPI-Host": "breachdirectory.p.rapidapi.com"
        }
        params = {"func": "auto", "term": email}

        try:
            proxies = self.config.get_proxy_settings()
            response = requests.get(url, headers=headers, params=params, timeout=10, proxies=proxies)

            if response.status_code == 200:
                data = response.json()
                if not data.get("success", True):
                    return {"error": data.get("message", "Неуспешна заявка")}
                findings = data.get("result", [])
                return {
                    "found": len(findings),
                    "results": findings
                }
            else:
                return {"error": f"Грешка {response.status_code}: {response.text}"}
        except Exception as e:
            return {"error": str(e)}


    def run_censys_lookup(self):
        """Търси в Censys"""
        if not hasattr(self, 'osint_results'):
            self.osint_results = {}
            
        target = self.censys_entry.get().strip()
        self.log_activity(f"OSINT Censys търсене: {term}")
        if not target:
            messagebox.showwarning("Грешка", "Моля, въведете IP или домейн.")
            return
            
        try:
            cid = self.config.config.get("OSINT_API_KEYS", "CENSYS_ID", fallback="")
            csec = self.config.config.get("OSINT_API_KEYS", "CENSYS_SECRET", fallback="")
            
            if not cid or not csec:
                messagebox.showerror("Грешка", "Няма API ключове за Censys в config.ini")
                return
                
            result = self.search_censys_host(target, cid, csec, proxies = self.config.get_proxy_settings())
            if result.get("error"):
                messagebox.showerror("Грешка", result["error"])
            else:
                self.osint_results["censys"] = result["results"]
                self.show_text_window("Censys Резултати", "\n".join(result["results"]))
        except Exception as e:
            messagebox.showerror("Грешка", f"Възникна грешка: {str(e)}")

    def run_onyphe_search(self):
        """Търси в Onyphe"""
        if not hasattr(self, 'osint_results'):
            self.osint_results = {}
            
        query = self.onyphe_entry.get().strip()
        self.log_activity(f"Onyphe заявка: {term}")
        if not query:
            messagebox.showwarning("Грешка", "Моля, въведете дума или IP за търсене.")
            return
            
        try:
            key = self.config.config.get("OSINT_API_KEYS", "ONYPHE_KEY", fallback="")
            if not key:
                messagebox.showerror("Грешка", "Няма API ключ за Onyphe в config.ini")
                return
                
            result = self.search_onyphe(query, key, proxies=self.config.get_proxy_settings())
            if result.get("error"):
                messagebox.showerror("Грешка", result["error"])
            else:
                self.osint_results["onyphe"] = result["results"]
                self.show_text_window("Onyphe Резултати", "\n".join(result["results"]))
        except Exception as e:
            messagebox.showerror("Грешка", f"Възникна грешка: {str(e)}")

    def run_publicwww_search(self):
        """Търси в PublicWWW"""
        if not hasattr(self, 'osint_results'):
            self.osint_results = {}
            
        q = self.publicwww_entry.get().strip()
        if not q:
            messagebox.showwarning("Грешка", "Моля, въведете ключова дума за търсене.")
            return
            
        try:
            result = self.search_publicwww(q)
            if result.get("error"):
                messagebox.showerror("Грешка", result["error"])
            else:
                self.osint_results["publicwww"] = result["results"]
                self.show_text_window("PublicWWW Резултати", "\n".join(result["results"]))
        except Exception as e:
            messagebox.showerror("Грешка", f"Възникна грешка: {str(e)}")

    # API методи

    #@staticmethod
    def check_password_pwned(self, password):
        """Проверява парола в HIBP"""
        try:
            proxies = self.config.get_proxy_settings()
            sha1 = hashlib.sha1(password.encode("utf-8")).hexdigest().upper()
            prefix, suffix = sha1[:5], sha1[5:]
            
            headers = {"User-Agent": "SOC-Tool-v1"}
            r = requests.get(
                f"https://api.pwnedpasswords.com/range/{prefix}",
                headers=headers,
                proxies=proxies,
                timeout=10
            )
            r.raise_for_status()
            
            for line in r.text.splitlines():
                if line.startswith(suffix):
                    return {"pwned": True, "count": int(line.split(":")[1])}
            return {"pwned": False}
        except Exception as e:
            return {"error": f"Грешка при проверка на парола: {str(e)}"}

    @staticmethod
    def check_email_dehashed(email, username, api_key):
        url = f"https://api.dehashed.com/search?query={email}"
        try:
            response = requests.get(url, auth=HTTPBasicAuth(username, api_key), timeout=10)
            if response.status_code == 200:
                data = response.json()
                results = []
                for row in data.get("entries", []):
                    result = f"🔎 Източник: {row.get('source', 'N/A')}\n"
                    result += f"👤 Потребител: {row.get('username', 'N/A')}\n"
                    result += f"📧 Имейл: {row.get('email', 'N/A')}\n"
                    result += f"🔑 Парола: {row.get('password', 'N/A')}\n"
                    results.append(result)
                return {"found": True, "results": results}
            elif response.status_code == 404:
                return {"found": False}
            else:
                return {"error": f"DeHashed: код {response.status_code}"}
        except Exception as e:
            return {"error": str(e)}

    #@staticmethod
    def get_crtsh_domains(self, domain):
        """Търси поддомейни в CRT.sh"""
        try:
            headers = {"User-Agent": "SOC-Tool-v1"}
            r = requests.get(
                f"https://crt.sh/?q={domain}&output=json",
                headers=headers,
                proxies=self.config.get_proxy_settings(),
                timeout=15
            )
            r.raise_for_status()
            
            data = r.json()
            domains = set()
            for entry in data:
                for d in entry.get("name_value", "").split("\n"):
                    domains.add(d.strip().lower())
            return {"results": sorted(domains)}
        except Exception as e:
            return {"error": f"Грешка при търсене в CRT.sh: {str(e)}"}

    #@staticmethod
    def lookup_threatcrowd(self, query):
        """Търси в ThreatCrowd"""
        try:
            if "." in query and not any(x in query for x in [":", "/", "@"]):
                qtype, key = "domain", "domain"
            elif len(query) in [32, 40, 64]:
                qtype, key = "hash", "resource"
            else:
                qtype, key = "ip", "ip"
                
            headers = {"User-Agent": "SOC-Tool-v1"}
            url = f"https://www.threatcrowd.org/searchApi/v2/{qtype}/report/?{key}={query}"
            
            r = requests.get(url, headers=headers, proxies=self.config.get_proxy_settings(), timeout=15)
            r.raise_for_status()
            
            d = r.json()
            if d.get("response_code") == "0":
                return {"found": False}
            return {"found": True, "data": d}
        except Exception as e:
            return {"error": f"Грешка при търсене в ThreatCrowd: {str(e)}"}

    #@staticmethod
    def smart_pastebin_lookup(self, query, intelx_api_key="", proxies=None, use_intelx=True):
        """Търси в Pastebin с опция за Intelligence X"""
        try:
            proxies = self.config.get_proxy_settings()
            if use_intelx and intelx_api_key:
                # Използване на Intelligence X API
                headers = {
                    "User-Agent": "SOC-Tool-v1",
                    "Authorization": f"Bearer {intelx_api_key}"
                }
                data = {
                    "term": query,
                    "buckets": ["pastebin"],
                    "lookuplevel": 0,
                    "maxresults": 100,
                    "timeout": 5,
                    "datefrom": "",
                    "dateto": "",
                    "sort": 4,
                    "media": 0,
                    "terminate": []
                }
                
                r = requests.post(
                    "https://2.intelx.io/intelligent/search",
                    json=data,
                    headers=headers,
                    proxies = proxies,
                    timeout=15
                )
                r.raise_for_status()
                
                results = []
                for item in r.json().get("records", []):
                    if "pastebin.com" in item.get("name", ""):
                        results.append(f"https://pastebin.com/{item.get('name')}")
                return {"results": results[:10]}  # Ограничаваме до 10 резултата
            else:
                # Google dork за Pastebin като fallback
                headers = {"User-Agent": "Mozilla/5.0"}
                url = f"https://www.google.com/search?q=site:pastebin.com {quote_plus(query)}"
                
                r = requests.get(url, headers=headers, proxies=proxies, timeout=15)
                r.raise_for_status()
                
                links = {
                    line for line in r.text.split('"') 
                    if "pastebin.com/" in line 
                    and "/u/" not in line
                    and line.startswith("http")
                }
                return {"results": list(links)[:10]}
        except Exception as e:
            return {"error": f"Грешка при търсене в Pastebin: {str(e)}"}

    #@staticmethod
    def search_hunter_io(self, domain, api_key):
        """Търси имейли в Hunter.io"""
        try:
            headers = {
                "User-Agent": "SOC-Tool-v1",
                "Authorization": f"Bearer {api_key}"
            }
            
            r = requests.get(
                f"https://api.hunter.io/v2/domain-search?domain={domain}",
                headers=headers,
                proxies = self.config.get_proxy_settings(),
                timeout=15
            )
            r.raise_for_status()
            
            emails = r.json().get("data", {}).get("emails", [])
            return {
                "results": [
                    f"{e.get('value')} – {e.get('first_name', '')} {e.get('last_name', '')} ({e.get('position', '')})" 
                    for e in emails
                ]
            }
        except Exception as e:
            return {"error": f"Грешка при търсене в Hunter.io: {str(e)}"}

    #@staticmethod
    def search_github_code(self, query, token=None, proxies=None):
        """Търси код в GitHub"""
        proxies = self.config.get_proxy_settings()
        try:
            headers = {
                "Accept": "application/vnd.github.v3.text-match+json",
                "User-Agent": "SOC-Tool-v1"
            }
            if token:
                headers["Authorization"] = f"token {token}"
            
            r = requests.get(
                f"https://api.github.com/search/code?q={quote_plus(query)}+in:file",
                headers=headers,
                proxies=proxies,
                timeout=15,
            )
            
            if r.status_code == 403:
                return {"error": "GitHub API лимит достигнат. Опитайте по-късно или използвайте токен."}
            r.raise_for_status()
            
            items = r.json().get("items", [])[:10]  # Ограничаваме до 10 резултата
            return {
                "results": [
                    f"{i['name']} – {i['repository']['full_name']}:\n🔗 {i['html_url']}" 
                    for i in items
                ]
            }
        except Exception as e:
            return {"error": f"Грешка при търсене в GitHub: {str(e)}"}

    #@staticmethod
    def lookup_asn_ip_info(self, target):
        """Търси ASN информация"""
        try:
            headers = {"User-Agent": "SOC-Tool-v1"}
            
            if target.lower().startswith("as"):
                target = target.lower().replace("as", "")
                url = f"https://api.bgpview.io/asn/{target}"
            else:
                url = f"https://api.bgpview.io/ip/{target}"
                
            r = requests.get(url, headers=headers, proxies = self.config.get_proxy_settings(), timeout=15)
            r.raise_for_status()
            
            return {"data": r.json().get("data", {})}
        except Exception as e:
            return {"error": f"Грешка при търсене на ASN информация: {str(e)}"}

    #@staticmethod
    def search_shodan_host(self, target, api_key):
        """Търси в Shodan"""
        try:
            headers = {
                "User-Agent": "SOC-Tool-v1",
                "Authorization": f"Bearer {api_key}"
            }
            
            r = requests.get(
                f"https://api.shodan.io/shodan/host/{target}",
                headers=headers,
                proxies = self.config.get_proxy_settings(),
                timeout=15
            )
            
            if r.status_code == 404:
                return {"not_found": True}
            r.raise_for_status()
            
            d = r.json()
            out = [
                f"🌐 IP: {d.get('ip_str')}",
                f"🌍 Location: {d.get('country_name')} / {d.get('city')}",
                f"🏢 Организация: {d.get('org')}",
                f"🖥️ Операционна система: {d.get('os', 'N/A')}"
            ]
            out += [
                f"🛠️ Port {s.get('port')}: {s.get('data', '').splitlines()[0][:100]}" 
                for s in d.get("data", [])[:5]
            ]
            return {"results": out}
        except Exception as e:
            return {"error": f"Грешка при търсене в Shodan: {str(e)}"}

    #@staticmethod
    def search_censys_host(self, target, api_id, api_secret):
        """Търси в Censys"""
        try:
            headers = {
                "User-Agent": "SOC-Tool-v1",
                "Authorization": "Basic " + base64.b64encode(f"{api_id}:{api_secret}".encode()).decode(),
                "Content-Type": "application/json"
            }
            
            r = requests.post(
                "https://search.censys.io/api/v2/hosts/search",
                headers=headers,
                proxies = self.config.get_proxy_settings(),
                json={"q": target, "per_page": 5},
                timeout=15
            )
            r.raise_for_status()
            
            hits = r.json().get("result", {}).get("hits", [])
            results = []
            for h in hits:
                results.append(
                    f"🌐 IP: {h.get('ip')} ({h.get('location', {}).get('country', 'N/A')})"
                )
                for s in h.get("services", [])[:3]:
                    results.append(
                        f"   🔹 {s.get('transport_protocol', '').upper()}:{s.get('port')} - "
                        f"{s.get('service_name')} ({s.get('software', {}).get('version', '')})"
                    )
            return {"results": results}
        except Exception as e:
            return {"error": f"Грешка при търсене в Censys: {str(e)}"}

    #@staticmethod
    def search_onyphe(self, query, api_key):
        """Търси в Onyphe"""
        try:
            headers = {
                "User-Agent": "SOC-Tool-v1",
                "Authorization": f"apikey {api_key}"
            }
            
            r = requests.get(
                f"https://www.onyphe.io/api/v2/simple/search?q={quote_plus(query)}",
                headers=headers,
                proxies = self.config.get_proxy_settings(),
                timeout=15
            )
            r.raise_for_status()
            
            hits = r.json().get("results", [])[:10]
            return {
                "results": [
                    f"{h.get('summary')} – {h.get('source')}\n🔗 {h.get('link')}" 
                    for h in hits
                ]
            }
        except Exception as e:
            return {"error": f"Грешка при търсене в Onyphe: {str(e)}"}

    #@staticmethod
    def search_publicwww(query):
        """Търси в PublicWWW"""
        try:
            headers = {"User-Agent": "Mozilla/5.0"}
            r = requests.get(
                f"https://publicwww.com/websites/{quote_plus(query)}/",
                headers=headers,
                proxies = self.config.get_proxy_settings(),
                timeout=15
            )
            r.raise_for_status()
            
            soup = BeautifulSoup(r.text, "html.parser")
            return {
                "results": [
                    a["href"].replace("/site/", "").strip("/") 
                    for a in soup.find_all("a", href=True) 
                    if "/site/" in a["href"]
                ][:10]
            }
        except Exception as e:
            return {"error": f"Грешка при търсене в PublicWWW: {str(e)}"}
            
    def create_email_analysis_tab(self):
        
        tab = ttk.Frame(self.notebook)
        self.email_analysis_tab = tab
        self.notebook.add(tab, text="Email Threat Analysis")
        self.email_tab = tab

        self.email_file_path = tk.StringVar()
        self.auto_analyze_var = tk.BooleanVar(value=True)
        self.include_email_context = tk.BooleanVar(value=True)

        # --- Горна част: Email текст ---
        self.email_textbox = scrolledtext.ScrolledText(tab, width=100, height=15, wrap=tk.WORD)
        self.email_textbox.pack(padx=5, pady=5, fill=tk.BOTH, expand=True)

        # --- Ред 1: Зареждане + Auto ---
        row1 = ttk.Frame(tab)
        row1.pack(padx=5, pady=2, fill=tk.X)
        ttk.Button(row1, text="Load Email (.eml/.msg)", command=self.browse_email_file).pack(side=tk.LEFT, padx=2)
        ttk.Checkbutton(row1, text="Auto-analyze", variable=self.auto_analyze_var).pack(side=tk.LEFT, padx=5)

        # --- Ред 2: OpenAI + VT + AbuseIPDB ---
        row2 = ttk.Frame(tab)
        row2.pack(padx=5, pady=2, fill=tk.X)
        ttk.Button(row2, text="Analyze with OpenAI", command=self.analyze_email_with_openai).pack(side=tk.LEFT, padx=2)
        ttk.Button(row2, text="Free Analyze (GPT)", command=self.analyze_email_with_free_gpt).pack(side=tk.LEFT, padx=2)
        ttk.Button(row2, text="Scan Links (VT)", command=self.scan_email_links).pack(side=tk.LEFT, padx=2)
        ttk.Button(row2, text="Scan IPs (AbuseIPDB)", command=self.scan_email_ips).pack(side=tk.LEFT, padx=2)
        ttk.Button(row2, text="Scan Attachments", command=self.scan_email_attachments).pack(side=tk.LEFT, padx=2)

        # --- Ред 3: Запазване/Копиране ---
        row3 = ttk.Frame(tab)
        row3.pack(padx=5, pady=2, fill=tk.X)
        ttk.Button(row3, text="Copy Summary", command=self.copy_threat_summary).pack(side=tk.LEFT, padx=2)
        ttk.Button(row3, text="Save Summary", command=self.save_threat_summary).pack(side=tk.LEFT, padx=2)

        # --- Custom GPT Prompt ---
        ttk.Label(tab, text="Custom GPT Query:").pack(pady=(10, 2))
        self.gpt_custom_input = scrolledtext.ScrolledText(tab, height=3, width=100, wrap=tk.WORD)
        self.gpt_custom_input.pack(padx=5, pady=2, fill=tk.X)

        row4 = ttk.Frame(tab)
        row4.pack(padx=5, pady=2, fill=tk.X)
        ttk.Checkbutton(row4, text="Include email in context", variable=self.include_email_context).pack(side=tk.LEFT, padx=5)
        ttk.Button(row4, text="Ask GPT", command=self.ask_gpt_custom_prompt).pack(side=tk.LEFT, padx=5)
        ttk.Button(row4, text="Free Ask GPT", command=self.ask_gpt_custom_prompt_free).pack(side=tk.LEFT, padx=5)

        # --- GPT Output Box ---
        self.gpt_custom_output = scrolledtext.ScrolledText(tab, height=6, width=100, wrap=tk.WORD, bg="#1e1e1e", fg="cyan")
        self.gpt_custom_output.pack(padx=5, pady=5, fill=tk.BOTH, expand=True)

        # Модел за OpenAI (официален)
        self.selected_gpt_model = tk.StringVar(value="gpt-4")
        model_row = ttk.Frame(tab)
        model_row.pack(padx=5, pady=2, fill=tk.X)
        ttk.Label(model_row, text="OpenAI Model:").pack(side=tk.LEFT)
        ttk.Combobox(model_row, textvariable=self.selected_gpt_model, values=["gpt-3.5-turbo", "gpt-4", "gpt-4o"], width=15).pack(side=tk.LEFT, padx=5)


        # --- GPT History + Save ---
        ttk.Label(tab, text="Prompt History:").pack(pady=(5, 2))
        self.gpt_prompt_history = tk.Listbox(tab, height=5, width=100)
        self.gpt_prompt_history.pack(padx=5, pady=2, fill=tk.X)
        self.gpt_prompt_history.bind("<<ListboxSelect>>", self.load_gpt_history_entry)

        ttk.Button(tab, text="Save GPT Answer", command=self.save_gpt_response).pack(padx=5, pady=2)

        # --- Email Result Output ---
        ttk.Label(tab, text="Full Threat Result:").pack(pady=(5, 2))
        self.email_result_text = scrolledtext.ScrolledText(tab, width=100, height=6, wrap=tk.WORD, bg="#1e1e1e", fg="lightgreen")
        self.email_result_text.pack(padx=5, pady=5, fill=tk.BOTH, expand=True)

        # --- Вътрешни данни ---
        self.loaded_attachments = []
        self.last_openai_result = ""
        self.last_urls = []
        self.last_ips = []
        self.last_attachment_results = []
        self.gpt_qa_history = []

    def get_proxy(self):
        try:
            return {
                "http": self.config.get("Proxy", "http", fallback=None),
                "https": self.config.get("Proxy", "https", fallback=None)
            }
        except:
            return None


    def browse_email_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Email Files", "*.eml *.msg")])
        self.log_activity(f"Зареден email файл")
        if file_path:
            self.email_file_path.set(file_path)
            self.load_email_file(file_path)

    def load_email_file(self, path):
        content = ""
        self.loaded_attachments = []
        self.last_attachment_results = []

        try:
            self.email_result_text.insert(tk.END, f"[Loading] Parsing file: {path}\n")
            self.root.update()

            if path.endswith(".eml"):
                with open(path, 'rb') as f:
                    msg = BytesParser(policy=policy.default).parse(f)
                    subject = msg["subject"]
                    sender = msg["from"]
                    to = msg["to"]
                    body = msg.get_body(preferencelist=('plain', 'html')).get_content()
                    content = f"From: {sender}\nTo: {to}\nSubject: {subject}\n\n{body}"
            elif path.endswith(".msg"):
                if not extract_msg:
                    messagebox.showerror("Missing Module", "extract_msg not found. Run: pip install extract-msg")
                    return
                if not os.path.exists("./temp_eml"):
                    os.makedirs("./temp_eml")
                msg = extract_msg.Message(path)
                content = f"From: {msg.sender}\nSubject: {msg.subject}\n\n{msg.body}"
                for att in msg.attachments:
                    att.save(customPath="./temp_eml")
                    self.loaded_attachments.append(os.path.join("./temp_eml", att.longFilename))
            else:
                content = "Unsupported format."

        except Exception as e:
            content = f"Error reading email file: {str(e)}"

        self.email_textbox.delete(1.0, tk.END)
        self.email_textbox.insert(tk.END, content)
        self.email_result_text.delete(1.0, tk.END)
        self.email_result_text.insert(tk.END, "[Loaded] File content inserted.\n")

        if self.auto_analyze_var.get():
            self.analyze_email_with_openai()
            self.scan_email_links()
            self.scan_email_ips()
            self.scan_email_attachments()

    def analyze_email_with_openai(self):
        text = self.email_textbox.get(1.0, tk.END).strip()
        if not text:
            messagebox.showwarning("Error", "Load an email first.")
            return

        self.email_result_text.insert(tk.END, "[OpenAI] Sending email content for analysis...\n")
        self.log_activity("Анализ на съдържание на email с OpenAI GPT")
        self.root.update()

        try:
            api_key = self.config.get_api_key("OPENAI_API_KEY")
            model_name = self.selected_gpt_model.get()
            if not api_key or not api_key.startswith("sk-"):
                raise Exception("Missing or invalid OpenAI key (must start with sk-)")

            from openai import OpenAI  # local import to avoid global issues
            client = OpenAI(api_key=api_key)  # without http_client

            prompt = f"""
    You are an email threat detection AI. Classify the following email by estimating the likelihood
    (0–100%) that it belongs to each of the following categories:

    - Legitimate
    - Marketing
    - Spam
    - Phishing
    - Malware

    Explain briefly why you gave each rating.

    EMAIL CONTENT:
    {text[:3000]}
    """

            response = client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": "You are a cybersecurity threat analyst."},
                    {"role": "user", "content": prompt}
                ]
            )

            result = response.choices[0].message.content.strip()
            self.last_openai_result = result
            self.email_result_text.insert(tk.END, f"[OpenAI Result ({model_name})]\n{result}\n\n")

        except Exception as e:
            self.last_openai_result = f"OpenAI error: {str(e)}"
            self.email_result_text.insert(tk.END, self.last_openai_result + "\n")


    def analyze_email_with_free_gpt(self):
        text = self.email_textbox.get(1.0, tk.END).strip()
        if not text:
            messagebox.showwarning("Error", "Load an email first.")
            return

        self.email_result_text.insert(tk.END, "[Free GPT] Sending email content to free GPT endpoint...\n")
        self.log_activity("Анализ на съдържание на email с Free GPT")
        self.root.update()

        try:
            model_name = self.selected_gpt_model.get()
            messages = [
                {"role": "system", "content": "You are a cybersecurity threat detection assistant."},
                {"role": "user", "content": f"""
    Classify the following email by estimating the likelihood (0–100%) for:

    - Legitimate
    - Marketing
    - Spam
    - Phishing
    - Malware

    Explain briefly why. EMAIL:
    {text[:3000]}
    """}
            ]

            proxies = self.config.get_proxy_settings()
            auth = HttpNegotiateAuth()

            response = requests.post(
                "https://gpt-api.puter.com/completion",
                json={"model": model_name, "messages": messages},
                proxies=proxies if proxies else None,
                auth=auth,
                timeout=15
            )

            if response.status_code != 200 or not response.text.strip().startswith("{"):
                raise Exception(f"Invalid response ({response.status_code}): {response.text[:100]}")

            result = response.json()["choices"][0]["message"]["content"]
            self.email_result_text.insert(tk.END, f"[Free GPT Result ({model_name})]\n{result}\n\n")

        except Exception as e:
            self.email_result_text.insert(tk.END, f"[Free GPT Error] {e}\n")


    def ask_gpt_custom_prompt_free(self):
        question = self.gpt_custom_input.get(1.0, tk.END).strip()
        self.log_activity("Въпрос към Free GPT")
        if not question:
            messagebox.showwarning("GPT Query", "Please enter a question.")
            return

        try:
            model_name = self.selected_gpt_model.get()
            context = ""
            if self.include_email_context.get():
                context = self.email_textbox.get(1.0, tk.END).strip()[:3000]

            messages = [{"role": "system", "content": "You are a helpful security assistant."}]
            if context:
                messages.append({"role": "user", "content": f"Email content:\n{context}"})
            messages.append({"role": "user", "content": question})

            proxies = self.config.get_proxy_settings()
            auth = HttpNegotiateAuth()

            response = requests.post(
                "https://gpt-api.puter.com/completion",
                json={"model": model_name, "messages": messages},
                proxies=proxies if proxies else None,
                auth=auth,
                timeout=15
            )

            if response.status_code != 200 or not response.text.strip().startswith("{"):
                raise Exception(f"Invalid response ({response.status_code}): {response.text[:100]}")

            reply = response.json()["choices"][0]["message"]["content"]

            self.gpt_custom_output.delete(1.0, tk.END)
            self.gpt_custom_output.insert(tk.END, reply)

            self.gpt_qa_history.append((question, reply))
            display_q = question.replace('\n', ' ').strip()
            self.gpt_prompt_history.insert(tk.END, display_q[:80] + ("..." if len(display_q) > 80 else ""))

            # === Рискова оценка (по избор) ===
            risks = {"legit": 0, "marketing": 0, "spam": 0, "phishing": 0, "malware": 0}
            for line in reply.splitlines():
                line = line.lower()
                for key in risks:
                    if key in line:
                        match = re.search(r'(\d{1,3})\s*%?', line)
                        if match:
                            risks[key] = int(match.group(1))

            highest = max(risks.items(), key=lambda x: x[1])
            label = highest[0].capitalize()
            value = highest[1]

            if value >= 80 and label in ["Phishing", "Malware"]:
                color = "red"
            elif value >= 50 and label in ["Spam", "Marketing"]:
                color = "orange"
            elif value <= 30 and label == "Legit":
                color = "green"
            else:
                color = "yellow"

            tag_name = "gpt_risk_tag"
            self.gpt_custom_output.tag_config(tag_name, background=color, foreground="black")
            self.gpt_custom_output.insert(tk.END, f"\n[RISK LEVEL] Most likely: {label} ({value}%)\n", tag_name)

        except Exception as e:
            self.gpt_custom_output.delete(1.0, tk.END)
            self.gpt_custom_output.insert(tk.END, f"[Free GPT Error] {str(e)}")




    def scan_email_links(self):
        text = self.email_textbox.get(1.0, tk.END)
        self.last_urls = re.findall(r"https?://[\w./?=#&%-]+", text)
        self.log_activity("Анализ на съдържание на линковете в email")
        if not self.last_urls:
            self.email_result_text.insert(tk.END, "[Info] No URLs found.\n")
            return
        for url in self.last_urls:
            result = self.security_tools.scan_url_virustotal(url)
            if result.get("error"):
                self.email_result_text.insert(tk.END, f"[VT] {url}: {result['error']}\n")
            else:
                self.email_result_text.insert(tk.END, f"[VT] {url}: {result['detections_str']}\n{result['permalink']}\n")

    def scan_email_ips(self):
        text = self.email_textbox.get(1.0, tk.END)
        self.last_ips = re.findall(r"\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b", text)
        self.log_activity("Проверка на мейл ip адреси")
        if not self.last_ips:
            self.email_result_text.insert(tk.END, "[Info] No IPs found.\n")
            return
        for ip in self.last_ips:
            result = self.security_tools.check_abuseipdb(ip)
            if result.get("error"):
                self.email_result_text.insert(tk.END, f"[AbuseIPDB] {ip}: {result['error']}\n")
            else:
                self.email_result_text.insert(tk.END, f"[AbuseIPDB] {ip}: Score {result['abuseConfidenceScore']}, Reports: {result['totalReports']}\n")

    def scan_email_attachments(self):
        if not self.loaded_attachments:
            self.email_result_text.insert(tk.END, "[Info] No attachments found.\n")
            return
        self.last_attachment_results.clear()
        self.log_activity("Сканиране на мейл прикачен файл")
        for file_path in self.loaded_attachments:
            result = self.security_tools.scan_file_virustotal(file_path)
            if result.get("error"):
                self.email_result_text.insert(tk.END, f"[VT] {file_path}: {result['error']}\n")
            else:
                self.last_attachment_results.append((file_path, result))
                self.email_result_text.insert(tk.END, f"[VT] {file_path}: {result['detections_str']}\n{result['permalink']}\n")

    def get_threat_summary_text(self):
        summary = "--- EMAIL THREAT SUMMARY ---\n"
        summary += self.last_openai_result.strip() + "\n\n"
        if self.last_urls:
            summary += "URLs:\n" + "\n".join(self.last_urls) + "\n\n"
        if self.last_ips:
            summary += "IP Addresses:\n" + "\n".join(self.last_ips) + "\n\n"
        if self.last_attachment_results:
            summary += "Attachments:\n"
            for fpath, res in self.last_attachment_results:
                summary += f"{os.path.basename(fpath)}: {res['detections_str']}\n{res['permalink']}\n"
        return summary.strip()

    def copy_threat_summary(self):
        if not pyperclip:
            messagebox.showerror("Missing Module", "pyperclip not found. Run: pip install pyperclip")
            return
        pyperclip.copy(self.get_threat_summary_text())
        messagebox.showinfo("Copied", "Threat summary copied to clipboard.")

    def save_threat_summary(self):
        summary = self.get_threat_summary_text()
        filepath = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt")])
        if filepath:
            try:
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(summary)
                messagebox.showinfo("Saved", f"Summary saved to {filepath}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save summary: {e}")

    def ask_gpt_custom_prompt(self):
        question = self.gpt_custom_input.get(1.0, tk.END).strip()
        self.log_activity("Въпрос към OpenAI GPT")
        if not question:
            messagebox.showwarning("GPT Query", "Please enter a question or prompt.")
            return

        try:
            api_key = self.config.get_api_key("OPENAI_API_KEY")
            if not api_key or not api_key.startswith("sk-"):
                messagebox.showerror("OpenAI API", "Invalid or missing API key. Must start with 'sk-'.")
                return

            from openai import OpenAI
            client = OpenAI(api_key=api_key)

            context = ""
            if self.include_email_context.get():
                context = self.email_textbox.get(1.0, tk.END).strip()[:3000]

            messages = [{"role": "system", "content": "You are a cybersecurity and phishing analysis assistant."}]
            if context:
                messages.append({"role": "user", "content": f"Email content:\n{context}"})
            messages.append({"role": "user", "content": question})

            model_name = self.selected_gpt_model.get()
            response = client.chat.completions.create(
                model=model_name,
                messages=messages
            )
            # === DEBUG START ===
            print("[DEBUG] Sending OpenAI request...")
            print("[DEBUG] API Key:", "SET" if api_key else "MISSING")
            print("[DEBUG] Model:", model)
            print("[DEBUG] Proxies:", proxies)
            print("[DEBUG] Headers:", headers)
            print("[DEBUG] Payload:", json.dumps(payload)[:300], "...")
            # === DEBUG END ===
            reply = response.choices[0].message.content.strip()

            self.gpt_custom_output.delete(1.0, tk.END)
            self.gpt_custom_output.insert(tk.END, reply + "\n\n")

            # История
            self.gpt_qa_history.append((question, reply))
            display_q = question.replace('\n', ' ').strip()
            self.gpt_prompt_history.insert(tk.END, display_q[:80] + ("..." if len(display_q) > 80 else ""))

            # === Рискова оценка ===
            risks = {"legit": 0, "marketing": 0, "spam": 0, "phishing": 0, "malware": 0}
            for line in reply.splitlines():
                line = line.lower()
                for key in risks.keys():
                    if key in line:
                        match = re.search(r'(\d{1,3})\s*%?', line)
                        if match:
                            risks[key] = int(match.group(1))

            highest = max(risks.items(), key=lambda x: x[1])
            label = highest[0].capitalize()
            value = highest[1]

            if value >= 80 and label in ["Phishing", "Malware"]:
                color = "red"
            elif value >= 50 and label in ["Spam", "Marketing"]:
                color = "orange"
            elif value <= 30 and label == "Legit":
                color = "green"
            else:
                color = "yellow"

            tag_name = "gpt_risk_tag"
            self.gpt_custom_output.tag_config(tag_name, background=color, foreground="black")
            self.gpt_custom_output.insert(tk.END, f"[RISK LEVEL] Most likely: {label} ({value}%)\n", tag_name)

        except Exception as e:
            self.gpt_custom_output.delete(1.0, tk.END)
            self.gpt_custom_output.insert(tk.END, f"GPT error: {str(e)}")



    def load_gpt_history_entry(self, event):
        selection = event.widget.curselection()
        if not selection:
            return
        index = selection[0]
        _, answer = self.gpt_qa_history[index]
        self.gpt_custom_output.delete(1.0, tk.END)
        self.gpt_custom_output.insert(tk.END, answer)

    def save_gpt_response(self):
        text = self.gpt_custom_output.get(1.0, tk.END).strip()
        if not text:
            messagebox.showinfo("Info", "No GPT response to save.")
            return
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt")])
        if path:
            try:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(text)
                messagebox.showinfo("Saved", f"Saved to {path}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save: {e}")


    def open_file(self):
        """Отваря файл с резултати"""
        file_path = filedialog.askopenfilename(
            filetypes=[("Текстови файлове", ".txt"), ("Всички файлове", ".*")]
        )
        if file_path:
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()

                # Показване на съдържанието в подходящото текстово поле
                current_tab = self.notebook.tab(self.notebook.select(), "text")
                if current_tab == "IP Анализ":
                    self.ip_text_results.delete(1.0, tk.END)
                    self.ip_text_results.insert(tk.END, content)
                elif current_tab == "URL Анализ":
                    self.url_text_results.delete(1.0, tk.END)
                    self.url_text_results.insert(tk.END, content)
                elif current_tab == "Файлов Анализ":
                    self.file_text_results.delete(1.0, tk.END)
                    self.file_text_results.insert(tk.END, content)

                self.status_bar.config(
                    text=f"Файлът {os.path.basename(file_path)} е зареден"
                )
            except Exception as e:
                messagebox.showerror("Грешка", f"Грешка при отваряне на файл: {str(e)}")
                self.status_bar.config(text="Грешка при отваряне на файл")


    def save_results(self):
        """Запазва резултатите във файл"""
        current_tab = self.notebook.tab(self.notebook.select(), "text")
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Текстови файлове", ".txt"), ("CSV файлове", ".csv")],
        )
        if not file_path:
            return

        try:
            if current_tab == "Пакетна IP Проверка" and file_path.endswith(".csv"):
                self._save_batch_results_to_csv(file_path)
            else:
                with open(file_path, "w", encoding="utf-8") as f:
                    if current_tab == "IP Анализ":
                        content = self.ip_text_results.get(1.0, tk.END)
                    elif current_tab == "URL Анализ":
                        content = self.url_text_results.get(1.0, tk.END)
                    elif current_tab == "Файлов Анализ":
                        content = self.file_text_results.get(1.0, tk.END)
                    elif current_tab == "Пакетна IP Проверка":
                        content = "\n".join(
                            f"{row[0]}\t{row[1]}\t{row[2]}\t{row[3]}\t{row[4]}"
                            for row in self.batch_ip_treeview.get_children()
                        )
                    else:
                        content = ""

                    f.write(content)

            self.status_bar.config(
                text=f"Резултатите са запазени във {os.path.basename(file_path)}"
            )
        except Exception as e:
            messagebox.showerror("Грешка", f"Грешка при запис на файл: {str(e)}")
            self.status_bar.config(text="Грешка при запазване на файл")

    def _save_batch_results_to_csv(self, file_path):
        """Запазва резултатите от пакетната проверка във CSV файл"""
        with open(file_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["IP", "Държава", "Доставчик", "Репут.", "Детайли"])
            for row in self.batch_ip_treeview.get_children():
                values = self.batch_ip_treeview.item(row)["values"]
                writer.writerow(values)

    def open_settings(self):
        settings_win = tk.Toplevel(self.root)
        settings_win.title("Настройки")
        settings_win.geometry("600x500")
        settings_win.resizable(True, True)

        main_frame = ttk.Frame(settings_win)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # ТЕМА
        theme_frame = ttk.LabelFrame(main_frame, text="Тема")
        theme_frame.pack(fill=tk.X, pady=5)
        self.theme_var = tk.StringVar(value=self.config.config.get("SETTINGS", "THEME", fallback="dark"))
        for theme_key in self.themes:
            theme_name = theme_key.capitalize()
            ttk.Radiobutton(theme_frame, text=theme_name, variable=self.theme_var, value=theme_key).pack(anchor="w", padx=10)

        # ЕЗИК
        lang_frame = ttk.LabelFrame(main_frame, text="Език")
        lang_frame.pack(fill=tk.X, pady=5)
        self.language_var = tk.StringVar(value=self.config.config.get("SETTINGS", "LANGUAGE", fallback="bg"))
        language_names = {"bg": "Български", "en": "Английски"}
        for lang_key in self.languages:
            lang_label = language_names.get(lang_key, lang_key)
            ttk.Radiobutton(lang_frame, text=lang_label, variable=self.language_var, value=lang_key).pack(anchor="w", padx=10)


        # ПРОКСИ
        proxy_frame = ttk.LabelFrame(main_frame, text="Прокси")
        proxy_frame.pack(fill=tk.X, pady=5)
        self.proxy_enabled = tk.BooleanVar(value=self.config.config.getboolean("SETTINGS", "PROXY_ENABLED", fallback=False))
        ttk.Checkbutton(proxy_frame, text="Активирай прокси", variable=self.proxy_enabled).pack(anchor="w", padx=10, pady=5)

        ttk.Label(proxy_frame, text="HTTP прокси:").pack(anchor="w", padx=10)
        self.http_proxy = tk.StringVar(value=self.config.config.get("PROXY", "HTTP_PROXY", fallback=""))
        ttk.Entry(proxy_frame, textvariable=self.http_proxy).pack(fill=tk.X, padx=10)

        ttk.Label(proxy_frame, text="HTTPS прокси:").pack(anchor="w", padx=10)
        self.https_proxy = tk.StringVar(value=self.config.config.get("PROXY", "HTTPS_PROXY", fallback=""))
        ttk.Entry(proxy_frame, textvariable=self.https_proxy).pack(fill=tk.X, padx=10)

        # БУТОНИ
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=10)
        ttk.Button(btn_frame, text="Отказ", command=settings_win.destroy).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frame, text="Запази", command=lambda: self.save_settings(settings_win)).pack(side=tk.RIGHT, padx=5)

        
    def save_settings(self, settings_window):
        """Запазва настройките и рестартира приложението"""
        import sys
        import subprocess
        try:
            self.config.config.set("SETTINGS", "THEME", self.theme_var.get())
            self.config.config.set("SETTINGS", "LANGUAGE", self.language_var.get())
            self.config.config.set("SETTINGS", "PROXY_ENABLED", str(self.proxy_enabled.get()))
            self.config.config.set("PROXY", "HTTP_PROXY", self.http_proxy.get())
            self.config.config.set("PROXY", "HTTPS_PROXY", self.https_proxy.get())

            with open(CONFIG_FILE, "w") as configfile:
                self.config.config.write(configfile)

            messagebox.showinfo("Настройки", "Настройките са запазени. Приложението ще се рестартира за да приложи промените.")
            settings_window.destroy()
            subprocess.Popen([sys.executable] + sys.argv)
            self.root.destroy()
        except Exception as e:
            messagebox.showerror("Грешка", f"Неуспешно запазване на настройки: {str(e)}")
        except Exception as e:
            messagebox.showerror("Грешка", f"Неуспешно запазване на настройки: {str(e)}")


    def open_api_settings(self):
        """Отваря прозорец за управление на API ключове"""
        api_window = tk.Toplevel(self.root)
        api_window.title("Управление на API ключове")
        api_window.geometry("800x700")

        # Създаване на Notebook за различните API ключове
        notebook = ttk.Notebook(api_window)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # VirusTotal API
        vt_frame = ttk.Frame(notebook)
        notebook.add(vt_frame, text="VirusTotal")
        ttk.Label(vt_frame, text="VirusTotal API ключ:").pack(pady=5)
        self.vt_api_entry = ttk.Entry(vt_frame, width=50)
        self.vt_api_entry.pack(pady=5)
        self.vt_api_entry.insert(0, self.config.get_api_key("VIRUSTOTAL_API_KEY"))

        # Hybrid Analysis API
        ha_frame = ttk.Frame(notebook)
        notebook.add(ha_frame, text="Hybrid Analysis")
        ttk.Label(ha_frame, text="Hybrid Analysis API ключ:").pack(pady=5)
        self.ha_api_entry = ttk.Entry(ha_frame, width=50)
        self.ha_api_entry.pack(pady=5)
        self.ha_api_entry.insert(
            0, self.config.get_api_key("HYBRIDANALYSIS_API_KEY")
        )

        # AbuseIPDB API
        abuse_frame = ttk.Frame(notebook)
        notebook.add(abuse_frame, text="AbuseIPDB")
        ttk.Label(abuse_frame, text="AbuseIPDB API ключ:").pack(pady=5)
        self.abuse_api_entry = ttk.Entry(abuse_frame, width=50)
        self.abuse_api_entry.pack(pady=5)
        self.abuse_api_entry.insert(0, self.config.get_api_key("ABUSEIP_API_KEY"))

        # Shodan API
        shodan_frame = ttk.Frame(notebook)
        notebook.add(shodan_frame, text="Shodan")
        ttk.Label(shodan_frame, text="Shodan API ключ:").pack(pady=5)
        self.shodan_api_entry = ttk.Entry(shodan_frame, width=50)
        self.shodan_api_entry.pack(pady=5)
        self.shodan_api_entry.insert(0, self.config.get_api_key("SHODAN_API_KEY"))

        # MXToolbox API
        mxtb_frame = ttk.Frame(notebook)
        notebook.add(mxtb_frame, text="MXToolbox")
        ttk.Label(mxtb_frame, text="MXToolbox API ключ:").pack(pady=5)
        self.mxtb_api_entry = ttk.Entry(mxtb_frame, width=50)
        self.mxtb_api_entry.pack(pady=5)
        self.mxtb_api_entry.insert(0, self.config.get_api_key("MXTOOLBOX_API_KEY"))

        ttk.Label(mxtb_frame, text="OpenAI API ключ:").pack(pady=5)
        self.opnai_api_entry = ttk.Entry(mxtb_frame, width=50)
        self.opnai_api_entry.pack(pady=5)
        self.opnai_api_entry.insert(0, self.config.get_api_key("OPENAI_API_KEY"))


        # URLSCAN API
        urlsc_frame = ttk.Frame(notebook)
        notebook.add(urlsc_frame, text="URL Scan")
        ttk.Label(urlsc_frame, text="URL Scan API ключ:").pack(pady=5)
        self.urlsc_api_entry = ttk.Entry(urlsc_frame, width=50)
        self.urlsc_api_entry.pack(pady=5)
        self.urlsc_api_entry.insert(0, self.config.get_api_key("URLSCAN_API_KEY"))

        # IPGeolocation
        ipgeo_frame = ttk.Frame(notebook)
        notebook.add(ipgeo_frame, text="IP Geolocation")
        ttk.Label(ipgeo_frame, text="Ip Geolocation API ключ:").pack(pady=5)
        self.ipgeo_api_entry = ttk.Entry(ipgeo_frame, width=50)
        self.ipgeo_api_entry.pack(pady=5)
        self.ipgeo_api_entry.insert(0, self.config.get_api_key("GEO_API_KEY"))
        
        # IPInfo
        ipinfo_frame = ttk.Frame(notebook)
        notebook.add(ipinfo_frame, text="IPinfo")
        ttk.Label(ipinfo_frame, text="IPinfo API ключ:").pack(pady=5)
        self.ipinfo_api_entry = ttk.Entry(ipinfo_frame, width=50)
        self.ipinfo_api_entry.pack(pady=5)
        self.ipinfo_api_entry.insert(0, self.config.get_api_key("IPINFO_API_KEY"))
        
        # FindIP
        findip_frame = ttk.Frame(notebook)
        notebook.add(findip_frame, text="FindIP")
        ttk.Label(findip_frame, text="FindIP API ключ:").pack(pady=5)
        self.findip_api_entry = ttk.Entry(findip_frame, width=50)
        self.findip_api_entry.pack(pady=5)
        self.findip_api_entry.insert(0, self.config.get_api_key("FINDIP_API_KEY"))
        
        # IP2Location
        ip2loc_frame = ttk.Frame(notebook)
        notebook.add(ip2loc_frame, text="IP2Location")
        ttk.Label(ip2loc_frame, text="IP2Location API ключ:").pack(pady=5)
        self.ip2location_api_entry = ttk.Entry(ip2loc_frame, width=50)
        self.ip2location_api_entry.pack(pady=5)
        self.ip2location_api_entry.insert(0, self.config.get_api_key("IP2LOCATION_API_KEY"))

        # MaxMind
        maxmind_frame = ttk.Frame(notebook)
        notebook.add(maxmind_frame, text="MaxMind")
        ttk.Label(maxmind_frame, text="MaxMind Account ID:").pack(pady=5)
        self.maxmind_id_entry = ttk.Entry(maxmind_frame, width=50)
        self.maxmind_id_entry.pack(pady=5)
        self.maxmind_id_entry.insert(0, self.config.get_api_key("MAXMIND_ACCOUNT_ID"))

        ttk.Label(maxmind_frame, text="MaxMind License Key:").pack(pady=5)
        self.maxmind_license_entry = ttk.Entry(maxmind_frame, width=50)
        self.maxmind_license_entry.pack(pady=5)
        self.maxmind_license_entry.insert(0, self.config.get_api_key("MAXMIND_LICENSE_KEY"))
        
        # Phishtank
        phishtank_frame = ttk.Frame(notebook)
        notebook.add(phishtank_frame, text="Phishing")
        ttk.Label(phishtank_frame, text="Phishtank API ключ:").pack(pady=5)
        self.phishtank_api_entry = ttk.Entry(phishtank_frame, width=50)
        self.phishtank_api_entry.pack(pady=5)
        self.phishtank_api_entry.insert(0, self.config.get_api_key("PHISHTANK_API_KEY"))

        # OSINT API
        osint_frame = ttk.Frame(notebook)
        notebook.add(osint_frame, text="OSINT")

        ttk.Label(osint_frame, text="Hunter.io API ключ:").pack(pady=2)
        self.hunter_api_entry = ttk.Entry(osint_frame, width=50)
        self.hunter_api_entry.pack()
        self.hunter_api_entry.insert(0, self.config.get_osint_key("HUNTER_IO"))

        ttk.Label(osint_frame, text="IntelligenceX API ключ:").pack(pady=2)
        self.intelx_api_entry = ttk.Entry(osint_frame, width=50)
        self.intelx_api_entry.pack()
        self.intelx_api_entry.insert(0, self.config.get_osint_key("INTELLIGENCE_X"))

        ttk.Label(osint_frame, text="GitHub Token:").pack(pady=2)
        self.github_api_entry = ttk.Entry(osint_frame, width=50)
        self.github_api_entry.pack()
        self.github_api_entry.insert(0, self.config.get_osint_key("GITHUB_TOKEN"))

        ttk.Label(osint_frame, text="Censys ID:").pack(pady=2)
        self.censys_id_entry = ttk.Entry(osint_frame, width=50)
        self.censys_id_entry.pack()
        self.censys_id_entry.insert(0, self.config.get_osint_key("CENSYS_ID"))

        ttk.Label(osint_frame, text="Censys Secret:").pack(pady=2)
        self.censys_secret_entry = ttk.Entry(osint_frame, width=50)
        self.censys_secret_entry.pack()
        self.censys_secret_entry.insert(0, self.config.get_osint_key("CENSYS_SECRET"))

        ttk.Label(osint_frame, text="Onyphe API ключ:").pack(pady=2)
        self.onyphe_api_entry = ttk.Entry(osint_frame, width=50)
        self.onyphe_api_entry.pack()
        self.onyphe_api_entry.insert(0, self.config.get_osint_key("ONYPHE_KEY"))

        ttk.Label(osint_frame, text="Използвай IntelligenceX? (True/False):").pack(pady=2)
        self.use_intelx_entry = ttk.Entry(osint_frame, width=50)
        self.use_intelx_entry.pack()
        self.use_intelx_entry.insert(0, self.config.get_osint_key("USE_INTELLIGENCE_X", fallback="True"))
        
        # BreachDirectory API
        ttk.Label(osint_frame, text="BreachDirectory API ключ:").pack(pady=5)
        self.breachdirectory_key_entry = ttk.Entry(osint_frame, width=50)
        self.breachdirectory_key_entry.pack(pady=5)
        self.breachdirectory_key_entry.insert(0, self.config.get_osint_key("BREACHDIRECTORY_KEY", fallback=""))




        # Бутони за запазване и отказ
        btn_frame = ttk.Frame(api_window)
        btn_frame.pack(pady=10)

        ttk.Button(
            btn_frame,
            text="Запази",
            command=lambda: self.save_api_settings(api_window),
        ).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Отказ", command=api_window.destroy).pack(
            side=tk.LEFT, padx=5
        )



    def save_api_settings(self, api_window):
        """Запазва API настройките"""
        try:
            self.config.config.set("API", "VIRUSTOTAL_API_KEY", self.vt_api_entry.get())
            self.config.config.set("API", "HYBRIDANALYSIS_API_KEY", self.ha_api_entry.get())
            self.config.config.set("API", "ABUSEIP_API_KEY", self.abuse_api_entry.get())
            self.config.config.set("API", "SHODAN_API_KEY", self.shodan_api_entry.get())
            self.config.config.set("API", "MXTOOLBOX_API_KEY", self.mxtb_api_entry.get())
            self.config.config.set("API", "OPENAI_API_KEY", self.opnai_api_entry.get())
            self.config.config.set("API", "URLSCAN_API_KEY", self.urlsc_api_entry.get())
            self.config.config.set("API", "GEO_API_KEY", self.ipgeo_api_entry.get())
            self.config.config.set("API", "IPINFO_API_KEY", self.ipinfo_api_entry.get())
            self.config.config.set("API", "FINDIP_API_KEY", self.findip_api_entry.get())
            self.config.config.set("API", "IP2LOCATION_API_KEY", self.ip2location_api_entry.get())
            self.config.config.set("API", "MAXMIND_ACCOUNT_ID", self.maxmind_id_entry.get())
            self.config.config.set("API", "MAXMIND_LICENSE_KEY", self.maxmind_license_entry.get())
            self.config.config.set("API", "PHISHTANK_API_KEY", self.phishtank_api_entry.get())

            self.config.config.set("OSINT_API_KEYS", "HUNTER_IO", self.hunter_api_entry.get())
            self.config.config.set("OSINT_API_KEYS", "INTELLIGENCE_X", self.intelx_api_entry.get())
            self.config.config.set("OSINT_API_KEYS", "GITHUB_TOKEN", self.github_api_entry.get())
            self.config.config.set("OSINT_API_KEYS", "CENSYS_ID", self.censys_id_entry.get())
            self.config.config.set("OSINT_API_KEYS", "CENSYS_SECRET", self.censys_secret_entry.get())
            self.config.config.set("OSINT_API_KEYS", "ONYPHE_KEY", self.onyphe_api_entry.get())
            self.config.config.set("OSINT_API_KEYS", "USE_INTELLIGENCE_X", self.use_intelx_entry.get())
            self.config.config.set("OSINT_API_KEYS", "BREACHDIRECTORY_KEY", self.breachdirectory_key_entry.get())




            with open(CONFIG_FILE, "w") as configfile:
                self.config.config.write(configfile)

            messagebox.showinfo("Успех", "API ключовете са запазени успешно.")
            api_window.destroy()
        except Exception as e:
            messagebox.showerror(
                "Грешка", f"Грешка при запазване на API ключовете:\n{str(e)}"
            )

    def open_email_settings(self):
        email_win = tk.Toplevel(self.root)
        email_win.title("Настройки на имейл")
        email_win.geometry("450x400")
        email_win.resizable(True, True)

        main_frame = ttk.Frame(email_win)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        fields = [
            ("SMTP сървър:", "SMTP_SERVER"),
            ("SMTP порт:", "SMTP_PORT"),
            ("Потребителско име:", "EMAIL_USER"),
            ("Парола:", "EMAIL_PASSWORD"),
            ("Получател по подразбиране:", "DEFAULT_RECIPIENT"),
            ("Hybrid Analysis Email:", "ANALYSIS_NOTIFICATION_EMAIL")
        ]

        self.email_entries = {}

        for label, key in fields:
            ttk.Label(main_frame, text=label).pack(anchor="w", pady=2)
            entry = ttk.Entry(main_frame, show="*" if "Парола" in label else None)
            entry.insert(0, self.config.config.get("EMAIL", key, fallback=""))
            entry.pack(fill=tk.X, padx=10, pady=2)
            self.email_entries[key] = entry

        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(pady=10)

        ttk.Button(btn_frame, text="Запази", command=lambda: self.save_email_settings(email_win)).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Тест", command=self.test_email_settings).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Отказ", command=email_win.destroy).pack(side=tk.LEFT, padx=5)

    def save_email_settings(self, email_window):
        """Запазва имейл настройките"""
        try:
            for key in self.email_entries:
                value = self.email_entries[key].get()
                self.config.config.set("EMAIL", key, value)

            with open(CONFIG_FILE, "w") as configfile:
                self.config.config.write(configfile)

            messagebox.showinfo("Успех", "Имейл настройките са запазени успешно.")
            email_window.destroy()
        except Exception as e:
            messagebox.showerror("Грешка", f"Грешка при запазване на имейл настройките:{str(e)}")
        except Exception as e:
            messagebox.showerror("Грешка", f"Грешка при запазване на имейл настройките:\n{str(e)}")


    def test_email_settings(self):
        """Тества имейл настройките"""
        try:
            smtp_server = self.smtp_server_entry.get()
            smtp_port = int(self.smtp_port_entry.get())
            email_user = self.email_user_entry.get()
            email_pass = self.email_pass_entry.get()
            recipient = self.default_recipient_entry.get() or email_user

            msg = MIMEMultipart()
            msg["From"] = email_user
            msg["To"] = recipient
            msg["Subject"] = "Тестов имейл от SOC Tool"
            body = "Това е тестов имейл за проверка на настройките."
            msg.attach(MIMEText(body, "plain"))

            with smtplib.SMTP(smtp_server, smtp_port) as server:
                server.starttls()
                server.login(email_user, email_pass)
                server.send_message(msg)

            messagebox.showinfo("Успех", "Тестовият имейл е изпратен успешно!")
        except Exception as e:
            messagebox.showerror(
                "Грешка", f"Неуспешно изпращане на тестови имейл:\n{str(e)}"
            )

    def test_connection(self):
        """Тества връзката с интернет"""
        try:
            response = requests.get("https://www.google.com", proxies=self.config.get_proxy_settings(), timeout=10)
            if response.status_code == 200:
                messagebox.showinfo(
                    "Тест на връзката", "Връзката с интернет е налична"
                )
            else:
                messagebox.showwarning(
                    "Тест на връзката",
                    "Връзката с интернет е налична, но има проблеми",
                )
        except Exception as e:
            messagebox.showerror(
                "Тест на връзката", f"Грешка при тестване на връзката:\n{str(e)}"
            )

    def show_documentation(self):
        """Показва документация"""
        doc_window = tk.Toplevel(self.root)
        doc_window.title("Документация")
        doc_window.geometry("800x600")
        doc_text = scrolledtext.ScrolledText(doc_window, wrap=tk.WORD)
        doc_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        documentation = """
        Документация за SOC Tool

        1. IP Анализ
        -Проверете един IP адрес с AbuseIPDB, VirusTotal, Shodan, GeoIP и други.

        2. URL Анализ
        - Въведете URL за проверка в VirusTotal и UrlScan 
        - Получавате информация за зловредност и свързани ресурси

        3. Файлов Анализ
        - Изберете файл за анализ с VirusTotal или Hybrid Analysis
        - Получавате информация за хеш, зловредност и поведение

        4. Пакетна IP Проверка
        - Въведете множество IP адреси (по един на ред или разделени със запетая)
        - Проверява ги с избраните услуги и показва резултатите в таблица
        - Асинхронна проверка на до 10 000 адреса
        - Използвайте контекстното меню за копиране, изтриване или докладване.

        5. Фишинг Доклади
        - Генерирайте шаблони за доклади за фишинг
        - Докладвайте директно към различни услуги

        6. Мрежови Инструменти
        - Ping, traceroute, Port scanner и други инструменти за мрежова диагностика

        7. Threat Intelligence
        - Проверка на IOC (Indicators of Compromise)
        - Анализ на домейни, IP адреси, хешове и URL-и

        8. MX Toolbox
        - Проверка на DNS записи, blacklist статус и други мейл свързани проверки

        9. PowerShell таб:
        - Изпълнявай команди локално или чрез AD. Поддържа .ps1 drag&drop.
        
        10.  Active Directory таб 
        – Справки по потребители, групи и lockouts
        

        11. OSINT таб за търсене в Hunter.io, BreachDirectory, Shodan, Censys, Onyphe и други
        
        12. Email анализ - Зареждане на .eml/.msg, анализ на мейли, линкове, прикачени файлове
    
        11. Допълнителни опции:
        - Настройки за SMTP/Outlook, API ключове, теми, език и прокси
        - Dashboard с история на действията
        - Drag & drop анализ
        - Имейл нотификации при анализ
        - SSPI базиран достъп до AD без парола
        - Кеширане и multithreading  
        - Конфигурация: config.ini
        
        Разработено от Мартин Стефанов
        """

        doc_text.insert(tk.END, documentation)

    def show_about(self):
        """Показва информация за програмата"""
        about_window = tk.Toplevel(self.root)
        about_window.title("Относно")
        about_window.geometry("600x600")

        about_text = f"""
        SOC Tool v2.1

        Разработено от екипа по киберсигурност на банка за вътрешно ползване.
        Приложението предоставя цялостен GUI инструментариум за анализ на IP адреси, URL-и, файлове, домейни и подозрителна активност в мрежата
        
        Основни функции:

        Анализ на IP адреси чрез AbuseIPDB, VirusTotal, Shodan, ip-api, ipgeolocation и др.
        Асинхронна и многопоточна проверка на големи списъци с IP адреси и опция за докладване
        
        Сканиране на файлове и URL-и чрез VirusTotal, Hybrid Analysis и URL Scan
        
        Фишинг таб с анализ и генериране на репорти и докладване
        Вграден имейл модул с SMTP и Outlook поддръжка
         
        Мрежови инструменти (ping, traceroute, dns resolve).
        Многопоточен портов скенер и traceroute с proxy поддръжка

        PowerShell таб с изпълнение на скриптове и често използвани команди
        
        Търсене и справки в Active Directory без логин (SSPI/Kerberos)

        OSINT таб за търсене в Hunter.io, BreachDirectory, Shodan, Censys, Onyphe и други
        
        Email анализ - Зареждане на .eml/.msg, анализ на мейли, линкове, прикачени файлове
        
        Поддръжка на теми, езикови настройки, API ключове, кеширане и прокси

        Създател: Martin Stefanov
        Версия: 2.1
        Дата: {datetime.now().strftime("%Y-%m-%d")}

        """

        ttk.Label(about_window, text=about_text, justify=tk.LEFT).pack(
            padx=20, pady=20
        )
        ttk.Button(about_window, text="Затвори", command=about_window.destroy).pack(
            pady=10
        )




    
    def navigate_ps_history_up(self, event):
        if self.ps_history and self.ps_history_index > 0:
            self.ps_history_index -= 1
            self.ps_entry.delete("1.0", tk.END)
            self.ps_entry.insert("1.0", self.ps_history[self.ps_history_index])
        return "break"



    
    
    def navigate_ps_history_down(self, event):
        if self.ps_history and self.ps_history_index < len(self.ps_history) - 1:
            self.ps_history_index += 1
            self.ps_entry.delete("1.0", tk.END)
            self.ps_entry.insert("1.0", self.ps_history[self.ps_history_index])
        elif self.ps_history_index == len(self.ps_history) - 1:
            self.ps_entry.delete("1.0", tk.END)
        return "break"



if __name__ == "__main__":
    # Инициализация на Tkinter с поддръжка за drag-and-drop
    # Инициализация с подобрена обработка на грешки
    try:
        tkdnd_path = get_tkdnd_path()
        os.environ['TKDND_LIBRARY'] = tkdnd_path
        root = TkinterDnD.Tk()
        root.title("SOC Tool v2.1")
        # Зареждане на икона
        # Зареждане на икона
        try:
            icon_data = base64.b64decode(
                ConfigManager().config.get("UI", "ICON_BASE64", fallback="")
            )
            with open(ICON_FILE, "wb") as icon_file:
                icon_file.write(icon_data)
            root.iconbitmap(ICON_FILE)
            os.remove(ICON_FILE)
        except Exception as e:
            print(f"Грешка при зареждане на икона: {e}")

        # Настройки на прозореца
        root.geometry("1024x600+100+50")
        root.minsize(1024, 600)
        root.maxsize(2700, 1600)
        root.resizable(True, True)
        
        #screen_width = root.winfo_screenwidth()
        #screen_height = root.winfo_screenheight()
        #scale_factor = min(screen_width / 1920, screen_height / 1080)
        #root.tk.call('tk', 'scaling', scale_factor)


        # Създаване и стартиране на приложението
        app = SOCGUI(root)
    
        # Стартиране на фонова нишка за мониторинг
        threading.Thread(target=app.background_tasks, daemon=True).start()
    
        root.mainloop()
    
    except Exception as e:
        print(f"Критична грешка: {e}")
        messagebox.showerror("Критична грешка", 
            f"Приложението срещна критична грешка:\n{str(e)}\n\nМоля, свържете се с поддръжката.")


